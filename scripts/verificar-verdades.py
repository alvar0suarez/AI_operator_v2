#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
verificar-verdades.py — test de regresión del curso "CX + IA".

Reconstruye las cinco verdades escondidas **sin leer `dataset/SOLUCIONES/`**,
solo desde los ficheros que se le entregan a la alumna, y las compara con los
porcentajes de `dataset/ESPECIFICACION-DATASET.md` con tolerancias declaradas.

Por qué existe (ESPECIFICACION.md §4 y §5.3): como los datos los generamos
nosotros, existe respuesta correcta y se puede corregir sin profesor. Si una
verdad deja de ser derivable desde los ficheros publicados, el ejercicio central
del bloque 4 (`b4-m10-analisis-completo`) se queda sin solución. Este script es
la garantía de que eso no pasa en silencio.

Qué hace, en orden:

  V1  Recalcula todas las líneas de `pedidos.xlsx` con el cálculo correcto y
      comprueba las invariantes del fallo de redondeo. Reclasifica los 800
      tickets desde `descripcion` —como tendría que hacerlo ella— y comprueba
      que la causa pesa el 38 %.
  V2  Normaliza teléfono y dirección de `clientes.xlsx` y busca los 12 pares.
  V3  Agrupa pedidos por (cliente, fecha, producto, cantidad), localiza los 62
      duplicados de diciembre y recalcula el pico con y sin ellos.
  V4  Cuenta tickets por cliente, calcula margen bruto y coste de contacto de los
      tres primeros y comprueba que salen en negativo.
  V5  Separa "entrega sin aviso" de "entrega con retraso", comprueba que pesa el
      9 %, que son todos de hostelería, y que PR-07 está en el manual y no se
      menciona ni una sola vez en tickets ni en correos.

Además comprueba la **suciedad contratada** (cuotas de la especificación), la
**coherencia cruzada** (todo `id_pedido` de un ticket existe, es del mismo
cliente y de fecha compatible; todo remitente de correo está en el maestro) y,
con `--reproducibilidad`, que dos ejecuciones del generador producen el mismo
contenido: los .xlsx y el .docx llevan marcas de tiempo dentro del zip, así que
se comparan **las celdas y el texto**, no los bytes.

Uso:
    python3 scripts/verificar-verdades.py [--rapido] [--reproducibilidad]
                                          [--dataset RUTA] [--json]

Salida: 0 si todo pasa, 1 si algo falla. Dependencias: openpyxl, python-docx y
biblioteca estándar. Python 3.11.
"""

from __future__ import annotations

import argparse
import email.parser
import email.policy
import hashlib
import json
import re
import shutil
import subprocess
import sys
import tempfile
import unicodedata
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

try:
    from openpyxl import load_workbook
except ImportError:  # pragma: no cover
    sys.stderr.write("Falta openpyxl: pip install openpyxl\n")
    raise SystemExit(2)

try:
    import docx
except ImportError:  # pragma: no cover
    sys.stderr.write("Falta python-docx: pip install python-docx\n")
    raise SystemExit(2)


RAIZ = Path(__file__).resolve().parent.parent
GENERADOR = RAIZ / "scripts" / "generar-dataset.py"


# ═══════════════════════════════════════════════════════════════════════════════
# Contrato: las cifras vienen de dataset/ESPECIFICACION-DATASET.md, nunca de
# dataset/SOLUCIONES/. Cada una con su tolerancia declarada.
# ═══════════════════════════════════════════════════════════════════════════════

ESPERADO = {
    "clientes_filas": 300,
    "clientes_reales": 288,
    "duplicados": 12,
    "duplicados_dto_distinto": 5,
    "tickets_filas": 800,
    "correos": 200,
    "pedidos_fantasma": 62,
    "devoluciones": 45,
    "hosteleria": 110,
    "particular": 190,
    "dias_laborables": 129,
    "procedimientos": 9,
    # Porcentajes de las verdades, en tanto por uno.
    "v1_tickets_pct": 0.380,
    "v1_tickets_con_pedido_pct": 0.860,
    "v1_lineas_dto_con_desviacion_pct_min": 0.850,
    "v3_subida_real": 0.120,
    "v3_subida_registrada": 0.410,
    "v4_tickets_pct": 0.220,
    "v5_tickets_pct": 0.090,
}

# Tolerancias. Se declaran aquí, a la vista, para que se pueda discutir si una
# es demasiado laxa. Las que son "exacto" no admiten desviación ninguna.
TOLERANCIA = {
    "v1_tickets_pct": 0.020,              # ±2 puntos
    "v1_tickets_con_pedido_pct": 0.030,   # ±3 puntos
    "v3_subida_real": 0.030,              # ±3 puntos
    "v3_subida_registrada": 0.030,        # ±3 puntos
    "v4_tickets_pct": 0.015,              # ±1,5 puntos
    "v5_tickets_pct": 0.015,              # ±1,5 puntos
    "clasificador_ruido": 0.06,           # ≤6 % de tickets sin clasificar
}

# Cuotas de suciedad de la especificación, en tanto por uno, con ±2,5 puntos.
SUCIEDAD = {
    "clientes_sin_email": (0.18, 0.025),
    "clientes_email_sin_arroba": (0.02, 0.020),
    "clientes_tipo_vacio": (0.03, 0.025),
    "clientes_cp_numerico": (0.05, 0.025),
    "clientes_ruta_vacia": (0.02, 0.020),
    "clientes_dto_vacio": (0.04, 0.025),
    "clientes_observaciones_vacias": (0.70, 0.060),
    "tickets_sin_cliente": (0.04, 0.020),
    "tickets_sin_pedido": (0.52, 0.040),
    "tickets_sin_subcategoria": (0.78, 0.040),
    "tickets_sin_agente": (0.06, 0.025),
    "tickets_sin_tiempo": (0.22, 0.040),
    "tickets_categoria_otros": (0.11, 0.030),
    "tickets_categoria_vacia": (0.05, 0.025),
}

COSTE_CONTACTO = Decimal("11.00")
CENT = Decimal("0.01")

# Márgenes brutos por producto. Están en ESPECIFICACION-DATASET.md (que no es
# SOLUCIONES) y la alumna los recibe en el nodo b4-m4.
MARGENES = {
    "AG-05": Decimal("0.22"), "AG-15": Decimal("0.24"), "AG-GA": Decimal("0.26"),
    "GF-19": Decimal("0.31"), "RE-05": Decimal("0.15"), "RE-NA": Decimal("0.15"),
    "CE-33": Decimal("0.13"), "ZU-20": Decimal("0.19"), "FU-AL": Decimal("0.65"),
    "PO-10": Decimal("0.10"),
}

VENTANA_V3 = (date(2024, 12, 9), date(2024, 12, 13))
MES_PICO = (2024, 12)


# ═══════════════════════════════════════════════════════════════════════════════
# Utilidades
# ═══════════════════════════════════════════════════════════════════════════════

def quitar_acentos(texto: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFD", texto)
                   if unicodedata.category(c) != "Mn")


def arreglar_mojibake(texto: str) -> str:
    if "Ã" not in texto and "Â" not in texto:
        return texto
    try:
        return texto.encode("latin-1").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return texto


ABREVIATURAS_VIA = (
    (r"\bavda\b", "avenida"), (r"\bavd\b", "avenida"), (r"\bav\b", "avenida"),
    (r"\bc\b", "calle"), (r"\bcl\b", "calle"), (r"\bcalle\b", "calle"),
    (r"\bpza\b", "plaza"), (r"\bpl\b", "plaza"),
    (r"\bp\b", "paseo"), (r"\bpso\b", "paseo"), (r"\bpaseo\b", "paseo"),
    (r"\btrav\b", "travesia"), (r"\bbo\b", "barrio"), (r"\bbº\b", "barrio"),
    (r"\bnum\b", ""), (r"\bn\b", ""), (r"\bs/n\b", ""),
)


def normalizar_direccion(direccion: str) -> str:
    texto = quitar_acentos(arreglar_mojibake(str(direccion))).lower()
    texto = texto.replace("/", " ").replace("º", " ").replace("ª", " ")
    texto = re.sub(r"[^a-z0-9\s]", " ", texto)
    texto = re.sub(r"\s+", " ", texto).strip()
    for patron, reemplazo in ABREVIATURAS_VIA:
        texto = re.sub(patron, reemplazo, texto)
    return re.sub(r"\s+", " ", texto).strip()


def normalizar_telefono(telefono) -> str:
    digitos = re.sub(r"\D", "", str(telefono or ""))
    if len(digitos) == 11 and digitos.startswith("34"):
        digitos = digitos[2:]
    if len(digitos) == 12 and digitos.startswith("0034"):
        digitos = digitos[4:]
    return digitos


MESES_ES = ("ene", "feb", "mar", "abr", "may", "jun",
            "jul", "ago", "sep", "oct", "nov", "dic")


def leer_fecha(valor) -> date | None:
    """Los tres formatos de `tickets.xlsx` más el ISO de `pedidos.xlsx`."""
    if valor is None or valor == "":
        return None
    if isinstance(valor, datetime):
        return valor.date()
    if isinstance(valor, date):
        return valor
    texto = str(valor).strip()
    m = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", texto)
    if m:
        return date(int(m[1]), int(m[2]), int(m[3]))
    m = re.fullmatch(r"(\d{1,2})/(\d{1,2})/(\d{4})", texto)
    if m:
        return date(int(m[3]), int(m[2]), int(m[1]))
    m = re.fullmatch(r"(\d{1,2})-([a-zA-Z]{3})-(\d{4})", texto)
    if m and m[2].lower() in MESES_ES:
        return date(int(m[3]), MESES_ES.index(m[2].lower()) + 1, int(m[1]))
    return None


def leer_hoja(ruta: Path) -> tuple[list[str], list[dict]]:
    libro = load_workbook(ruta, data_only=True)
    hoja = libro.active
    filas = list(hoja.iter_rows(values_only=True))
    cabecera = [str(c) for c in filas[0]]
    registros = [dict(zip(cabecera, fila)) for fila in filas[1:]]
    libro.close()
    return cabecera, registros


def dec(valor) -> Decimal:
    return Decimal(str(valor)).quantize(CENT, rounding=ROUND_HALF_UP)


def euros(valor: Decimal) -> str:
    texto = f"{valor.quantize(CENT, rounding=ROUND_HALF_UP):,.2f}"
    entero, decimales = texto.split(".")
    return f"{entero.replace(',', '.')},{decimales} €"


def pct(valor: float) -> str:
    return f"{100 * valor:.2f}".replace(".", ",") + " %"


# ═══════════════════════════════════════════════════════════════════════════════
# El reclasificador. Es lo que ella tiene que construir en el bloque 4: leer la
# `descripcion` y decidir la categoría, porque la columna `categoria` no sirve.
# Aquí se hace con reglas de palabra clave, que es exactamente el nivel al que se
# llega a mano; si el dataset solo fuera clasificable con un modelo grande, el
# ejercicio no sería resoluble por ella.
# ═══════════════════════════════════════════════════════════════════════════════

MARCAS = {
    # Facturación que NO viene del fallo: se comprueba primero porque comparte
    # vocabulario con la de abajo y es la trampa deliberada de la especificación.
    "facturacion-otro": (
        "dos veces la misma factura", "duplicad", "el iva", "iva de la factura",
        "no es el que acordamos", "que yo no he pedido nunca",
        "nombre del antiguo dueno", "ya pague en efectivo",
        "numero de factura", "me la habeis pasado dos veces",
    ),
    "facturacion-redondeo": (
        "no me cuadra", "no cuadra", "descuadre", "no da lo mismo",
        "de mas en la", "cobrado", "descuento no esta bien", "la cuenta no sale",
        "no coincide", "de diferencia", "una diferencia", "no sale de sumar",
        "no entiendo la factura", "no es el que me habeis facturado",
        "mas caro de lo que me toca", "calculadora", "no me da vuestro total",
        "el importe no es ese", "el importe esta mal", "arriba siempre",
        "el descuento que me aplicais", "he sumado el albaran",
    ),
    "entrega-sin-aviso": (
        "avis", "estaba cerrado", "esta cerrado", "no sabia",
        "perdido la entrega", "nadie me dijo nada", "no habia nadie",
        "ya se habia ido",
    ),
    "entrega-retraso": (
        "retras", "tarde", "tardisimo", "no llega", "no ha llegado",
        "sin llegar", "sigo esperando", "estoy esperando", "llevo esperando",
        "dias de retraso",
    ),
    "producto-defectuoso": (
        "rota", "rotas", "caducad", "sabe raro", "mal estado", "estropead",
        "sucia", "precinto", "aplastadas", "pierde agua", "fecha pasada",
        "perdiendo agua", "cajas mojadas",
    ),
    "pedido-erroneo": (
        "yo pedi", "os pedi", "no he pedido", "no las he pedido", "el doble de lo que pedi",
        "falta", "faltan", "incomplet", "de otro cliente", "no he hecho",
        "no corresponde", "sobran", "el albaran pone",
    ),
    "cambio-datos": (
        "cambiar la cuenta", "cambiado de direccion", "cambiar el telefono",
        "cambiar los datos", "a otro correo", "cambiado de numero",
        "cuenta bancaria", "actualiz", "nombre fiscal", "cambiad el nombre",
        "domicili", "apuntad",
    ),
    "informacion-producto": (
        "quiero saber", "cuanto cuesta", "que precio tiene", "teneis cerveza",
        "en otro formato", "que descuento hariais", "me podeis decir el precio",
        "si servis", "es retornable", "que dias pasais", "haceis reparto a",
        "cuanto tardais", "catalogo con todos los precios",
    ),
    "otros": (
        "darme de baja", "certificad", "catalogo actualizado", "vacaciones",
        "copia del contrato", "horario teneis", "muy amable", "no me traigais",
        "punto de entrega", "abrimos por la tarde", "hablar con",
    ),
}
# Desempate cuando dos categorías puntúan igual. El orden es deliberado: lo más
# específico primero.
PRIORIDAD = (
    "facturacion-otro", "producto-defectuoso", "entrega-sin-aviso",
    "entrega-retraso", "facturacion-redondeo", "cambio-datos", "pedido-erroneo",
    "informacion-producto", "otros",
)


def normalizar_texto(texto: str) -> str:
    return re.sub(r"\s+", " ", quitar_acentos(str(texto or "")).lower())


def clasificar(descripcion: str) -> str:
    """Categoría reconstruida desde el texto libre. '' si no hay señal."""
    texto = normalizar_texto(descripcion)
    puntos: dict[str, int] = {}
    for categoria, marcas in MARCAS.items():
        aciertos = sum(1 for marca in marcas if marca in texto)
        if aciertos:
            puntos[categoria] = aciertos
    if not puntos:
        return ""
    mejor = max(puntos.values())
    empatadas = [c for c, p in puntos.items() if p == mejor]
    for candidata in PRIORIDAD:
        if candidata in empatadas:
            return candidata
    return empatadas[0]


# ═══════════════════════════════════════════════════════════════════════════════
# Informe
# ═══════════════════════════════════════════════════════════════════════════════

class Informe:
    def __init__(self) -> None:
        self.fallos: list[str] = []
        self.avisos: list[str] = []
        self.datos: dict = {}
        self._bloque = ""

    def bloque(self, titulo: str) -> None:
        self._bloque = titulo
        print()
        print(titulo)
        print("─" * len(titulo))

    def ok(self, etiqueta: str, detalle: str = "") -> None:
        print(f"  [ok]    {etiqueta}" + (f"  ·  {detalle}" if detalle else ""))

    def fallo(self, etiqueta: str, detalle: str = "") -> None:
        self.fallos.append(f"{self._bloque} :: {etiqueta} — {detalle}")
        print(f"  [FALLO] {etiqueta}" + (f"  ·  {detalle}" if detalle else ""))

    def aviso(self, etiqueta: str, detalle: str = "") -> None:
        self.avisos.append(f"{self._bloque} :: {etiqueta} — {detalle}")
        print(f"  [aviso] {etiqueta}" + (f"  ·  {detalle}" if detalle else ""))

    def exacto(self, etiqueta: str, obtenido, esperado) -> None:
        if obtenido == esperado:
            self.ok(etiqueta, f"{obtenido}")
        else:
            self.fallo(etiqueta, f"obtenido {obtenido}, esperado {esperado}")

    def cerca(self, etiqueta: str, obtenido: float, esperado: float,
              tolerancia: float) -> None:
        diferencia = abs(obtenido - esperado)
        detalle = (f"{pct(obtenido)} frente a {pct(esperado)} "
                   f"(desvío {abs(100 * (obtenido - esperado)):.2f} p., "
                   f"tolerancia ±{100 * tolerancia:.1f} p.)")
        if diferencia <= tolerancia:
            self.ok(etiqueta, detalle)
        else:
            self.fallo(etiqueta, detalle)

    def condicion(self, etiqueta: str, condicion: bool, detalle: str = "") -> None:
        (self.ok if condicion else self.fallo)(etiqueta, detalle)


# ═══════════════════════════════════════════════════════════════════════════════
# Carga del dataset
# ═══════════════════════════════════════════════════════════════════════════════

class Dataset:
    def __init__(self, ficheros: Path) -> None:
        self.ruta = ficheros
        _, self.clientes = leer_hoja(ficheros / "clientes.xlsx")
        _, self.lineas = leer_hoja(ficheros / "pedidos.xlsx")
        _, self.tickets = leer_hoja(ficheros / "tickets.xlsx")

        self.por_cliente = {c["id_cliente"]: c for c in self.clientes}

        # Un pedido es el conjunto de sus líneas.
        self.pedidos: dict[str, dict] = {}
        for linea in self.lineas:
            pedido = self.pedidos.setdefault(linea["id_pedido"], {
                "id_pedido": linea["id_pedido"],
                "fecha": leer_fecha(linea["fecha"]),
                "id_cliente": linea["id_cliente"],
                "lineas": [],
            })
            pedido["lineas"].append(linea)

        analizador = email.parser.BytesParser(policy=email.policy.default)
        self.correos = []
        for ruta in sorted((ficheros / "correos").glob("*.eml")):
            with ruta.open("rb") as fichero:
                self.correos.append((ruta.name, analizador.parse(fichero)))

        documento = docx.Document(str(ficheros / "procedimientos.docx"))
        self.manual = "\n".join(p.text for p in documento.paragraphs)
        for tabla in documento.tables:
            for fila in tabla.rows:
                self.manual += "\n" + " | ".join(c.text for c in fila.cells)

    def texto_correo(self, mensaje) -> str:
        try:
            cuerpo = mensaje.get_content()
        except Exception:  # pragma: no cover
            cuerpo = ""
        return f"{mensaje['Subject'] or ''}\n{cuerpo}"

    def tipo_real(self, id_cliente: str) -> str:
        """`tipo` viene en seis grafías y a veces vacío. Se normaliza."""
        cliente = self.por_cliente.get(id_cliente)
        if not cliente:
            return ""
        crudo = normalizar_texto(cliente.get("tipo") or "")
        if crudo.startswith("host"):
            return "hosteleria"
        if crudo.startswith("part"):
            return "particular"
        return ""


# ═══════════════════════════════════════════════════════════════════════════════
# Comprobaciones
# ═══════════════════════════════════════════════════════════════════════════════

def comprobar_forma(d: Dataset, inf: Informe) -> None:
    inf.bloque("Forma de los ficheros")
    inf.exacto("clientes.xlsx: filas", len(d.clientes), ESPERADO["clientes_filas"])
    inf.exacto("tickets.xlsx: filas", len(d.tickets), ESPERADO["tickets_filas"])
    inf.exacto("correos/: ficheros .eml", len(d.correos), ESPERADO["correos"])
    inf.ok("pedidos.xlsx: líneas", f"{len(d.lineas)}")
    inf.ok("pedidos.xlsx: pedidos", f"{len(d.pedidos)}")

    mbox = (d.ruta / "bandeja.mbox").read_text(encoding="utf-8")
    inf.exacto("bandeja.mbox: mensajes",
               len(re.findall(r"(?m)^From \S+@", mbox)), ESPERADO["correos"])

    codigos = re.findall(r"PR-0[1-9]", d.manual)
    inf.exacto("procedimientos.docx: procedimientos",
               len(sorted(set(codigos))), ESPERADO["procedimientos"])
    inf.condicion("procedimientos.docx: portada v3 de marzo de 2019",
                  "marzo de 2019" in d.manual)

    tipos = [d.tipo_real(c["id_cliente"]) for c in d.clientes]
    # El 3 % con `tipo` vacío se reparte a ojo; se comprueba el orden de magnitud.
    inf.condicion("clientes: reparto hostelería / particular",
                  abs(tipos.count("hosteleria") - ESPERADO["hosteleria"]) <= 12,
                  f"{tipos.count('hosteleria')} hostelería (declarados), "
                  f"{tipos.count('particular')} particulares, "
                  f"{tipos.count('')} sin tipo")


def comprobar_suciedad(d: Dataset, inf: Informe) -> None:
    inf.bloque("Suciedad contratada")
    n_cli, n_tk = len(d.clientes), len(d.tickets)

    medidas = {
        "clientes_sin_email": sum(1 for c in d.clientes if not c["email"]) / n_cli,
        "clientes_email_sin_arroba":
            sum(1 for c in d.clientes if c["email"] and "@" not in str(c["email"])) / n_cli,
        "clientes_tipo_vacio": sum(1 for c in d.clientes if not c["tipo"]) / n_cli,
        "clientes_cp_numerico":
            sum(1 for c in d.clientes if isinstance(c["cp"], (int, float))) / n_cli,
        "clientes_ruta_vacia": sum(1 for c in d.clientes if not c["ruta"]) / n_cli,
        "clientes_dto_vacio":
            sum(1 for c in d.clientes if c["descuento_pct"] in (None, "")) / n_cli,
        "clientes_observaciones_vacias":
            sum(1 for c in d.clientes if not c["observaciones"]) / n_cli,
        "tickets_sin_cliente": sum(1 for t in d.tickets if not t["id_cliente"]) / n_tk,
        "tickets_sin_pedido": sum(1 for t in d.tickets if not t["id_pedido"]) / n_tk,
        "tickets_sin_subcategoria": sum(1 for t in d.tickets if not t["subcategoria"]) / n_tk,
        "tickets_sin_agente": sum(1 for t in d.tickets if not t["agente"]) / n_tk,
        "tickets_sin_tiempo":
            sum(1 for t in d.tickets if t["tiempo_dedicado_min"] in (None, "")) / n_tk,
        "tickets_categoria_otros":
            sum(1 for t in d.tickets if str(t["categoria"] or "").strip() == "Otros") / n_tk,
        "tickets_categoria_vacia": sum(1 for t in d.tickets if not t["categoria"]) / n_tk,
    }
    for clave, (esperado, tolerancia) in SUCIEDAD.items():
        inf.cerca(clave, medidas[clave], esperado, tolerancia)

    # Cuatro formatos de teléfono
    formatos = set()
    for c in d.clientes:
        crudo = str(c["telefono"])
        if crudo.startswith("+"):
            formatos.add("prefijo")
        elif "-" in crudo:
            formatos.add("guiones")
        elif " " in crudo:
            formatos.add("espacios")
        else:
            formatos.add("seguido")
    inf.exacto("clientes: formatos de teléfono distintos", len(formatos), 4)

    # Tres formatos de fecha en tickets
    formatos_fecha = {"iso": 0, "barras": 0, "mes": 0}
    for t in d.tickets:
        crudo = str(t["fecha_apertura"])
        if re.fullmatch(r"\d{4}-\d{2}-\d{2}", crudo):
            formatos_fecha["iso"] += 1
        elif "/" in crudo:
            formatos_fecha["barras"] += 1
        else:
            formatos_fecha["mes"] += 1
    inf.condicion("tickets: tres formatos de fecha conviviendo",
                  all(v > 40 for v in formatos_fecha.values()), str(formatos_fecha))

    # 14 etiquetas sucias + Otros + vacío
    etiquetas = {str(t["categoria"]).strip() for t in d.tickets if t["categoria"]}
    inf.condicion("tickets: 14 etiquetas sucias más `Otros`",
                  len(etiquetas) == 15, f"{len(etiquetas)} etiquetas no vacías")

    # Devoluciones camufladas como negativos
    negativos = {l["id_pedido"] for l in d.lineas if (l["cantidad"] or 0) < 0}
    inf.exacto("pedidos: devoluciones en negativo sin marcar", len(negativos),
               ESPERADO["devoluciones"])
    marcas = [l for l in d.lineas
              if (l["cantidad"] or 0) < 0
              and re.search(r"abono|devoluc", normalizar_texto(l["descripcion"]))]
    inf.condicion("pedidos: las devoluciones no llevan ninguna marca", not marcas)

    # Cierres imposibles y tiempos absurdos
    imposibles = 0
    for t in d.tickets:
        apertura, cierre = leer_fecha(t["fecha_apertura"]), leer_fecha(t["fecha_cierre"])
        if apertura and cierre and cierre < apertura:
            imposibles += 1
    inf.exacto("tickets: cierres anteriores a la apertura", imposibles, 9)
    absurdos = sum(1 for t in d.tickets
                   if isinstance(t["tiempo_dedicado_min"], (int, float))
                   and t["tiempo_dedicado_min"] in (0, 999, 1440))
    inf.exacto("tickets: tiempos absurdos", absurdos, 7)

    # Mojibake
    mojibake = sum(1 for c in d.clientes
                   if "Ã" in str(c["nombre"]) or "Ã" in str(c["direccion"]))
    inf.condicion("clientes: acentos mal codificados", mojibake >= 15,
                  f"{mojibake} fichas con mojibake")


def comprobar_coherencia(d: Dataset, inf: Informe) -> None:
    inf.bloque("Coherencia cruzada")

    rotos, otro_cliente, futuro = [], [], []
    for t in d.tickets:
        if not t["id_pedido"]:
            continue
        pedido = d.pedidos.get(t["id_pedido"])
        if pedido is None:
            rotos.append(t["id_ticket"])
            continue
        if t["id_cliente"] and pedido["id_cliente"] != t["id_cliente"]:
            otro_cliente.append(t["id_ticket"])
        apertura = leer_fecha(t["fecha_apertura"])
        if apertura and pedido["fecha"] and pedido["fecha"] > apertura:
            futuro.append(t["id_ticket"])
    inf.condicion("tickets: todo `id_pedido` existe en pedidos.xlsx",
                  not rotos, f"{len(rotos)} rotos")
    inf.condicion("tickets: el pedido es del mismo cliente",
                  not otro_cliente, f"{len(otro_cliente)} descuadrados")
    inf.condicion("tickets: el pedido es anterior a la incidencia",
                  not futuro, f"{len(futuro)} posteriores")

    huerfanos = [l["id_pedido"] for l in d.lineas
                 if l["id_cliente"] not in d.por_cliente]
    inf.condicion("pedidos: todo `id_cliente` existe en clientes.xlsx",
                  not huerfanos, f"{len(sorted(set(huerfanos)))} huérfanos")

    correos_validos = {str(c["email"]).lower() for c in d.clientes
                       if c["email"] and "@" in str(c["email"])}
    desconocidos = []
    for nombre, mensaje in d.correos:
        remitente = re.search(r"<([^>]+)>", mensaje["From"] or "")
        direccion = (remitente[1] if remitente else (mensaje["From"] or "")).lower()
        if direccion not in correos_validos:
            desconocidos.append(nombre)
    inf.condicion("correos: todo remitente está en clientes.xlsx",
                  not desconocidos, f"{len(desconocidos)} desconocidos")

    # Los .eml parsean y traen las cabeceras que exige la especificación.
    faltan = [n for n, m in d.correos
              if not (m["From"] and m["To"] and m["Subject"] is not None
                      and m["Date"] and m["Message-ID"])]
    inf.condicion("correos: cabeceras RFC 5322 completas", not faltan,
                  f"{len(faltan)} incompletos")

    ids = {str(m["Message-ID"]) for _, m in d.correos}
    rotos_hilo = [n for n, m in d.correos
                  if m["In-Reply-To"] and str(m["In-Reply-To"]) not in ids]
    inf.condicion("correos: hay hilos rotos (suciedad contratada)",
                  len(rotos_hilo) >= 20, f"{len(rotos_hilo)} `In-Reply-To` colgando")

    adjuntos = [n for n, m in d.correos
                if "adjunt" in normalizar_texto(d.texto_correo(m))
                and not list(m.iter_attachments())]
    inf.condicion("correos: adjuntos mencionados que no están",
                  len(adjuntos) >= 20, f"{len(adjuntos)} correos")


def comprobar_v1(d: Dataset, inf: Informe, clasificadas: dict[str, str]) -> dict:
    inf.bloque("V1 — el fallo de facturación en pedidos con descuento")

    sin_dto_con_desvio, con_dto, con_dto_y_desvio, negativas = 0, 0, 0, 0
    desviacion_total = Decimal("0.00")
    peor = Decimal("0.00")
    pedidos_afectados: set[str] = set()
    for linea in d.lineas:
        precio = Decimal(str(linea["precio_ud"]))
        dto = int(linea["dto_pct"] or 0)
        cantidad = int(linea["cantidad"])
        facturado = dec(linea["importe_linea"])
        correcto = (precio * (Decimal(100) - Decimal(dto)) / Decimal(100)
                    * Decimal(cantidad)).quantize(CENT, rounding=ROUND_HALF_UP)
        desvio = facturado - correcto
        if dto == 0:
            if desvio != 0:
                sin_dto_con_desvio += 1
        else:
            con_dto += 1
            if desvio > 0:
                con_dto_y_desvio += 1
            if desvio < 0:
                negativas += 1
        if not str(linea["id_pedido"]).startswith("PED-9"):
            desviacion_total += desvio
            if desvio > 0:
                pedidos_afectados.add(linea["id_pedido"])
        peor = max(peor, desvio)

    inf.condicion("sin descuento no hay desviación (la pista)",
                  sin_dto_con_desvio == 0,
                  f"{sin_dto_con_desvio} líneas con dto_pct = 0 y desviación")
    inf.condicion("con descuento la desviación nunca es negativa",
                  negativas == 0, f"{negativas} líneas negativas")
    ratio = con_dto_y_desvio / con_dto
    inf.condicion("≥85 % de las líneas con descuento sobrefacturan",
                  ratio >= ESPERADO["v1_lineas_dto_con_desviacion_pct_min"],
                  f"{pct(ratio)} de {con_dto} líneas")
    inf.ok("sobrefacturación acumulada (pedidos reales)",
           f"{euros(desviacion_total)} en {len(pedidos_afectados)} pedidos; "
           f"peor línea {euros(peor)}")

    redondeo = [t for t in d.tickets if clasificadas[t["id_ticket"]] == "facturacion-redondeo"]
    inf.cerca("tickets reclasificados como el fallo",
              len(redondeo) / len(d.tickets), ESPERADO["v1_tickets_pct"],
              TOLERANCIA["v1_tickets_pct"])

    con_pedido = [t for t in redondeo if t["id_pedido"]]
    inf.cerca("de ellos, con `id_pedido` informado",
              len(con_pedido) / len(redondeo), ESPERADO["v1_tickets_con_pedido_pct"],
              TOLERANCIA["v1_tickets_con_pedido_pct"])

    malos = []
    for t in con_pedido:
        pedido = d.pedidos[t["id_pedido"]]
        dtos = {int(l["dto_pct"] or 0) for l in pedido["lineas"]}
        if dtos == {0}:
            malos.append(t["id_ticket"])
    inf.condicion("todos apuntan a pedidos con descuento", not malos,
                  f"{len(malos)} apuntan a pedidos sin descuento")

    otros_fact = [t for t in d.tickets
                  if clasificadas[t["id_ticket"]] == "facturacion-otro"]
    inf.condicion("hay incidencias de facturación que NO son el fallo",
                  len(otros_fact) >= 10,
                  f"{len(otros_fact)} tickets (factura duplicada, IVA, precio)")

    sin_nombre = sum(1 for t in d.tickets
                     if "redonde" in normalizar_texto(t["descripcion"]))
    inf.condicion("ningún texto nombra el hallazgo", sin_nombre == 0,
                  f"{sin_nombre} descripciones lo dicen")

    # El puente pedagógico: la columna `categoria` tiene que ser inservible.
    etiquetas_fact = {"Facturación", "facturacion", "FACTURACION", "Facturas",
                      "Incidencia facturación"}
    bucket = [t for t in d.tickets if str(t["categoria"] or "").strip() in etiquetas_fact]
    puros = [t for t in bucket
             if clasificadas[t["id_ticket"]] == "facturacion-redondeo"]
    ingenua = len(bucket) / len(d.tickets)
    error = len(redondeo) / len(d.tickets) - ingenua
    inf.condicion("la columna `categoria` es inservible para dimensionarlo",
                  error >= 0.08,
                  f"agrupando por `categoria` sale {pct(ingenua)} en vez de "
                  f"{pct(len(redondeo) / len(d.tickets))}: "
                  f"{100 * error:.1f} puntos de error, pureza {pct(len(puros) / len(bucket))}")

    return {
        "desviacion_total": str(desviacion_total),
        "pedidos_afectados": len(pedidos_afectados),
        "lineas_con_descuento": con_dto,
        "lineas_con_descuento_y_desviacion": con_dto_y_desvio,
        "tickets": len(redondeo),
        "tickets_pct": len(redondeo) / len(d.tickets),
        "estimacion_ingenua_pct": ingenua,
    }


def comprobar_v2(d: Dataset, inf: Informe) -> dict:
    inf.bloque("V2 — los clientes duplicados")

    grupos: dict[tuple[str, str], list[dict]] = {}
    for c in d.clientes:
        clave = (normalizar_telefono(c["telefono"]), normalizar_direccion(c["direccion"]))
        grupos.setdefault(clave, []).append(c)
    duplicados = {k: v for k, v in sorted(grupos.items()) if len(v) > 1}

    pares = sum(len(v) - 1 for v in duplicados.values())
    inf.exacto("pares detectados por teléfono + dirección normalizados",
               pares, ESPERADO["duplicados"])
    inf.exacto("clientes reales tras deduplicar",
               len(d.clientes) - pares, ESPERADO["clientes_reales"])

    # Ninguno se detecta comparando literalmente.
    literales = {}
    for c in d.clientes:
        literales.setdefault((str(c["telefono"]), str(c["direccion"])), []).append(c)
    a_pelo = sum(len(v) - 1 for v in literales.values() if len(v) > 1)
    inf.condicion("la comparación literal no encuentra ninguno", a_pelo == 0,
                  f"{a_pelo} encontrados a pelo")

    dto_distinto, ambos_con_pedidos = 0, 0
    con_pedidos: dict[str, int] = {}
    for linea in d.lineas:
        con_pedidos[linea["id_cliente"]] = con_pedidos.get(linea["id_cliente"], 0) + 1
    for fichas in duplicados.values():
        descuentos = {int(f["descuento_pct"] or 0) for f in fichas}
        if len(descuentos) > 1:
            dto_distinto += 1
        if all(con_pedidos.get(f["id_cliente"], 0) > 0 for f in fichas):
            ambos_con_pedidos += 1
    inf.exacto("pares con el descuento distinto entre fichas",
               dto_distinto, ESPERADO["duplicados_dto_distinto"])
    inf.exacto("pares en los que ambas fichas tienen pedidos",
               ambos_con_pedidos, ESPERADO["duplicados"])

    nombres_iguales = sum(1 for fichas in duplicados.values()
                          if len({str(f["nombre"]) for f in fichas}) == 1)
    inf.condicion("ningún par comparte el nombre tal cual",
                  nombres_iguales == 0, f"{nombres_iguales} pares con nombre idéntico")

    return {"pares": pares, "pares_dto_distinto": dto_distinto}


def comprobar_v3(d: Dataset, inf: Informe) -> dict:
    inf.bloque("V3 — el pico de diciembre")

    # Agrupación por (cliente, fecha, producto, cantidad): así se ve la copia.
    huellas: dict[tuple, list[str]] = {}
    for linea in d.lineas:
        clave = (linea["id_cliente"], str(linea["fecha"]), linea["producto"],
                 linea["cantidad"])
        huellas.setdefault(clave, []).append(linea["id_pedido"])

    sospechosos: dict[str, int] = {}
    for pedidos_clave in huellas.values():
        unicos = sorted(set(pedidos_clave))
        if len(unicos) > 1:
            for id_pedido in unicos:
                sospechosos[id_pedido] = sospechosos.get(id_pedido, 0) + 1

    # Un pedido es copia si TODAS sus líneas tienen gemela en otro pedido.
    copias = []
    for id_pedido, pedido in sorted(d.pedidos.items()):
        if sospechosos.get(id_pedido, 0) == len(pedido["lineas"]):
            copias.append(id_pedido)
    parejas = {}
    for id_pedido in copias:
        pedido = d.pedidos[id_pedido]
        parejas.setdefault((pedido["id_cliente"], pedido["fecha"],
                            tuple(sorted((l["producto"], l["cantidad"])
                                         for l in pedido["lineas"]))), []).append(id_pedido)
    fantasma = []
    for grupo in parejas.values():
        if len(grupo) == 2:
            # El intruso es el del correlativo aparte.
            fantasma.extend([p for p in grupo if p.startswith("PED-9")]
                            or sorted(grupo)[1:])

    inf.exacto("pedidos duplicados detectados", len(fantasma),
               ESPERADO["pedidos_fantasma"])
    inf.condicion("todos llevan el correlativo aparte `PED-9xxxx`",
                  all(p.startswith("PED-9") for p in fantasma))
    fechas = {d.pedidos[p]["fecha"] for p in fantasma}
    inf.condicion("todos caen en la semana del 9 al 13 de diciembre",
                  all(VENTANA_V3[0] <= f <= VENTANA_V3[1] for f in fechas),
                  f"{len(fechas)} días distintos")

    con_ticket = {t["id_pedido"] for t in d.tickets if t["id_pedido"]}
    inf.condicion("los fantasma no generan ni un ticket",
                  not (set(fantasma) & con_ticket),
                  f"{len(set(fantasma) & con_ticket)} con ticket")

    def lineas_mes(anio: int, mes: int, excluir: set[str]) -> int:
        return sum(1 for l in d.lineas
                   if l["id_pedido"] not in excluir
                   and (f := leer_fecha(l["fecha"])) and (f.year, f.month) == (anio, mes))

    meses = sorted({(f.year, f.month) for l in d.lineas
                    if (f := leer_fecha(l["fecha"]))})
    otros = [m for m in meses if m != MES_PICO]
    base = sum(lineas_mes(*m, set()) for m in otros) / len(otros)
    registrado = lineas_mes(*MES_PICO, set())
    real = lineas_mes(*MES_PICO, set(fantasma))

    inf.cerca("subida real de diciembre (sin los fantasma)",
              real / base - 1, ESPERADO["v3_subida_real"], TOLERANCIA["v3_subida_real"])
    inf.cerca("subida aparente de diciembre (con los fantasma)",
              registrado / base - 1, ESPERADO["v3_subida_registrada"],
              TOLERANCIA["v3_subida_registrada"])
    inf.condicion("el pico aparente es varias veces el real",
                  (registrado / base - 1) / (real / base - 1) >= 2.5,
                  f"{(registrado / base - 1) / (real / base - 1):.1f}×")

    return {
        "fantasma": len(fantasma),
        "base_lineas": round(base, 2),
        "diciembre_real": real,
        "diciembre_registrado": registrado,
        "subida_real": real / base - 1,
        "subida_registrada": registrado / base - 1,
    }


def comprobar_v4(d: Dataset, inf: Informe) -> dict:
    inf.bloque("V4 — los tres clientes que cuestan dinero")

    conteo: dict[str, int] = {}
    for t in d.tickets:
        cliente = t["id_cliente"]
        if not cliente and t["id_pedido"] and t["id_pedido"] in d.pedidos:
            cliente = d.pedidos[t["id_pedido"]]["id_cliente"]   # se recupera por el pedido
        if cliente:
            conteo[cliente] = conteo.get(cliente, 0) + 1
    ranking = sorted(conteo.items(), key=lambda kv: (-kv[1], kv[0]))
    tres = ranking[:3]

    inf.cerca("los tres primeros concentran incidencias",
              sum(n for _, n in tres) / len(d.tickets), ESPERADO["v4_tickets_pct"],
              TOLERANCIA["v4_tickets_pct"])
    inf.condicion("se despegan claramente del cuarto",
                  tres[2][1] >= 2 * ranking[3][1],
                  f"tercero {tres[2][1]} tickets, cuarto {ranking[3][1]}")
    inf.condicion("los tres son de hostelería",
                  all(d.tipo_real(c) == "hosteleria" for c, _ in tres),
                  ", ".join(f"{c}:{d.tipo_real(c) or 'sin tipo'}" for c, _ in tres))
    inf.condicion("los tres tienen descuento alto (10–12)",
                  all(int(d.por_cliente[c]["descuento_pct"] or 0) >= 10 for c, _ in tres),
                  ", ".join(f"{c}:{d.por_cliente[c]['descuento_pct']}%" for c, _ in tres))

    detalle = []
    for id_cliente, n_tickets in tres:
        facturado, margen = Decimal("0.00"), Decimal("0.00")
        for linea in d.lineas:
            if linea["id_cliente"] != id_cliente or str(linea["id_pedido"]).startswith("PED-9"):
                continue
            importe = dec(linea["importe_linea"])
            coste = (Decimal(str(linea["precio_ud"]))
                     * (Decimal(1) - MARGENES[linea["producto"]])
                     * Decimal(int(linea["cantidad"]))).quantize(CENT, rounding=ROUND_HALF_UP)
            facturado += importe
            margen += importe - coste
        coste_contactos = COSTE_CONTACTO * n_tickets
        detalle.append({
            "id_cliente": id_cliente, "tickets": n_tickets,
            "facturado": str(facturado), "margen": str(margen),
            "coste_contactos": str(coste_contactos),
            "resultado": str(margen - coste_contactos),
        })
        inf.condicion(f"{id_cliente}: el margen no paga sus contactos",
                      margen < coste_contactos,
                      f"{n_tickets} tickets · facturado {euros(facturado)} · "
                      f"margen {euros(margen)} · coste {euros(coste_contactos)} · "
                      f"resultado {euros(margen - coste_contactos)}")

    return {"clientes": [c for c, _ in tres],
            "tickets": sum(n for _, n in tres),
            "tickets_pct": sum(n for _, n in tres) / len(d.tickets),
            "detalle": detalle}


def comprobar_v5(d: Dataset, inf: Informe, clasificadas: dict[str, str]) -> dict:
    inf.bloque("V5 — PR-07, el aviso previo que nadie da")

    sin_aviso = [t for t in d.tickets if clasificadas[t["id_ticket"]] == "entrega-sin-aviso"]
    retraso = [t for t in d.tickets if clasificadas[t["id_ticket"]] == "entrega-retraso"]

    inf.cerca("tickets de entrega sin aviso", len(sin_aviso) / len(d.tickets),
              ESPERADO["v5_tickets_pct"], TOLERANCIA["v5_tickets_pct"])
    inf.ok("separados de los de retraso", f"{len(retraso)} tickets de retraso")

    tipos = []
    for t in sin_aviso:
        cliente = t["id_cliente"]
        if not cliente and t["id_pedido"] and t["id_pedido"] in d.pedidos:
            cliente = d.pedidos[t["id_pedido"]]["id_cliente"]
        tipos.append(d.tipo_real(cliente) if cliente else "")
    particulares = tipos.count("particular")
    inf.condicion("ninguno es de un particular", particulares == 0,
                  f"{particulares} particulares, {tipos.count('')} sin tipo declarado")

    inf.condicion("PR-07 existe en el manual y describe el aviso previo",
                  "PR-07" in d.manual and "aviso previo" in normalizar_texto(d.manual))

    menciones_tickets = sum(1 for t in d.tickets
                            if "aviso previo" in normalizar_texto(t["descripcion"]))
    menciones_correos = sum(1 for _, m in d.correos
                            if "aviso previo" in normalizar_texto(d.texto_correo(m)))
    inf.condicion("cero menciones al aviso previo en tickets",
                  menciones_tickets == 0, f"{menciones_tickets}")
    inf.condicion("cero menciones al aviso previo en correos",
                  menciones_correos == 0, f"{menciones_correos}")

    # Contradicciones del manual con los datos: la fecha de 2019 no es casual.
    manual = normalizar_texto(d.manual)
    rutas_reales = {str(c["ruta"]) for c in d.clientes if c["ruta"]}
    inf.condicion("el manual dice tres rutas y en los datos hay cuatro",
                  "tres rutas" in manual and len(rutas_reales) == 4,
                  f"{len(rutas_reales)} rutas en clientes.xlsx")
    inf.condicion("el manual inventa un Departamento de Calidad",
                  "departamento de calidad" in manual)
    horas_reales = sum(1 for t in d.tickets
                       if re.search(r"\blas (7|siete)\b", normalizar_texto(t["descripcion"])))
    inf.condicion("el manual dice reparto de 8:00 a 16:00 y los clientes hablan "
                  "de las 7",
                  "8:00 a 16:00" in manual and horas_reales >= 5,
                  f"{horas_reales} tickets citan la hora real de entrada de ruta")
    inf.condicion("PR-06 describe el cálculo correcto del descuento",
                  "importe" in manual and "nunca" in manual and "PR-06" in d.manual)

    return {"tickets": len(sin_aviso), "tickets_pct": len(sin_aviso) / len(d.tickets),
            "retraso": len(retraso)}


def comprobar_clasificador(d: Dataset, inf: Informe) -> dict[str, str]:
    inf.bloque("Reclasificación desde `descripcion`")
    clasificadas = {t["id_ticket"]: clasificar(t["descripcion"]) for t in d.tickets}
    sin_señal = [k for k, v in clasificadas.items() if not v]
    ruido = len(sin_señal) / len(d.tickets)
    inf.condicion("los textos son clasificables a mano",
                  ruido <= TOLERANCIA["clasificador_ruido"],
                  f"{len(sin_señal)} tickets sin señal ({pct(ruido)}), "
                  f"tolerancia {pct(TOLERANCIA['clasificador_ruido'])}")
    reparto: dict[str, int] = {}
    for valor in clasificadas.values():
        reparto[valor or "(sin señal)"] = reparto.get(valor or "(sin señal)", 0) + 1
    for categoria, n in sorted(reparto.items(), key=lambda kv: -kv[1]):
        print(f"          {categoria:<24} {n:>4}  {pct(n / len(d.tickets)):>8}")
    return clasificadas


# ═══════════════════════════════════════════════════════════════════════════════
# Reproducibilidad
# ═══════════════════════════════════════════════════════════════════════════════

def huella_xlsx(ruta: Path) -> str:
    """
    sha256 del CONTENIDO de las celdas. Los .xlsx son ficheros zip y llevan la
    hora de creación dentro, así que comparar bytes daría siempre distinto: lo
    que tiene que ser idéntico es lo que hay escrito en las celdas.
    """
    libro = load_workbook(ruta, data_only=True)
    resumen = hashlib.sha256()
    for hoja in libro.worksheets:
        resumen.update(hoja.title.encode("utf-8"))
        for fila in hoja.iter_rows(values_only=True):
            resumen.update(repr(fila).encode("utf-8"))
    libro.close()
    return resumen.hexdigest()


def huella_docx(ruta: Path) -> str:
    """Igual que el .xlsx: se compara el texto, no el zip."""
    documento = docx.Document(str(ruta))
    resumen = hashlib.sha256()
    for parrafo in documento.paragraphs:
        resumen.update(parrafo.text.encode("utf-8"))
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                resumen.update(celda.text.encode("utf-8"))
    return resumen.hexdigest()


def huella_texto(ruta: Path) -> str:
    return hashlib.sha256(ruta.read_bytes()).hexdigest()


def huellas_ficheros(ficheros: Path) -> dict[str, str]:
    """Huella de contenido de todo lo que se le entrega a la alumna."""
    salida: dict[str, str] = {}
    for nombre in ("clientes.xlsx", "pedidos.xlsx", "tickets.xlsx"):
        salida[nombre] = huella_xlsx(ficheros / nombre)
    salida["procedimientos.docx"] = huella_docx(ficheros / "procedimientos.docx")
    for nombre in ("bandeja.mbox", "LEEME.md"):
        salida[nombre] = huella_texto(ficheros / nombre)
    for eml in sorted((ficheros / "correos").glob("*.eml")):
        salida[f"correos/{eml.name}"] = huella_texto(eml)
    return salida


def huellas_soluciones(soluciones: Path) -> dict[str, str]:
    """
    Huella de la clave de corrección. Solo se usa sobre directorios recién
    generados en un temporal, nunca sobre `dataset/SOLUCIONES/`: este script no
    abre las soluciones del repositorio ni para mirarlas de reojo.
    """
    return {f"SOLUCIONES/{r.name}": huella_texto(r)
            for r in sorted(soluciones.iterdir()) if r.is_file()}


def _generar_en(destino: Path, inf: Informe) -> bool:
    resultado = subprocess.run(
        [sys.executable, str(GENERADOR), "--salida", str(destino), "--silencioso"],
        capture_output=True, text=True)
    if resultado.returncode != 0:
        inf.fallo("el generador falla al regenerar", resultado.stderr.strip()[:400])
        return False
    return True


def comprobar_reproducibilidad(base: Path, inf: Informe) -> None:
    """
    Dos ejecuciones nuevas del generador en directorios aparte, y las dos
    comparadas entre sí y contra lo que hay versionado. Los .xlsx y el .docx son
    ficheros zip con la hora de creación dentro, así que comparar bytes daría
    siempre distinto: se compara el CONTENIDO (celdas y texto). Los .eml, el
    .mbox y los .csv/.json/.md sí son texto plano y se comparan byte a byte.
    """
    inf.bloque("Reproducibilidad")
    uno = Path(tempfile.mkdtemp(prefix="gemelo-a-"))
    dos = Path(tempfile.mkdtemp(prefix="gemelo-b-"))
    try:
        if not _generar_en(uno, inf) or not _generar_en(dos, inf):
            return

        a = huellas_ficheros(uno / "ficheros") | huellas_soluciones(uno / "SOLUCIONES")
        b = huellas_ficheros(dos / "ficheros") | huellas_soluciones(dos / "SOLUCIONES")

        faltan = sorted(set(a) ^ set(b))
        if faltan:
            inf.fallo("las dos ejecuciones no producen los mismos ficheros",
                      ", ".join(faltan[:6]))
            return
        distintos = sorted(k for k in a if a[k] != b[k])
        inf.condicion("dos ejecuciones nuevas dan contenido idéntico", not distintos,
                      f"{len(a)} ficheros comparados por sha256 de su contenido "
                      f"(incluidos los .csv y el .json de la clave de corrección)"
                      if not distintos else f"difieren: {', '.join(distintos[:8])}")

        versionado = huellas_ficheros(base / "ficheros")
        desfasados = sorted(k for k in versionado if versionado[k] != a.get(k))
        inf.condicion("lo versionado coincide con lo que produce el generador",
                      not desfasados,
                      f"{len(versionado)} ficheros de dataset/ficheros/"
                      if not desfasados
                      else f"desfasados: {', '.join(desfasados[:8])} "
                           f"(regenera con `make dataset`)")

        muestra = ("clientes.xlsx", "pedidos.xlsx", "tickets.xlsx",
                   "procedimientos.docx", "bandeja.mbox", "correos/correo-001.eml",
                   "SOLUCIONES/taxonomia-real.csv", "SOLUCIONES/mapa-duplicados.csv",
                   "SOLUCIONES/pedidos-fantasma.csv", "SOLUCIONES/cuentas-v4.csv")
        for nombre in muestra:
            print(f"          {nombre:<34} {a[nombre][:16]}…")
    finally:
        shutil.rmtree(uno, ignore_errors=True)
        shutil.rmtree(dos, ignore_errors=True)


# ═══════════════════════════════════════════════════════════════════════════════
# Programa principal
# ═══════════════════════════════════════════════════════════════════════════════

def main(argv: list[str] | None = None) -> int:
    analizador = argparse.ArgumentParser(
        description="Comprueba que las cinco verdades escondidas siguen siendo "
                    "derivables desde dataset/ficheros/.")
    analizador.add_argument("--dataset", default=str(RAIZ / "dataset"))
    analizador.add_argument("--rapido", action="store_true",
                            help="Salta la comprobación de reproducibilidad")
    analizador.add_argument("--reproducibilidad", action="store_true",
                            help="Solo la comprobación de reproducibilidad")
    analizador.add_argument("--json", action="store_true",
                            help="Vuelca las cifras reconstruidas en JSON")
    args = analizador.parse_args(argv)

    base = Path(args.dataset)
    if (base / "SOLUCIONES").exists():
        # Recordatorio explícito: este script no abre nada de ahí. Si alguna vez
        # lo hiciera, dejaría de ser un test y pasaría a ser una tautología.
        pass

    inf = Informe()
    print("Verificación del gemelo sintético — Aguas del Norte, S.L.")
    print(f"Dataset: {base / 'ficheros'}")
    print("Fuente de las cifras: dataset/ESPECIFICACION-DATASET.md "
          "(este script NO lee dataset/SOLUCIONES/)")

    if args.reproducibilidad:
        comprobar_reproducibilidad(base, inf)
    else:
        d = Dataset(base / "ficheros")
        comprobar_forma(d, inf)
        comprobar_suciedad(d, inf)
        comprobar_coherencia(d, inf)
        clasificadas = comprobar_clasificador(d, inf)
        inf.datos["v1"] = comprobar_v1(d, inf, clasificadas)
        inf.datos["v2"] = comprobar_v2(d, inf)
        inf.datos["v3"] = comprobar_v3(d, inf)
        inf.datos["v4"] = comprobar_v4(d, inf)
        inf.datos["v5"] = comprobar_v5(d, inf, clasificadas)
        if not args.rapido:
            comprobar_reproducibilidad(base, inf)

    print()
    print("─" * 78)
    if inf.avisos:
        print(f"{len(inf.avisos)} aviso(s):")
        for aviso in inf.avisos:
            print(f"  · {aviso}")
    if inf.fallos:
        print(f"FALLA: {len(inf.fallos)} comprobación(es) no pasan.")
        for fallo in inf.fallos:
            print(f"  · {fallo}")
        print()
        print("Un fallo aquí significa que el ejercicio central del bloque 4 "
              "(b4-m10) ha dejado de tener solución.")
        return 1

    print("Las cinco verdades escondidas siguen siendo derivables desde los "
          "ficheros de la alumna.")
    if args.json:
        print(json.dumps(inf.datos, indent=2, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
