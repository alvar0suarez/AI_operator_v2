#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generar-dataset.py — generador del gemelo sintético "Aguas del Norte, S.L.".

Implementa el contrato de dataset/ESPECIFICACION-DATASET.md, que a su vez deriva
de ESPECIFICACION.md §4. Produce:

    dataset/ficheros/    clientes.xlsx, pedidos.xlsx, tickets.xlsx,
                         correos/*.eml (200), bandeja.mbox, procedimientos.docx,
                         LEEME.md
    dataset/SOLUCIONES/  verdades-escondidas.md, taxonomia-real.csv,
                         mapa-duplicados.csv, pedidos-fantasma.csv,
                         cuentas-v4.csv, metricas-generacion.json

Reproducibilidad (requisito duro de la especificación): toda la aleatoriedad sale
de un único random.Random(SEMILLA). No se usa la hora del sistema, ni uuid, ni
hash(), ni el orden de iteración de un set o de un dict sin ordenar. Dos
ejecuciones producen ficheros idénticos en contenido. Los .xlsx y el .docx llevan
marcas de tiempo internas en el zip, así que la comprobación de reproducibilidad
compara el CONTENIDO (celdas, texto), no los bytes: la hace
scripts/verificar-verdades.py --reproducibilidad.

Dependencias: openpyxl, python-docx y biblioteca estándar. Python 3.11.
Nada de pandas ni numpy: no están instalados y no se quieren como dependencia.

Uso:
    python3 scripts/generar-dataset.py [--salida RUTA] [--silencioso]
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import unicodedata
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal, ROUND_CEILING, ROUND_HALF_UP
from pathlib import Path
import random

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font
except ImportError:  # pragma: no cover
    sys.stderr.write("Falta openpyxl: pip install openpyxl\n")
    raise SystemExit(2)

try:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:  # pragma: no cover
    sys.stderr.write("Falta python-docx: pip install python-docx\n")
    raise SystemExit(2)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. CONSTANTES DEL CONTRATO
#    Todo lo que hay aquí es normativo: sale de dataset/ESPECIFICACION-DATASET.md.
# ═══════════════════════════════════════════════════════════════════════════════

SEMILLA = 20250901

RAIZ = Path(__file__).resolve().parent.parent

FECHA_INICIO = date(2024, 9, 2)
FECHA_FIN = date(2025, 2, 28)
# 130 días entre semana en el periodo; el 25 de diciembre no trabaja nadie.
# 130 − 1 = 129 días laborables, que es la cifra de la especificación.
FESTIVOS = frozenset({date(2024, 12, 25)})
DIAS_LABORABLES_ESPERADOS = 129

N_CLIENTES = 300
N_DUPLICADOS = 12            # V2
N_HOSTELERIA = 110
N_TICKETS = 800
N_CORREOS = 200
N_FANTASMA = 62              # V3
N_DEVOLUCIONES = 45

COSTE_CONTACTO = Decimal("11.00")     # €/contacto, se deriva en el nodo b4-m4
CLIENTES_V4 = ("CLI-0042", "CLI-0118", "CLI-0233")
TICKETS_V4 = {"CLI-0042": 62, "CLI-0118": 58, "CLI-0233": 56}   # 176 = 22,0 %

EMPRESA = "Aguas del Norte, S.L."
DOMINIO_EMPRESA = "aguasdelnorte.es"
BUZON_EMPRESA = "atencioncliente@aguasdelnorte.es"

# Catálogo: código, descripción, precio unitario, unidad, margen bruto.
# El margen NO sale nunca en los ficheros de la alumna: solo en SOLUCIONES/.
CATALOGO = (
    ("AG-05", "Agua mineral 0,5 L (pack 24)", Decimal("0.42"), "botella", Decimal("0.22")),
    ("AG-15", "Agua mineral 1,5 L (pack 6)", Decimal("0.58"), "botella", Decimal("0.24")),
    ("AG-GA", "Agua con gas 1 L (pack 12)", Decimal("0.71"), "botella", Decimal("0.26")),
    ("GF-19", "Garrafa 19 L retornable", Decimal("6.85"), "garrafa", Decimal("0.31")),
    ("RE-05", "Refresco cola 0,33 L (pack 24)", Decimal("0.63"), "lata", Decimal("0.15")),
    ("RE-NA", "Refresco naranja 0,33 L (pack 24)", Decimal("0.63"), "lata", Decimal("0.15")),
    ("CE-33", "Cerveza 0,33 L (pack 24)", Decimal("0.79"), "botella", Decimal("0.13")),
    ("ZU-20", "Zumo 0,2 L (pack 18)", Decimal("0.55"), "brik", Decimal("0.19")),
    ("FU-AL", "Alquiler fuente refrigerada", Decimal("18.00"), "mes", Decimal("0.65")),
    ("PO-10", "Portes fuera de ruta", Decimal("9.50"), "servicio", Decimal("0.10")),
)
PRODUCTOS = {c[0]: {"descripcion": c[1], "precio": c[2], "unidad": c[3], "margen": c[4]}
             for c in CATALOGO}

# Plantilla de la empresa. La inconsistencia de grafía es suciedad deliberada.
PLANTILLA = (
    ("Marta Ibáñez", "Atención al cliente", ("M. Ibáñez", "Marta", "marta", "Marta Ibáñez")),
    ("Rubén Solana", "Atención al cliente", ("R. Solana", "Ruben", "rubén", "Rubén Solana")),
    ("Nieves Palacio", "Administración / facturación", ("N. Palacio", "Nieves", "nieves")),
    ("Chema Ortiz", "Reparto ruta 1 y 2", ("Chema", "Chema Ortiz", "chema")),
    ("Iván Cuadrado", "Reparto ruta 3 y 4", ("Iván", "Ivan", "I. Cuadrado")),
    ("Begoña Salces", "Gerencia", ("Begoña", "B. Salces", "begoña")),
)

# Taxonomía real. Estas cifras son las de la especificación y suman 800.
TAXONOMIA = (
    ("facturacion-redondeo", 304),
    ("entrega-sin-aviso", 72),
    ("entrega-retraso", 96),
    ("producto-defectuoso", 58),
    ("pedido-erroneo", 74),
    ("cambio-datos", 63),
    ("informacion-producto", 89),
    ("otros", 44),
)
N_FACTURACION_NO_REDONDEO = 17   # error de precio, factura duplicada, IVA mal

# ── Presupuesto de líneas de pedido por mes ────────────────────────────────────
# La media de los cinco meses que no son diciembre es 842,4 líneas.
#   diciembre real       = 944   →  +12,1 %  (estacionalidad de verdad)
#   líneas fantasma      = 244   →  la doble importación del 9 al 13 de diciembre
#   diciembre registrado = 1.188 →  +41,0 %  (el pico que parece demanda)
# Son exactamente los porcentajes de la especificación para V3.
PRESUPUESTO_LINEAS = {
    (2024, 9): 868,
    (2024, 10): 902,
    (2024, 11): 856,
    (2024, 12): 944,
    (2025, 1): 790,
    (2025, 2): 796,
}
LINEAS_FANTASMA = 244
VENTANA_FANTASMA_INI = date(2024, 12, 9)
VENTANA_FANTASMA_FIN = date(2024, 12, 13)

# Presupuesto de pedidos, proporcional al de líneas (≈2,84 líneas por pedido).
# Se fija igual que el de líneas para que diciembre suba lo mismo se mida como se
# mida —en pedidos o en líneas— y el hallazgo de V3 no dependa de qué unidad
# elija ella. Total: 1.815 pedidos reales + 62 fantasma ≈ 1.900.
PRESUPUESTO_PEDIDOS = {
    (2024, 9): 306,
    (2024, 10): 318,
    (2024, 11): 301,
    (2024, 12): 332,      # +12,0 % sobre la media de los otros cinco (296,6)
    (2025, 1): 278,
    (2025, 2): 280,
}
PEDIDOS_FANTASMA_LOTE = 62

# Tamaño de pedido en líneas. Distribuciones distintas según el tramo, para que
# los totales de línea y de pedido cuadren a la vez con la especificación
# (~1.900 pedidos y ~5.400 líneas) sin dejar de respetar los porcentajes de V3.
TAMANOS = (1, 2, 3, 4, 5, 6, 7)
PESOS_TAMANO_NORMAL = (0.21, 0.26, 0.23, 0.12, 0.09, 0.06, 0.03)      # media 2,92
PESOS_TAMANO_DICIEMBRE = (0.22, 0.27, 0.23, 0.14, 0.08, 0.05, 0.01)   # media 2,78
PESOS_TAMANO_LOTE = (0.02, 0.08, 0.20, 0.28, 0.22, 0.14, 0.06)        # media ~4,0

# 14 etiquetas sucias de `categoria` para 8 categorías reales, más Otros y vacío.
ETIQUETAS_SUCIAS = (
    "Facturación", "facturacion", "FACTURACION", "Facturas", "Incidencia facturación",
    "Entrega", "entregas", "Reparto",
    "Producto", "Calidad producto",
    "Pedidos", "Datos cliente", "Consulta", "Info",
)
ETIQUETAS_FACTURACION = frozenset({
    "Facturación", "facturacion", "FACTURACION", "Facturas", "Incidencia facturación",
})

# Reparto de cada categoría real sobre las etiquetas sucias. El objetivo
# pedagógico es que la columna `categoria` NO sirva: quien se fíe de ella
# infravalora el problema de facturación en más de diez puntos.
MEZCLA_ETIQUETAS = {
    "facturacion-redondeo": (
        ("Facturación", 0.16), ("facturacion", 0.10), ("FACTURACION", 0.06),
        ("Facturas", 0.12), ("Incidencia facturación", 0.08),
        ("Pedidos", 0.14), ("Consulta", 0.11), ("Entrega", 0.08),
        ("Otros", 0.09), ("", 0.06),
    ),
    "pedido-erroneo": (
        ("Pedidos", 0.52), ("Facturas", 0.06), ("Facturación", 0.06),
        ("Entrega", 0.14), ("Consulta", 0.10), ("Otros", 0.09), ("", 0.03),
    ),
    "entrega-retraso": (
        ("Entrega", 0.36), ("entregas", 0.22), ("Reparto", 0.22),
        ("Consulta", 0.07), ("Facturación", 0.02), ("Otros", 0.08), ("", 0.03),
    ),
    "entrega-sin-aviso": (
        ("Entrega", 0.31), ("entregas", 0.20), ("Reparto", 0.26),
        ("Consulta", 0.09), ("Pedidos", 0.04), ("Otros", 0.07), ("", 0.03),
    ),
    "producto-defectuoso": (
        ("Producto", 0.38), ("Calidad producto", 0.28), ("Entrega", 0.13),
        ("Consulta", 0.07), ("Otros", 0.11), ("", 0.03),
    ),
    "cambio-datos": (
        ("Datos cliente", 0.52), ("Consulta", 0.20), ("Info", 0.10),
        ("Otros", 0.12), ("", 0.06),
    ),
    "informacion-producto": (
        ("Info", 0.32), ("Consulta", 0.34), ("Producto", 0.16),
        ("Facturación", 0.04), ("Otros", 0.10), ("", 0.04),
    ),
    "otros": (
        ("Otros", 0.42), ("Consulta", 0.20), ("Info", 0.16),
        ("Entrega", 0.06), ("Facturación", 0.04), ("", 0.12),
    ),
}

MUNICIPIOS = (
    ("Santander", "39008", "RUTA-1"),
    ("Torrelavega", "39300", "RUTA-2"),
    ("Castro Urdiales", "39700", "RUTA-3"),
    ("Camargo", "39600", "RUTA-1"),
    ("Piélagos", "39470", "RUTA-2"),
    ("El Astillero", "39610", "RUTA-1"),
    ("Laredo", "39770", "RUTA-3"),
    ("Santoña", "39740", "RUTA-3"),
    ("Reinosa", "39200", "RUTA-4"),
    ("Los Corrales de Buelna", "39400", "RUTA-4"),
    ("San Vicente de la Barquera", "39540", "RUTA-4"),
    ("Comillas", "39520", "RUTA-4"),
    ("Llanes", "33500", "RUTA-3"),
    ("Ribadesella", "33560", "RUTA-3"),
)

CALLES = (
    "C/ Alta", "C/ Burgos", "Avda. de los Castros", "C/ Mayor", "Barrio La Iglesia",
    "C/ Ruamayor", "Plaza del Ayuntamiento", "C/ Cervantes", "Avda. de España",
    "C/ del Sol", "Barrio El Cruce", "C/ La Fuente", "C/ Santa Lucía",
    "C/ Menéndez Pelayo", "Pº de Pereda", "C/ San Fernando", "C/ Vargas",
    "C/ Isabel II", "C/ del Río", "Barrio Somahoz", "C/ General Dávila",
    "C/ La Barquera", "Travesía del Muelle", "C/ Los Escalones", "C/ Peña Herbosa",
)

TIPOS_LOCAL = (
    "Bar", "Restaurante", "Café", "Sidrería", "Mesón", "Cafetería", "Hotel",
    "Pensión", "Asador", "Taberna", "Chiringuito", "Bar Restaurante", "Pizzería",
    "Marisquería", "Cervecería",
)
NOMBRES_LOCAL = (
    "Manolo", "El Puerto", "La Marina", "Casa Pepe", "El Rincón", "La Bodega",
    "Los Arcos", "El Faro", "La Plaza", "El Cantábrico", "La Peña", "Santa Ana",
    "El Molino", "La Ría", "Casa Ramón", "El Muelle", "La Escalinata",
    "Los Pinos", "El Sardinero", "La Perla", "Casa Elena", "El Nogal",
    "La Cuadra", "El Mirador", "La Terraza", "Casa Aurora", "El Cruce",
    "La Gaviota", "El Pescador", "Casa Tino", "La Ermita", "El Bosque",
    "La Fuente", "El Cantón", "Casa Julián", "La Palma", "El Cabildo",
    "La Herradura", "El Tonel", "Casa Nieves",
)
NOMBRES_PILA = (
    "María", "Carmen", "Ana", "Isabel", "Laura", "Marta", "Elena", "Rosa",
    "Pilar", "Cristina", "Beatriz", "Lucía", "Sara", "Nuria", "Patricia",
    "José", "Antonio", "Manuel", "Francisco", "Juan", "Luis", "Javier",
    "Miguel", "Carlos", "Ángel", "Jesús", "Pedro", "Fernando", "Alberto",
    "Roberto", "Sergio", "Rubén", "Iván", "Óscar", "Raúl",
)
APELLIDOS = (
    "Gómez", "Fernández", "Solana", "Ibáñez", "Palacio", "Ortiz", "Cuadrado",
    "Salces", "Ruiz", "Cobo", "Setién", "Gutiérrez", "Bolado", "Higuera",
    "Lavín", "Trueba", "Ceballos", "Bustamante", "Quintana", "Diego",
    "Herrera", "Pérez", "Sáiz", "Villegas", "Obregón", "Mazo", "Puente",
    "Rivas", "Carrera", "Escalante", "Noriega", "Posada", "Corral", "Villar",
)
DOMINIOS_CORREO = (
    "gmail.com", "hotmail.com", "yahoo.es", "telefonica.net", "outlook.es",
    "movistar.es",
)

OBSERVACIONES = (
    "Llamar antes de ir", "Cerrado los lunes", "Paga a 30 días",
    "El reparto pasa sobre las 7 y media", "No hay nadie antes de las 9",
    "Entrar por la parte de atrás", "No hay nadie por la tarde",
    "Cliente antiguo", "Ojo con la factura", "Cambió de dueño en 2023",
    "Solo garrafas", "Tiene fuente en alquiler", "Pide albarán siempre",
    "Cerrado en febrero", "Abre solo fines de semana en invierno",
    "Preguntar por Charo", "Descarga complicada, calle estrecha",
)

SUBCATEGORIAS = (
    "revisar", "pendiente", "urgente", "cliente habitual", "2ª vez",
    "ver con Nieves", "hablado con gerencia", "repetido", "ver ruta",
    "pendiente de respuesta", "avisado", "",
)

# Palabras cuyo esqueleto no se corrompe al ensuciar el texto: son las que
# sostienen la reclasificación desde `descripcion`, que es el puente del verbo 1
# al bloque 4. Quien escribe con prisa suele comer tildes y comerse letras en
# los conectores, no en el sustantivo por el que llama.
RAICES_PROTEGIDAS = (
    # facturación
    "factur", "albaran", "cuadr", "descuadr", "descuent", "cobr", "diferenc",
    "importe", "total", "coincid", "suma", "sumad", "sumar", "caro", "toca",
    "calculadora", "arriba", "aplicais", "gestor", "iva", "duplicad", "acordamos",
    # entrega
    "avis", "cerrad", "sabia", "perdido", "retras", "tard", "lleg", "esper",
    "entrega", "reparto", "reparte", "repartidor",
    # producto
    "rot", "caducad", "estropead", "sucia", "sabe", "aplastad", "precinto",
    "pierde", "mojad", "estado",
    # pedido
    "pedido", "falta", "incomplet", "sobran", "corresponde", "doble", "abono",
    "devol",
    # datos de cliente
    "cambi", "actualiz", "domicili", "fiscal", "apuntad", "direccion", "telefon",
    "cuenta", "banco",
    # información y otros
    "precio", "cuesta", "teneis", "retornable", "hariais", "servis", "tardais",
    "pasais", "formato", "catalog", "baja", "certificad", "vacaciones",
    "contrato", "horario", "amable", "traigais",
    # producto y genéricos del negocio
    "garrafa", "agua", "fuente", "alquiler", "gas", "zumo", "cerveza",
)


# ═══════════════════════════════════════════════════════════════════════════════
# 2. UTILIDADES DETERMINISTAS
# ═══════════════════════════════════════════════════════════════════════════════

def quitar_acentos(texto: str) -> str:
    """Deja el texto sin tildes ni diéresis, en su forma base."""
    descompuesto = unicodedata.normalize("NFD", texto)
    return "".join(c for c in descompuesto if unicodedata.category(c) != "Mn")


def arreglar_mojibake(texto: str) -> str:
    """
    Deshace el clásico UTF-8 leído como latin-1 (`JosÃ©` → `José`).
    Si el texto no lo lleva, lo devuelve tal cual.
    """
    if "Ã" not in texto and "Â" not in texto:
        return texto
    try:
        return texto.encode("latin-1").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return texto


def romper_codificacion(texto: str) -> str:
    """Rompe la codificación: UTF-8 mostrado como latin-1. Suciedad deliberada."""
    try:
        return texto.encode("utf-8").decode("latin-1")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return texto


ABREVIATURAS_VIA = (
    (r"\bavda\b", "avenida"), (r"\bavd\b", "avenida"), (r"\bav\b", "avenida"),
    (r"\bc\b", "calle"), (r"\bcl\b", "calle"), (r"\bcalle\b", "calle"),
    (r"\bpza\b", "plaza"), (r"\bpl\b", "plaza"),
    (r"\bp\b", "paseo"), (r"\bpso\b", "paseo"), (r"\bpaseo\b", "paseo"),
    (r"\btrav\b", "travesia"), (r"\bbo\b", "barrio"), (r"\bbº\b", "barrio"),
    (r"\bnum\b", ""), (r"\bn\b", ""), (r"\bs/n\b", ""),
)


def normalizar_direccion(direccion: str) -> str:
    """
    Clave de comparación de direcciones: sin mojibake, sin tildes, en minúsculas,
    sin puntuación y con las abreviaturas de vía desplegadas.
    `C/ Alta 14` y `Calle Alta, 14` dan la misma clave.
    """
    texto = quitar_acentos(arreglar_mojibake(direccion)).lower()
    texto = texto.replace("/", " ").replace("º", " ").replace("ª", " ")
    texto = re.sub(r"[^a-z0-9\s]", " ", texto)
    texto = re.sub(r"\s+", " ", texto).strip()
    for patron, reemplazo in ABREVIATURAS_VIA:
        texto = re.sub(patron, reemplazo, texto)
    return re.sub(r"\s+", " ", texto).strip()


def normalizar_telefono(telefono: str) -> str:
    """Teléfono a sus 9 dígitos, venga en el formato que venga."""
    digitos = re.sub(r"\D", "", str(telefono))
    if len(digitos) == 11 and digitos.startswith("34"):
        digitos = digitos[2:]
    if len(digitos) == 12 and digitos.startswith("0034"):
        digitos = digitos[4:]
    return digitos


def repartir_entero(total: int, pesos: list[float]) -> list[int]:
    """Reparto entero de `total` proporcional a `pesos`, por restos mayores."""
    suma = sum(pesos)
    crudos = [total * p / suma for p in pesos]
    base = [int(c) for c in crudos]
    resto = total - sum(base)
    orden = sorted(range(len(pesos)), key=lambda i: (-(crudos[i] - base[i]), i))
    for i in orden[:resto]:
        base[i] += 1
    return base


def dias_laborables(inicio: date, fin: date) -> list[date]:
    """Días de lunes a viernes del periodo, quitando los festivos declarados."""
    dias: list[date] = []
    actual = inicio
    while actual <= fin:
        if actual.weekday() < 5 and actual not in FESTIVOS:
            dias.append(actual)
        actual += timedelta(days=1)
    return dias


def elegir_pesado(az: random.Random, opciones: tuple, pesos: tuple) -> object:
    """Elección ponderada sobre secuencias ordenadas. Determinista."""
    return az.choices(list(opciones), weights=list(pesos), k=1)[0]


def ajustar_tamanos(tamanos: list[int], objetivo: int) -> list[int]:
    """
    Retoca ±1 los tamaños de pedido hasta que sumen exactamente `objetivo`,
    respetando el rango 1–7 líneas. Recorre en orden, así que es determinista.
    """
    salida = list(tamanos)
    indice = 0
    vueltas = 0
    while sum(salida) != objetivo and vueltas < 200 * max(len(salida), 1):
        paso = 1 if sum(salida) < objetivo else -1
        posicion = indice % len(salida)
        nuevo = salida[posicion] + paso
        if 1 <= nuevo <= 7:
            salida[posicion] = nuevo
        indice += 1
        vueltas += 1
    if sum(salida) != objetivo:  # pragma: no cover
        raise AssertionError(
            f"No se pueden repartir {objetivo} líneas en {len(salida)} pedidos")
    return salida


def muestra_exacta(az: random.Random, poblacion, cuantos: int) -> set:
    """
    Subconjunto de tamaño exacto sobre una población ordenada. Las cuotas de
    suciedad de la especificación son cifras, no probabilidades: si se tiran
    dados por fila, el 18 % sale a veces 21 % y el contrato deja de cumplirse.
    El conjunto devuelto solo se usa para preguntar si algo está dentro, nunca
    para iterar, así que no introduce dependencia del orden de un `set`.
    """
    ordenada = sorted(poblacion)
    return set(az.sample(ordenada, min(cuantos, len(ordenada))))


def formatear_telefono(digitos: str, formato: int) -> str:
    """Los cuatro formatos de teléfono que conviven en el maestro de clientes."""
    if formato == 0:
        return f"{digitos[:3]} {digitos[3:5]} {digitos[5:7]} {digitos[7:9]}"
    if formato == 1:
        return f"+34{digitos}"
    if formato == 2:
        return digitos
    return f"{digitos[:3]}-{digitos[3:5]}-{digitos[5:7]}-{digitos[7:9]}"


MESES_ES = ("ene", "feb", "mar", "abr", "may", "jun",
            "jul", "ago", "sep", "oct", "nov", "dic")
MESES_ES_LARGO = ("enero", "febrero", "marzo", "abril", "mayo", "junio",
                  "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre")


def formatear_fecha(f: date, formato: int) -> str:
    """Los tres formatos de fecha que conviven en tickets.xlsx."""
    if formato == 0:
        return f.isoformat()                       # 2024-11-03
    if formato == 1:
        return f"{f.day:02d}/{f.month:02d}/{f.year}"   # 03/11/2024
    return f"{f.day}-{MESES_ES[f.month - 1]}-{f.year}"  # 3-nov-2024


def desfase_horario(momento: datetime) -> timezone:
    """
    Huso de la España peninsular sin depender de la máquina: CEST hasta el
    último domingo de octubre de 2024, CET a partir de ahí.
    """
    fin_verano = datetime(2024, 10, 27, 3, 0)
    return timezone(timedelta(hours=2 if momento < fin_verano else 1))


DIAS_EN = ("Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun")
MESES_EN = ("Jan", "Feb", "Mar", "Apr", "May", "Jun",
            "Jul", "Aug", "Sep", "Oct", "Nov", "Dec")


def cabecera_fecha(momento: datetime) -> str:
    """Cabecera Date: en RFC 5322, sin depender de la configuración regional."""
    huso = desfase_horario(momento)
    desplazamiento = huso.utcoffset(None)
    signo = "+" if desplazamiento >= timedelta(0) else "-"
    minutos = int(abs(desplazamiento).total_seconds()) // 60
    return (f"{DIAS_EN[momento.weekday()]}, {momento.day:02d} "
            f"{MESES_EN[momento.month - 1]} {momento.year} "
            f"{momento.hour:02d}:{momento.minute:02d}:{momento.second:02d} "
            f"{signo}{minutos // 60:02d}{minutos % 60:02d}")


def linea_from_mbox(momento: datetime) -> str:
    """Línea `From ` separadora del formato mbox, en asctime y sin locale."""
    return (f"{DIAS_EN[momento.weekday()]} {MESES_EN[momento.month - 1]} "
            f"{momento.day:2d} {momento.hour:02d}:{momento.minute:02d}:"
            f"{momento.second:02d} {momento.year}")


# ── Aritmética del dinero. Decimal, nunca float: el fallo de V1 es de céntimos ──

CENT = Decimal("0.01")


def precio_neto_exacto(precio_ud: Decimal, dto_pct: int) -> Decimal:
    return precio_ud * (Decimal(100) - Decimal(dto_pct)) / Decimal(100)


def precio_neto_facturado(precio_ud: Decimal, dto_pct: int) -> Decimal:
    """El fallo: se redondea AL ALZA el precio unitario antes de multiplicar."""
    return precio_neto_exacto(precio_ud, dto_pct).quantize(CENT, rounding=ROUND_CEILING)


def importe_facturado(precio_ud: Decimal, dto_pct: int, cantidad: int) -> Decimal:
    neto = precio_neto_facturado(precio_ud, dto_pct)
    return (neto * Decimal(cantidad)).quantize(CENT, rounding=ROUND_HALF_UP)


def importe_correcto(precio_ud: Decimal, dto_pct: int, cantidad: int) -> Decimal:
    exacto = precio_neto_exacto(precio_ud, dto_pct)
    return (exacto * Decimal(cantidad)).quantize(CENT, rounding=ROUND_HALF_UP)


# ═══════════════════════════════════════════════════════════════════════════════
# 3. GENERADOR DE TEXTO COLOQUIAL
#    Plantillas variadas + relleno variable + erratas. Nada de plantilla repetida
#    literal 300 veces. Ni un solo texto contiene la palabra que nombra el fallo.
# ═══════════════════════════════════════════════════════════════════════════════

APERTURAS = (
    "", "", "", "", "Buenas, ", "Hola, ", "Buenos dias, ", "Otra vez lo mismo, ",
    "A ver, ", "Perdona q moleste pero ", "Mira, ", "Nada, q ", "Os escribo porque ",
    "Llamo porque ", "Ya van varias veces, ",
)
CIERRES = (
    "", "", "", "", "", " A ver si lo mirais.", " Gracias.", " Un saludo.",
    " Decidme algo.", " Espero respuesta.", " Que no se repita por favor.",
    " Gracias y perdona.", " Me llamais cuando podais.", " Un saludo y gracias.",
    " Estoy hasta el gorro la verdad.",
)

PLANTILLAS = {
    "facturacion-redondeo": (
        "la factura de {mes} no me cuadra con el albaran",
        "he sumado el albaran y la factura y no da lo mismo",
        "me habeis cobrado {dif} de mas en la ultima factura",
        "otra vez descuadre en la factura de {mes}",
        "el descuento no esta bien aplicado, tengo un {dto}% y la cuenta no sale",
        "el importe de la factura no coincide con el del albaran",
        "reviso la factura y hay {dif} de diferencia y siempre a vuestro favor",
        "no entiendo la factura, el total no sale de sumar las lineas",
        "el precio con mi descuento no es el que me habeis facturado",
        "en la factura el {prod} sale mas caro de lo que me toca con mi descuento",
        "me sale una diferencia de {dif} entre lo pedido y lo facturado",
        "vuelvo a tener un descuadre en la factura, van {n} meses seguidos",
        "he echado la cuenta con la calculadora y no me da vuestro total",
        "la linea del {prod} esta mal, el importe no es ese",
        "el descuento que me aplicais en la factura no cuadra con el que tengo",
        "cada mes lo mismo con la factura, {dif} arriba siempre",
        "mi gestor ha revisado la factura y dice que el importe esta mal",
        "el total facturado no coincide con lo que suma el albaran de {mes}",
    ),
    "facturacion-otro": (
        "me habeis mandado dos veces la misma factura del mes de {mes}",
        "tengo una factura duplicada en el banco, me la habeis pasado dos veces",
        "el iva de la factura esta mal puesto, no es el tipo que corresponde",
        "el precio del {prod} no es el que acordamos con {agente} en su dia",
        "me habeis facturado un {prod} que yo no he pedido nunca",
        "la factura viene a nombre del antiguo dueño del local",
        "me pasais al banco una factura que ya pague en efectivo",
        "el numero de factura esta repetido, tengo dos con el mismo numero",
    ),
    "entrega-sin-aviso": (
        "vinisteis y estaba cerrado, nadie me dijo nada",
        "no sabia q veniais hoy y he perdido la entrega",
        "otra vez pasasteis sin avisar y no habia nadie en el local",
        "el repartidor dice que vino ayer, yo no sabia nada",
        "si me avisarais antes de venir no pasaria esto",
        "he estado toda la mañana fuera y justo vinisteis vosotros, nadie me aviso",
        "me lo habeis dejado en la puerta porq no estaba, nadie me avisa",
        "no puedo estar aqui las 24 horas, si nadie me avisa no puedo recibirlo",
        "otra vez la misma historia, pasais y esta cerrado",
        "vine a abrir a las once y el repartidor ya se habia ido, nadie me aviso",
        "nadie me avisa de cuando toca reparto y asi es imposible",
        "he perdido la entrega otra vez porque no sabia q era hoy",
        "estamos de descanso los martes y venis los martes sin avisar",
        "el chico del reparto se fue sin dejar nada y sin avisar",
        "vinisteis a las 7 y media y yo abro a las diez, nadie me aviso",
        "pasais a primera hora y a esa hora esta cerrado, avisad antes",
    ),
    "entrega-retraso": (
        "llevo esperando el pedido desde el {dia} y no llega",
        "el reparto viene tardisimo, ayer aparecio a {hora}",
        "me dijisteis lunes y estamos a {dia} y el pedido sin llegar",
        "el pedido lleva {n} dias de retraso",
        "sigo esperando las garrafas, se me ha acabado el agua",
        "otra vez llega tarde el reparto y tengo el bar sin bebida",
        "esto se retrasa cada semana un poco mas",
        "quede en que venia el {dia} por la mañana y aun estoy esperando",
        "no ha llegado el pedido y ya es el {dia}",
        "el reparto de esta semana viene con retraso otra vez",
        "llevo dos dias sin agua con gas porque no llega el pedido",
        "me prometisteis que llegaba antes del finde y sigo esperando",
        "el reparto entra a las 7 y hoy a las 3 no habia venido nadie",
        "acabais la ruta a las 3 y hoy me quede sin pedido otra vez",
    ),
    "producto-defectuoso": (
        "las garrafas venian rotas, dos perdiendo agua",
        "la ultima caja de {prod} estaba caducada",
        "el agua sabe raro, la he tenido que tirar",
        "vienen botellas rotas en el palet, casi media caja",
        "el {prod} viene en mal estado, las cajas mojadas",
        "he abierto la caja y estaba todo estropeado",
        "la garrafa venia sucia por fuera, no la puedo poner",
        "el precinto venia abierto y falta producto dentro",
        "cuatro botellas de {prod} rotas dentro de la caja",
        "la fuente pierde agua por abajo desde que la cambiasteis",
        "el zumo tiene la fecha pasada, no lo puedo servir",
        "me llego el pedido con las cajas aplastadas",
    ),
    "pedido-erroneo": (
        "me habeis traido {prod} y yo pedi otra cosa",
        "faltan {n} cajas del pedido de ayer",
        "el pedido viene incompleto, falta la mitad",
        "me han dejado el doble de lo que pedi",
        "he pedido {n} garrafas y me han dejado {n2}, falta el resto",
        "me han traido el pedido de otro cliente",
        "en el albaran pone una cosa y en las cajas hay otra",
        "falta el {prod} en el pedido de esta semana",
        "os pedi {prod} y me habeis dejado {prod2}",
        "he recibido un pedido que yo no he hecho",
        "el albaran no corresponde con lo que me han descargado",
        "me sobran cajas de {prod}, no las he pedido",
    ),
    "cambio-datos": (
        "os quiero cambiar la cuenta del banco para los recibos",
        "hemos cambiado de direccion, apuntadla bien",
        "cambiar el telefono de contacto, el viejo ya no funciona",
        "el local ha cambiado de titular, hay que cambiar los datos",
        "quiero que me mandeis las facturas a otro correo",
        "he cambiado de numero, apuntad el nuevo",
        "cambiamos de cuenta bancaria a partir de {mes}",
        "actualizadme la direccion de entrega, nos hemos mudado",
        "el nombre fiscal esta mal escrito en las facturas",
        "cambiad el nombre del contacto, ahora lo lleva mi hija",
        "quiero domiciliar los recibos en otro banco",
    ),
    "informacion-producto": (
        "quiero saber si teneis agua con gas en garrafa",
        "cuanto cuesta el alquiler de la fuente al mes",
        "hacéis reparto a {pueblo}?",
        "que precio tiene el {prod} si cojo palet entero",
        "teneis cerveza sin alcohol?",
        "el zumo lo teneis en otro formato mas grande?",
        "que descuento hariais si subo el pedido de {prod}",
        "me podeis decir el precio del {prod} para el año que viene",
        "necesito saber si servis los sabados",
        "la garrafa es retornable? hay que devolver el envase?",
        "que dias pasais por {pueblo}",
        "teneis catalogo con todos los precios",
        "cuanto tardais en servir un pedido nuevo",
    ),
    "otros": (
        "quiero darme de baja como cliente",
        "necesito el certificado del agua para sanidad",
        "podeis mandarme el catalogo actualizado",
        "cerramos por vacaciones del {dia} al {dia2}",
        "quiero hablar con {agente} sobre el contrato",
        "necesito una copia del contrato de la fuente",
        "que horario teneis en oficina",
        "el local cierra un mes, no me traigais nada",
        "queria felicitaros, el chico del reparto es muy amable",
        "ahora abrimos por la tarde, os lo digo por el reparto",
        "me podeis dar de alta otro punto de entrega",
    ),
}

RELLENOS_DIF = ("0,43", "1,12", "0,87", "2,05", "0,29", "1,68", "3,40", "0,76",
                "1,95", "0,52", "2,31", "0,64", "4,10", "1,27")
RELLENOS_HORA = ("las 7 y cuarto", "las 7 y media", "las 8", "las 9 y media",
                 "las 13:30", "las 14", "las 14:40", "media mañana")
RELLENOS_PROD = ("agua de litro y medio", "agua de medio", "cola", "naranja",
                 "cerveza", "zumo", "agua con gas", "garrafas", "el agua chica")


MESES_DEL_PERIODO = ("septiembre", "octubre", "noviembre", "diciembre",
                    "enero", "febrero")


def _dia_suelto(az: random.Random) -> str:
    return f"{az.randint(1, 28)} de {az.choice(MESES_DEL_PERIODO)}"


def rellenar(az: random.Random, plantilla: str, contexto: dict) -> str:
    """Sustituye los huecos de la plantilla con relleno variable."""
    valores = {
        "mes": az.choice(MESES_DEL_PERIODO),
        "dif": az.choice(RELLENOS_DIF),
        "dto": str(contexto.get("dto", az.choice([3, 5, 8, 10, 12]))),
        "prod": az.choice(RELLENOS_PROD),
        "prod2": az.choice(RELLENOS_PROD),
        "n": str(az.randint(2, 9)),
        "n2": str(az.randint(2, 9)),
        "dia": _dia_suelto(az),
        "dia2": _dia_suelto(az),
        "hora": az.choice(RELLENOS_HORA),
        "pueblo": contexto.get("poblacion", az.choice(MUNICIPIOS)[0]),
        "agente": az.choice([p[0].split()[0] for p in PLANTILLA]),
    }
    salida = plantilla
    for clave, valor in valores.items():
        salida = salida.replace("{" + clave + "}", valor)
    return salida


def _protegida(palabra: str) -> bool:
    base = quitar_acentos(palabra.lower()).strip(".,;:¿?¡!()\"'")
    return any(base.startswith(raiz) for raiz in RAICES_PROTEGIDAS)


def ensuciar(az: random.Random, texto: str) -> str:
    """
    Le mete al texto las prisas de quien escribe desde el móvil: tildes que se
    comen, 'q' por 'que', alguna letra bailada, alguna palabra en mayúsculas.
    Las palabras que sostienen la reclasificación no se tocan a nivel de letra:
    quien escribe con prisa se come la tilde, no el sustantivo por el que llama.
    """
    if az.random() < 0.55:
        texto = quitar_acentos(texto)

    palabras = texto.split(" ")
    salida: list[str] = []
    for palabra in palabras:
        limpia = palabra.lower()
        if limpia.startswith("que") and az.random() < 0.35:
            palabra = "q" + palabra[3:]
        elif limpia.startswith("porque") and az.random() < 0.20:
            palabra = "xq" + palabra[6:]
        elif limpia.startswith("tambien") and az.random() < 0.30:
            palabra = "tb" + palabra[7:]
        elif not _protegida(palabra) and len(palabra) > 4 and az.random() < 0.045:
            pos = az.randrange(1, len(palabra) - 1)
            modo = az.randrange(3)
            if modo == 0:
                palabra = palabra[:pos] + palabra[pos + 1:]           # se come una
            elif modo == 1:
                palabra = palabra[:pos] + palabra[pos] + palabra[pos:]  # la repite
            else:
                palabra = palabra[:pos] + palabra[pos + 1] + palabra[pos] + palabra[pos + 2:]
        if az.random() < 0.020 and len(palabra) > 3:
            palabra = palabra.upper()
        salida.append(palabra)
    texto = " ".join(salida)

    if az.random() < 0.14:
        texto = texto.replace(" ", "  ", 1)
    return texto


def redactar(az: random.Random, categoria: str, contexto: dict,
             frases: int = 0) -> str:
    """Descripción de 1 a 3 frases en español coloquial, con erratas."""
    plantillas = PLANTILLAS[categoria]
    if frases <= 0:
        frases = elegir_pesado(az, (1, 2, 3), (0.46, 0.40, 0.14))
    elegidas = az.sample(list(plantillas), min(frases, len(plantillas)))
    trozos = [rellenar(az, p, contexto) for p in elegidas]

    texto = az.choice(APERTURAS) + trozos[0]
    for extra in trozos[1:]:
        texto += az.choice((". ", ". ", ", y ademas ", ". Y otra cosa, ", ". "))
        texto += extra
    texto += az.choice(CIERRES)

    texto = ensuciar(az, texto)
    if az.random() < 0.25:
        texto = texto[0].lower() + texto[1:] if texto else texto
    else:
        texto = texto[0].upper() + texto[1:] if texto else texto
    if az.random() < 0.35:
        texto = texto.rstrip(". ")
    elif not texto.endswith((".", "?", "!")):
        texto += "."
    return re.sub(r"\s+$", "", texto)


# ═══════════════════════════════════════════════════════════════════════════════
# 4. CLIENTES  (fichero 1 + V2)
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class Cliente:
    indice: int
    id_cliente: str
    nombre_real: str
    nombre: str
    tipo_real: str            # "hosteleria" | "particular"
    tipo: str                 # lo que se escribe en la celda, sucio
    direccion: str
    poblacion: str
    cp: str
    cp_como_numero: bool
    telefono_digitos: str
    telefono: str
    email: str
    email_valido: bool
    fecha_alta: date
    ruta: str
    forma_pago: str
    descuento_pct: int
    descuento_celda: object
    observaciones: str
    es_duplicado_de: str = ""


def generar_clientes(az: random.Random) -> tuple[list[Cliente], list[tuple[str, str]]]:
    """300 filas, 288 clientes reales y 12 duplicados (V2)."""
    indices_v4 = [int(c.split("-")[1]) - 1 for c in CLIENTES_V4]

    # ── Tipo real: 110 hostelería, 190 particulares. Los tres de V4 son hostelería.
    resto = [i for i in range(N_CLIENTES) if i not in indices_v4]
    hosteleria = set(indices_v4) | set(az.sample(resto, N_HOSTELERIA - len(indices_v4)))
    tipos = ["hosteleria" if i in hosteleria else "particular" for i in range(N_CLIENTES)]

    # ── Pares de duplicados: el duplicado siempre muy posterior al original.
    candidatos_dup = [i for i in range(150, N_CLIENTES) if i not in indices_v4]
    duplicados = sorted(az.sample(candidatos_dup, N_DUPLICADOS))
    originales: dict[int, int] = {}
    usados: set[int] = set()
    for dup in duplicados:
        posibles = [i for i in range(0, 120)
                    if i not in indices_v4 and i not in usados and tipos[i] == tipos[dup]]
        elegido = az.choice(posibles)
        usados.add(elegido)
        originales[dup] = elegido
    # En 5 de los 12 pares el descuento difiere: es el daño real del duplicado.
    dup_con_dto_distinto = set(az.sample(duplicados, 5))

    # ── Nombres únicos.
    nombres_hosteleria = []
    for tipo_local in TIPOS_LOCAL:
        for nombre in NOMBRES_LOCAL:
            nombres_hosteleria.append(f"{tipo_local} {nombre}")
    az.shuffle(nombres_hosteleria)
    nombres_particular = []
    for pila in NOMBRES_PILA:
        for ap1 in APELLIDOS:
            nombres_particular.append(f"{pila} {ap1}")
    az.shuffle(nombres_particular)

    # ── Direcciones únicas (calle, número, población).
    direcciones_usadas: set[tuple[str, int, str]] = set()
    telefonos_usados: set[str] = set()

    clientes: list[Cliente] = []
    it_host = iter(nombres_hosteleria)
    it_part = iter(nombres_particular)

    for i in range(N_CLIENTES):
        id_cliente = f"CLI-{i + 1:04d}"
        tipo_real = tipos[i]
        nombre_real = next(it_host) if tipo_real == "hosteleria" else next(it_part)

        municipio, cp, ruta_base = az.choice(MUNICIPIOS)
        while True:
            calle = az.choice(CALLES)
            numero = az.randint(1, 84)
            if (calle, numero, municipio) not in direcciones_usadas:
                direcciones_usadas.add((calle, numero, municipio))
                break
        direccion = f"{calle} {numero}"

        while True:
            prefijo = az.choice(("942", "942", "942", "985", "6" + str(az.randint(10, 99)),
                                 "6" + str(az.randint(10, 99)), "722"))
            digitos = prefijo + "".join(str(az.randint(0, 9)) for _ in range(9 - len(prefijo)))
            if digitos not in telefonos_usados:
                telefonos_usados.add(digitos)
                break

        fecha_alta = date(2015, 1, 1) + timedelta(days=az.randint(0, 3500))
        forma_pago = elegir_pesado(az, ("domiciliado", "transferencia", "efectivo"),
                                   (0.62, 0.22, 0.16))

        if tipo_real == "hosteleria":
            descuento = elegir_pesado(az, (0, 3, 5, 8, 10, 12),
                                      (0.18, 0.16, 0.22, 0.20, 0.14, 0.10))
        else:
            descuento = elegir_pesado(az, (0, 3, 5), (0.92, 0.05, 0.03))

        clientes.append(Cliente(
            indice=i, id_cliente=id_cliente, nombre_real=nombre_real, nombre=nombre_real,
            tipo_real=tipo_real, tipo="", direccion=direccion, poblacion=municipio,
            cp=cp, cp_como_numero=False, telefono_digitos=digitos,
            telefono=formatear_telefono(digitos, az.randrange(4)),
            email="", email_valido=False, fecha_alta=fecha_alta, ruta=ruta_base,
            forma_pago=forma_pago, descuento_pct=descuento, descuento_celda=descuento,
            observaciones="",
        ))

    # ── Los tres de V4: hostelería con descuento alto (10–12).
    for indice, dto in zip(indices_v4, (12, 10, 12)):
        clientes[indice].descuento_pct = dto
        clientes[indice].descuento_celda = dto

    # ── V2: el duplicado hereda dirección y teléfono del original y los disfraza.
    pares: list[tuple[str, str]] = []
    for dup in duplicados:
        orig = originales[dup]
        c_orig, c_dup = clientes[orig], clientes[dup]

        telefonos_usados.discard(c_dup.telefono_digitos)
        c_dup.telefono_digitos = c_orig.telefono_digitos
        # Formato distinto del original: por eso la comparación literal falla.
        formatos = [f for f in range(4)
                    if formatear_telefono(c_orig.telefono_digitos, f) != c_orig.telefono]
        c_dup.telefono = formatear_telefono(c_orig.telefono_digitos, az.choice(formatos))

        c_dup.poblacion = c_orig.poblacion
        c_dup.cp = c_orig.cp
        c_dup.ruta = c_orig.ruta
        calle, numero = c_orig.direccion.rsplit(" ", 1)
        variantes = [
            f"{calle}, {numero}",
            f"{calle.replace('C/', 'Calle').replace('Avda.', 'Avenida')} {numero}",
            f"{calle.replace('C/', 'Calle').replace('Avda.', 'Avenida')}, {numero}",
            f"{calle} nº {numero}",
            f"{calle.upper()} {numero}",
        ]
        c_dup.direccion = az.choice(variantes)

        # Nombre: forma jurídica, abreviatura, mayúsculas, mojibake o inversión.
        # Un particular no lleva forma jurídica: a él se le abrevia el nombre o
        # se le invierte con el apellido, que es como lo teclea otra persona.
        base = c_orig.nombre_real
        if c_orig.tipo_real == "hosteleria":
            variantes_nombre = [
                base.upper() + az.choice((" S.L.", " SL")),
                (base.replace("Bar Restaurante", "Bar Rest.")
                     .replace("Restaurante", "Rest.")
                     .replace("Cafetería", "Cafet.")
                     .replace("Marisquería", "Marisq.")
                     .replace("Cervecería", "Cervec.")
                     .replace("Chiringuito", "Chiring.")),
                base.upper(),
                base + " S.L.",
            ]
            partes = base.split(" ", 1)
            if len(partes) == 2:
                variantes_nombre.append(f"{partes[1]}, {partes[0]}")
        else:
            partes = base.split(" ", 1)
            variantes_nombre = [base.upper()]
            if len(partes) == 2:
                variantes_nombre += [
                    f"{partes[1]}, {partes[0]}",
                    f"{partes[0][0]}. {partes[1]}",
                    f"{partes[1].upper()}, {partes[0]}",
                ]
        variantes_nombre = [v for v in variantes_nombre if v != base]
        c_dup.nombre_real = az.choice(variantes_nombre)

        c_dup.fecha_alta = min(
            c_orig.fecha_alta + timedelta(days=az.randint(400, 2400)), date(2024, 8, 20))
        c_dup.tipo_real = c_orig.tipo_real
        c_dup.es_duplicado_de = c_orig.id_cliente

        if dup in dup_con_dto_distinto:
            # El segundo alta se dio con otras condiciones. Para un particular
            # las condiciones posibles son las de un particular, no las de un bar.
            posibles = ((0, 3, 5, 8, 10, 12) if c_orig.tipo_real == "hosteleria"
                        else (0, 3, 5))
            opciones = [d for d in posibles if d != c_orig.descuento_pct]
            c_dup.descuento_pct = az.choice(opciones)
        else:
            c_dup.descuento_pct = c_orig.descuento_pct
        c_dup.descuento_celda = c_dup.descuento_pct
        pares.append((c_orig.id_cliente, c_dup.id_cliente))

    # ── Suciedad de presentación, ya con las identidades cerradas, y con las
    #    cuotas exactas que fija la especificación.
    grafias_host = ("Hostelería", "HOSTELERIA", "hosteleria")
    grafias_part = ("Particular", "part.", "Particular")
    indices_dup = set(duplicados)
    fichas_pareja = indices_dup | set(originales.values())

    def con_acentos(indice: int, texto: str) -> bool:
        return any(ord(caracter) > 127 for caracter in texto)

    mojibake_nombre = muestra_exacta(
        az, [c.indice for c in clientes if con_acentos(c.indice, c.nombre_real)],
        round(0.08 * N_CLIENTES))
    # La dirección de las fichas emparejadas se deja limpia: la codificación rota
    # es otra suciedad y no debe estorbar a la normalización que revela V2.
    mojibake_direccion = muestra_exacta(
        az, [c.indice for c in clientes
             if c.indice not in fichas_pareja and con_acentos(c.indice, c.direccion)],
        round(0.08 * N_CLIENTES))
    tipo_vacio = muestra_exacta(az, range(N_CLIENTES), round(0.03 * N_CLIENTES))
    cp_numerico = muestra_exacta(az, range(N_CLIENTES), round(0.05 * N_CLIENTES))
    ruta_vacia = muestra_exacta(az, range(N_CLIENTES), round(0.02 * N_CLIENTES))
    con_observaciones = muestra_exacta(az, range(N_CLIENTES), round(0.30 * N_CLIENTES))

    indices_v4_set = {int(c.split("-")[1]) - 1 for c in CLIENTES_V4}
    candidatos_email = [c.indice for c in clientes if c.indice not in indices_v4_set]
    sin_email = muestra_exacta(az, candidatos_email, round(0.18 * N_CLIENTES))
    email_roto = muestra_exacta(az, [i for i in candidatos_email if i not in sin_email],
                                round(0.02 * N_CLIENTES))
    dto_vacio = muestra_exacta(az, [c.indice for c in clientes if c.descuento_pct == 0],
                               round(0.04 * N_CLIENTES))

    for c in clientes:
        c.nombre = (romper_codificacion(c.nombre_real) if c.indice in mojibake_nombre
                    else c.nombre_real)
        if c.indice in mojibake_direccion:
            c.direccion = romper_codificacion(c.direccion)

        if c.indice in tipo_vacio:
            c.tipo = ""
        else:
            c.tipo = az.choice(grafias_host if c.tipo_real == "hosteleria" else grafias_part)

        c.cp_como_numero = c.indice in cp_numerico
        if c.indice in ruta_vacia:
            c.ruta = ""

        base_email = quitar_acentos(c.nombre_real).lower()
        base_email = re.sub(r"[^a-z0-9]+", "", base_email)[:22]
        if c.tipo_real == "hosteleria" and az.random() < 0.45:
            dominio = base_email[:18] + ".es"
            direccion_email = f"info@{dominio}"
        else:
            direccion_email = f"{base_email}{c.indice + 1}@{az.choice(DOMINIOS_CORREO)}"
        if c.indice in sin_email:
            c.email, c.email_valido = "", False
        elif c.indice in email_roto:
            c.email, c.email_valido = direccion_email.replace("@", "."), False
        else:
            # Los tres de V4 caen siempre aquí: están sobrerrepresentados en el
            # buzón y sin correo válido no podrían escribir.
            c.email, c.email_valido = direccion_email, True

        # `descuento_pct` vacío equivale a 0: solo se vacía donde de verdad es 0.
        if c.indice in dto_vacio:
            c.descuento_celda = ""
        c.observaciones = (az.choice(OBSERVACIONES) if c.indice in con_observaciones
                           else "")

    return clientes, pares


# ═══════════════════════════════════════════════════════════════════════════════
# 5. PEDIDOS  (fichero 2 + V1 + V3 + devoluciones)
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class Linea:
    producto: str
    cantidad: int
    precio_ud: Decimal
    dto_pct: int
    importe: Decimal
    importe_ok: Decimal

    @property
    def desviacion(self) -> Decimal:
        return self.importe - self.importe_ok


@dataclass
class Pedido:
    id_pedido: str
    fecha: date
    id_cliente: str
    lineas: list[Linea] = field(default_factory=list)
    es_fantasma: bool = False
    es_devolucion: bool = False
    del_lote: bool = False          # entró por la carga automática de diciembre
    original_de: str = ""

    @property
    def total(self) -> Decimal:
        return sum((l.importe for l in self.lineas), Decimal("0.00"))

    @property
    def desviacion(self) -> Decimal:
        return sum((l.desviacion for l in self.lineas), Decimal("0.00"))


CANTIDADES = {
    "AG-05": {"hosteleria": (24, 48, 72, 96, 120, 144), "particular": (24, 24, 48)},
    "AG-15": {"hosteleria": (12, 18, 24, 36, 48, 60), "particular": (6, 12, 12, 18, 24)},
    "AG-GA": {"hosteleria": (12, 24, 36, 48), "particular": (12, 12, 24)},
    "GF-19": {"hosteleria": (3, 4, 5, 6, 8, 10, 12), "particular": (2, 3, 4, 4, 6)},
    "RE-05": {"hosteleria": (24, 48, 72, 96), "particular": (24, 24, 48)},
    "RE-NA": {"hosteleria": (24, 48, 72), "particular": (24, 24, 48)},
    "CE-33": {"hosteleria": (24, 48, 72, 96, 120), "particular": (24, 48)},
    "ZU-20": {"hosteleria": (18, 36, 54), "particular": (18, 18, 36)},
    "FU-AL": {"hosteleria": (1,), "particular": (1,)},
    "PO-10": {"hosteleria": (1,), "particular": (1,)},
}
PESOS_PRODUCTO = {
    "hosteleria": (("AG-05", 0.20), ("AG-15", 0.12), ("AG-GA", 0.10), ("GF-19", 0.10),
                   ("RE-05", 0.12), ("RE-NA", 0.10), ("CE-33", 0.16), ("ZU-20", 0.06),
                   ("FU-AL", 0.02), ("PO-10", 0.02)),
    "particular": (("AG-05", 0.10), ("AG-15", 0.25), ("AG-GA", 0.08), ("GF-19", 0.35),
                   ("RE-05", 0.07), ("RE-NA", 0.05), ("CE-33", 0.05), ("ZU-20", 0.03),
                   ("FU-AL", 0.01), ("PO-10", 0.01)),
}
PESOS_DIA_SEMANA = (1.25, 1.05, 1.00, 1.00, 0.85)   # lunes … viernes


def generar_pedidos(az: random.Random,
                    clientes: list[Cliente]) -> tuple[list[Pedido], list[Pedido]]:
    """
    Pedidos reales + los 62 fantasma de V3. Devuelve (todos, fantasma).
    El presupuesto de líneas por mes es fijo: es lo que hace que los porcentajes
    de V3 salgan clavados sin depender del azar.
    """
    dias = dias_laborables(FECHA_INICIO, FECHA_FIN)
    assert len(dias) == DIAS_LABORABLES_ESPERADOS, (
        f"El periodo tiene {len(dias)} días laborables y la especificación dice "
        f"{DIAS_LABORABLES_ESPERADOS}")

    ventana = [d for d in dias if VENTANA_FANTASMA_INI <= d <= VENTANA_FANTASMA_FIN]
    assert len(ventana) == 5

    # ── Reparto de los presupuestos de pedidos y de líneas, mes a mes ─────────
    # Los dos son cifras fijas: así diciembre sube exactamente lo mismo se cuente
    # en pedidos o en líneas, y la doble importación de V3 se ve igual de clara
    # con cualquiera de las dos medidas.
    huecos: list[tuple[date, int, bool]] = []   # (día, nº líneas, es del lote)

    for (anio, mes), presupuesto in sorted(PRESUPUESTO_LINEAS.items()):
        dias_mes = [d for d in dias if d.year == anio and d.month == mes]
        es_diciembre = (anio, mes) == (2024, 12)
        lineas_libres = presupuesto - LINEAS_FANTASMA if es_diciembre else presupuesto
        pedidos_libres = PRESUPUESTO_PEDIDOS[(anio, mes)]
        if es_diciembre:
            pedidos_libres -= PEDIDOS_FANTASMA_LOTE

        pesos = []
        for d in dias_mes:
            peso = PESOS_DIA_SEMANA[d.weekday()] * (0.90 + 0.20 * az.random())
            # En la semana del lote hay menos pedidos "a mano": casi todo entró
            # por la carga automática. Por eso la ventana pesa menos aquí.
            if es_diciembre and VENTANA_FANTASMA_INI <= d <= VENTANA_FANTASMA_FIN:
                peso *= 0.40
            pesos.append(peso)
        pedidos_por_dia = repartir_entero(pedidos_libres, pesos)

        pesos_tam = PESOS_TAMANO_DICIEMBRE if es_diciembre else PESOS_TAMANO_NORMAL
        tamanos_mes: list[int] = []
        dias_de_cada = []
        for d, cuantos in zip(dias_mes, pedidos_por_dia):
            for _ in range(cuantos):
                tamanos_mes.append(elegir_pesado(az, TAMANOS, pesos_tam))
                dias_de_cada.append(d)
        tamanos_mes = ajustar_tamanos(tamanos_mes, lineas_libres)
        for d, tam in zip(dias_de_cada, tamanos_mes):
            huecos.append((d, tam, False))

    # ── El lote del 9 al 13 de diciembre: 62 pedidos, 244 líneas exactas ──────
    tamanos_lote = ajustar_tamanos(
        [elegir_pesado(az, TAMANOS, PESOS_TAMANO_LOTE) for _ in range(N_FANTASMA)],
        LINEAS_FANTASMA)
    reparto_lote = repartir_entero(N_FANTASMA, [1.15, 1.10, 1.00, 0.95, 0.80])
    cursor = 0
    for d, cuantos in zip(ventana, reparto_lote):
        for _ in range(cuantos):
            huecos.append((d, tamanos_lote[cursor], True))
            cursor += 1

    huecos.sort(key=lambda h: (h[0], not h[2]))

    # ── Asignación de clientes ────────────────────────────────────────────────
    por_id = {c.id_cliente: c for c in clientes}
    asignado: list[str] = [""] * len(huecos)
    usados_dia: dict[date, set[str]] = {}

    # Los tres de V4 primero: muchos pedidos pequeños. Volumen medio-bajo y
    # muchísimo contacto es justo el perfil que sale deficitario.
    objetivos_v4 = {"CLI-0042": 44, "CLI-0118": 40, "CLI-0233": 38}
    for id_cliente in CLIENTES_V4:
        candidatos = [k for k, (d, t, lote) in enumerate(huecos)
                      if t <= 2 and not asignado[k]]
        az.shuffle(candidatos)
        puestos = 0
        for k in candidatos:
            if puestos >= objetivos_v4[id_cliente]:
                break
            dia = huecos[k][0]
            if id_cliente in usados_dia.setdefault(dia, set()):
                continue
            asignado[k] = id_cliente
            usados_dia[dia].add(id_cliente)
            puestos += 1

    # Las 24 fichas implicadas en los 12 pares duplicados (V2) tienen que tener
    # pedidos las dos: si una estuviera vacía, el duplicado se detectaría por ahí
    # y el ejercicio perdería la gracia. La especificación lo exige.
    fichas_duplicadas: list[str] = []
    for c in clientes:
        if c.es_duplicado_de:
            fichas_duplicadas.append(c.id_cliente)
            fichas_duplicadas.append(c.es_duplicado_de)
    for id_cliente in sorted(set(fichas_duplicadas)):
        libres = [k for k, _ in enumerate(huecos) if not asignado[k]]
        az.shuffle(libres)
        for k in libres:
            dia = huecos[k][0]
            if id_cliente in usados_dia.setdefault(dia, set()):
                continue
            asignado[k] = id_cliente
            usados_dia[dia].add(id_cliente)
            break

    pool = [c.id_cliente for c in clientes if c.id_cliente not in CLIENTES_V4]
    pesos_pool = [5.0 if por_id[i].tipo_real == "hosteleria" else 1.0 for i in pool]
    for k, (dia, _tam, _lote) in enumerate(huecos):
        if asignado[k]:
            continue
        ocupados = usados_dia.setdefault(dia, set())
        for _ in range(60):
            candidato = az.choices(pool, weights=pesos_pool, k=1)[0]
            if candidato not in ocupados:
                break
        else:  # pragma: no cover - con 300 clientes y ~15 pedidos/día no ocurre
            candidato = next(i for i in pool if i not in ocupados)
        asignado[k] = candidato
        ocupados.add(candidato)

    # ── Construcción de los pedidos ───────────────────────────────────────────
    pedidos: list[Pedido] = []
    for k, (dia, tam, es_lote) in enumerate(huecos):
        id_cliente = asignado[k]
        cliente = por_id[id_cliente]
        pedido = Pedido(id_pedido="", fecha=dia, id_cliente=id_cliente)
        codigos, pesos = zip(*PESOS_PRODUCTO[cliente.tipo_real])
        elegidos: list[str] = []
        while len(elegidos) < tam:
            codigo = az.choices(list(codigos), weights=list(pesos), k=1)[0]
            if codigo not in elegidos:
                elegidos.append(codigo)
        for codigo in elegidos:
            opciones = CANTIDADES[codigo][cliente.tipo_real]
            cantidad = az.choice(opciones)
            if id_cliente in CLIENTES_V4:
                cantidad = min(opciones)     # pedidos pequeños y frecuentes
            precio = PRODUCTOS[codigo]["precio"]
            dto = cliente.descuento_pct
            pedido.lineas.append(Linea(
                producto=codigo, cantidad=cantidad, precio_ud=precio, dto_pct=dto,
                importe=importe_facturado(precio, dto, cantidad),
                importe_ok=importe_correcto(precio, dto, cantidad),
            ))
        pedido.del_lote = es_lote
        pedidos.append(pedido)

    # ── Devoluciones: ~45 pedidos en negativo, sin marca ninguna ──────────────
    # Solo de clientes sin descuento, para no romper la invariante de V1
    # (`dto_pct > 0 ⇒ desviación ≥ 0`, que es la pista de que el fallo va con el
    # descuento). Un abono con descuento invertiría el signo de la desviación.
    candidatos_dev = [k for k, p in enumerate(pedidos)
                      if len(p.lineas) <= 2
                      and por_id[p.id_cliente].descuento_pct == 0
                      and p.id_cliente not in CLIENTES_V4
                      and not (VENTANA_FANTASMA_INI <= p.fecha <= VENTANA_FANTASMA_FIN)]
    for k in az.sample(candidatos_dev, N_DEVOLUCIONES):
        pedido = pedidos[k]
        pedido.es_devolucion = True
        for linea in pedido.lineas:
            linea.cantidad = -linea.cantidad
            linea.importe = importe_facturado(linea.precio_ud, linea.dto_pct, linea.cantidad)
            linea.importe_ok = importe_correcto(linea.precio_ud, linea.dto_pct, linea.cantidad)

    # ── Numeración correlativa por fecha ──────────────────────────────────────
    pedidos.sort(key=lambda p: (p.fecha, p.id_cliente))
    for n, pedido in enumerate(pedidos, start=1):
        pedido.id_pedido = f"PED-{n:05d}"

    # ── V3: la carga del 9 al 13 se importó dos veces ─────────────────────────
    # El fichero de carga de esa semana se volvió a subir. Cada pedido del lote
    # se repite tal cual, con otro id_pedido de un correlativo aparte.
    originales_lote = sorted((p for p in pedidos if p.del_lote),
                             key=lambda p: p.id_pedido)
    assert len(originales_lote) == N_FANTASMA, (
        f"Se han identificado {len(originales_lote)} pedidos del lote y deben ser "
        f"{N_FANTASMA}")

    fantasma: list[Pedido] = []

    for n, original in enumerate(originales_lote, start=1):
        copia = Pedido(
            id_pedido=f"PED-9{n:04d}", fecha=original.fecha,
            id_cliente=original.id_cliente, es_fantasma=True,
            original_de=original.id_pedido,
            lineas=[Linea(l.producto, l.cantidad, l.precio_ud, l.dto_pct,
                          l.importe, l.importe_ok) for l in original.lineas],
        )
        fantasma.append(copia)

    return pedidos + fantasma, fantasma


# ═══════════════════════════════════════════════════════════════════════════════
# 6. TICKETS  (fichero 3 + V1 + V4 + V5)
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class Ticket:
    id_ticket: str
    fecha_apertura: date
    formato_fecha: int
    canal: str
    id_cliente: str
    id_cliente_visible: str
    categoria_real: str
    subtipo: str
    categoria_sucia: str
    subcategoria: str
    descripcion: str
    id_pedido: str
    estado: str
    fecha_cierre: object
    agente: str
    tiempo: object


GRAFIAS_CANAL = {
    "telefono": ("Teléfono", "TELEFONO", "tel", "telefono", "Telefono"),
    "email": ("Email", "email", "correo", "e-mail", "EMAIL"),
    "whatsapp": ("WhatsApp", "whatsapp", "wasap", "WA"),
    "presencial": ("Presencial", "presencial", "mostrador"),
}
GRAFIAS_ESTADO = {
    "cerrado": ("cerrado", "Cerrado", "CERRADO", "cerrada", "Cerrado "),
    "abierto": ("abierto", "Abierto", "ABIERTO"),
    "pendiente": ("pendiente", "Pendiente", "en curso"),
}

MEZCLA_V4 = {
    "CLI-0042": (("facturacion-redondeo", 26), ("entrega-sin-aviso", 8),
                 ("entrega-retraso", 10), ("producto-defectuoso", 4),
                 ("pedido-erroneo", 7), ("cambio-datos", 2),
                 ("informacion-producto", 3), ("otros", 2)),
    "CLI-0118": (("facturacion-redondeo", 24), ("entrega-sin-aviso", 7),
                 ("entrega-retraso", 9), ("producto-defectuoso", 4),
                 ("pedido-erroneo", 7), ("cambio-datos", 2),
                 ("informacion-producto", 3), ("otros", 2)),
    "CLI-0233": (("facturacion-redondeo", 22), ("entrega-sin-aviso", 7),
                 ("entrega-retraso", 9), ("producto-defectuoso", 4),
                 ("pedido-erroneo", 7), ("cambio-datos", 2),
                 ("informacion-producto", 3), ("otros", 2)),
}
# Tickets con id_pedido informado, por categoría. Los de V1 llevan el 86 %.
CON_PEDIDO = {
    "facturacion-redondeo": 261,
    "pedido-erroneo": 45,
    "entrega-retraso": 30,
    "producto-defectuoso": 28,
    "entrega-sin-aviso": 20,
}


def generar_tickets(az: random.Random, clientes: list[Cliente],
                    pedidos: list[Pedido]) -> list[Ticket]:
    por_id = {c.id_cliente: c for c in clientes}
    dias = dias_laborables(FECHA_INICIO, FECHA_FIN)
    dias_set = set(dias)

    reales = [p for p in pedidos if not p.es_fantasma]
    con_desviacion: dict[str, list[Pedido]] = {}
    cualquiera: dict[str, list[Pedido]] = {}
    for p in reales:
        cualquiera.setdefault(p.id_cliente, []).append(p)
        if p.desviacion > 0:
            con_desviacion.setdefault(p.id_cliente, []).append(p)

    hosteleria = sorted(c.id_cliente for c in clientes if c.tipo_real == "hosteleria")
    clientes_v1 = sorted(k for k, v in con_desviacion.items() if v)

    # ── Reparto de las 800 incidencias por categoría y cliente ────────────────
    encargos: list[tuple[str, str]] = []   # (categoría real, id_cliente)
    pendientes = {cat: n for cat, n in TAXONOMIA}

    for id_cliente in CLIENTES_V4:
        for categoria, cuantos in MEZCLA_V4[id_cliente]:
            for _ in range(cuantos):
                encargos.append((categoria, id_cliente))
                pendientes[categoria] -= 1

    for categoria, restantes in sorted(pendientes.items()):
        for _ in range(restantes):
            if categoria == "facturacion-redondeo":
                id_cliente = az.choice([c for c in clientes_v1 if c not in CLIENTES_V4])
            elif categoria == "entrega-sin-aviso":
                id_cliente = az.choice([c for c in hosteleria if c not in CLIENTES_V4])
            else:
                todos = [c.id_cliente for c in clientes if c.id_cliente not in CLIENTES_V4]
                id_cliente = az.choice(todos)
            encargos.append((categoria, id_cliente))
    az.shuffle(encargos)
    assert len(encargos) == N_TICKETS

    # ── Los 17 tickets de facturación que NO vienen del fallo ─────────────────
    indices_pedido_erroneo = [k for k, (cat, _) in enumerate(encargos)
                              if cat == "pedido-erroneo"]
    indices_fact_otro = set(az.sample(indices_pedido_erroneo, N_FACTURACION_NO_REDONDEO))

    # ── Cuántos de cada categoría llevan id_pedido ────────────────────────────
    cupo_pedido = dict(CON_PEDIDO)

    tickets: list[Ticket] = []
    for k, (categoria, id_cliente) in enumerate(encargos):
        cliente = por_id[id_cliente]
        subtipo = "facturacion-otro" if k in indices_fact_otro else categoria

        # Fecha. Los de facturación caen tras el cierre mensual (PR-06): la
        # factura sale a principio de mes y el teléfono suena esa misma semana.
        pedido_ligado: Pedido | None = None
        if categoria == "facturacion-redondeo":
            candidatos = [p for p in con_desviacion.get(id_cliente, [])
                          if p.fecha < date(2025, 2, 1)]
            if not candidatos:
                candidatos = con_desviacion.get(id_cliente, [])
            pedido_ligado = az.choice(candidatos) if candidatos else None
            if pedido_ligado is not None:
                mes_siguiente = (pedido_ligado.fecha.replace(day=1)
                                 + timedelta(days=32)).replace(day=1)
                fecha = mes_siguiente + timedelta(days=az.randint(0, 13))
            else:
                fecha = az.choice(dias)
        else:
            candidatos = [p for p in cualquiera.get(id_cliente, []) if not p.es_devolucion]
            if candidatos and cupo_pedido.get(categoria, 0) > 0:
                pedido_ligado = az.choice(candidatos)
                fecha = pedido_ligado.fecha + timedelta(days=az.randint(0, 5))
            else:
                pedido_ligado = None
                fecha = az.choice(dias)

        while fecha not in dias_set:
            fecha += timedelta(days=1)
            if fecha > FECHA_FIN:
                fecha = az.choice(dias)
        if fecha < FECHA_INICIO or fecha > FECHA_FIN:
            fecha = az.choice(dias)

        # id_pedido: se informa según cupo y solo si el pedido es coherente.
        id_pedido = ""
        if pedido_ligado is not None and cupo_pedido.get(categoria, 0) > 0 \
                and pedido_ligado.fecha <= fecha:
            id_pedido = pedido_ligado.id_pedido
            cupo_pedido[categoria] -= 1

        contexto = {"dto": cliente.descuento_pct, "poblacion": cliente.poblacion}
        descripcion = redactar(az, subtipo, contexto)

        # Etiqueta sucia
        if subtipo == "facturacion-otro":
            etiquetas = (("Facturación", 0.30), ("Facturas", 0.24),
                         ("Incidencia facturación", 0.18), ("facturacion", 0.16),
                         ("Otros", 0.08), ("", 0.04))
        else:
            etiquetas = MEZCLA_ETIQUETAS[categoria]
        nombres, pesos = zip(*etiquetas)
        categoria_sucia = az.choices(list(nombres), weights=list(pesos), k=1)[0]

        canal_real = elegir_pesado(az, ("telefono", "email", "whatsapp", "presencial"),
                                   (0.46, 0.28, 0.19, 0.07))
        canal = az.choice(GRAFIAS_CANAL[canal_real])

        estado_real = elegir_pesado(az, ("cerrado", "abierto", "pendiente"),
                                    (0.80, 0.09, 0.11))
        estado = az.choice(GRAFIAS_ESTADO[estado_real])

        if estado_real == "cerrado":
            if az.random() < 0.14:
                cierre = ""
            else:
                cierre = fecha + timedelta(days=az.randint(0, 12))
        else:
            cierre = ""

        agente = az.choice(az.choice(PLANTILLA)[2])
        tiempo = elegir_pesado(az, (3, 5, 8, 10, 12, 15, 20, 25, 30, 45),
                               (0.06, 0.12, 0.16, 0.16, 0.12, 0.12, 0.10, 0.07,
                                0.06, 0.03))
        subcategoria = az.choice(SUBCATEGORIAS)

        tickets.append(Ticket(
            id_ticket="", fecha_apertura=fecha, formato_fecha=0, canal=canal, id_cliente=id_cliente, id_cliente_visible=id_cliente,
            categoria_real=categoria, subtipo=subtipo, categoria_sucia=categoria_sucia,
            subcategoria=subcategoria, descripcion=descripcion, id_pedido=id_pedido,
            estado=estado, fecha_cierre=cierre, agente=agente, tiempo=tiempo,
        ))

    tickets.sort(key=lambda t: (t.fecha_apertura, t.id_cliente))
    for n, t in enumerate(tickets, start=1):
        t.id_ticket = f"TCK-{n:04d}"

    # ── Suciedad final, con las cuotas exactas de la especificación ───────────
    # Tres formatos de fecha conviviendo en la misma columna: 55 / 30 / 15 %.
    posiciones = list(range(N_TICKETS))
    en_barras = muestra_exacta(az, posiciones, round(0.30 * N_TICKETS))
    resto = [p for p in posiciones if p not in en_barras]
    con_mes = muestra_exacta(az, resto, round(0.15 * N_TICKETS))
    for p in posiciones:
        tickets[p].formato_fecha = 1 if p in en_barras else (2 if p in con_mes else 0)

    sin_agente = muestra_exacta(az, posiciones, round(0.06 * N_TICKETS))
    sin_tiempo = muestra_exacta(az, posiciones, round(0.22 * N_TICKETS))
    con_subcategoria = muestra_exacta(az, posiciones, round(0.22 * N_TICKETS))
    for p in posiciones:
        if p in sin_agente:
            tickets[p].agente = ""
        if p in sin_tiempo:
            tickets[p].tiempo = ""
        if p not in con_subcategoria:
            tickets[p].subcategoria = ""

    # ── Suciedad final ────────────────────────────────────────────────────────
    # 9 filas con cierre anterior a la apertura.
    cerrados = [t for t in tickets if isinstance(t.fecha_cierre, date)]
    for t in az.sample(cerrados, 9):
        t.fecha_cierre = t.fecha_apertura - timedelta(days=az.randint(1, 20))

    # 7 valores absurdos de tiempo dedicado.
    con_tiempo = [t for t in tickets if t.tiempo != ""]
    for t, valor in zip(az.sample(con_tiempo, 7), (999, 0, 1440, 999, 0, 1440, 999)):
        t.tiempo = valor

    # ~4 % sin id_cliente. Los de los tres clientes de V4 se dejan intactos: la
    # concentración de V4 es una cifra exacta de la especificación y no se puede
    # ir diluyendo con ruido.
    candidatos_sin_cliente = [t for t in tickets if t.id_cliente not in CLIENTES_V4]
    for t in az.sample(candidatos_sin_cliente, 32):
        t.id_cliente_visible = ""

    return tickets


# ═══════════════════════════════════════════════════════════════════════════════
# 7. CORREOS  (fichero 4)
# ═══════════════════════════════════════════════════════════════════════════════

@dataclass
class Correo:
    numero: int
    id_mensaje: str
    de_nombre: str
    de_email: str
    asunto: str
    cuerpo: str
    momento: datetime
    responde_a: str
    patologia: str
    id_cliente: str
    categoria: str


ASUNTOS_MUDOS = ("Consulta", "Buenos días", "-", "Pregunta", "(sin asunto)",
                 "Hola", "Duda", "Buenas tardes")
ASUNTOS = {
    "facturacion-redondeo": ("Factura de {mes}", "Descuadre en la factura",
                             "Revisar factura", "La factura no cuadra",
                             "Duda con la factura de {mes}", "Factura mal"),
    "facturacion-otro": ("Factura duplicada", "Problema con la factura",
                         "IVA de la factura", "Factura repetida"),
    "entrega-sin-aviso": ("Reparto de ayer", "No estaba nadie", "Entrega perdida",
                          "Vinisteis y estaba cerrado", "Reparto"),
    "entrega-retraso": ("Pedido sin llegar", "Retraso del pedido",
                        "El reparto llega tarde", "Sigo esperando"),
    "producto-defectuoso": ("Garrafas rotas", "Producto en mal estado",
                            "Cajas rotas", "Problema con el género"),
    "pedido-erroneo": ("Pedido mal servido", "Falta género", "Pedido incompleto",
                       "Error en el pedido"),
    "cambio-datos": ("Cambio de datos", "Nueva cuenta bancaria",
                     "Cambio de dirección", "Actualizar datos"),
    "informacion-producto": ("Consulta de precios", "Información", "Precio garrafa",
                             "Pregunta sobre la fuente"),
    "otros": ("Baja", "Certificado", "Catálogo", "Vacaciones", "Consulta general"),
}
FIRMAS_AJENAS = ("Charo", "Mi hija Laura", "Toñi", "Javi (el encargado)",
                 "Roberto, el socio", "Marisa", "El cocinero, Andrés",
                 "Pili", "Su hermano, Fernando")

REPARTO_CORREOS = (
    ("facturacion-redondeo", 74),
    ("facturacion-otro", 6),
    ("entrega-sin-aviso", 28),
    ("entrega-retraso", 24),
    ("producto-defectuoso", 16),
    ("pedido-erroneo", 18),
    ("cambio-datos", 12),
    ("informacion-producto", 14),
    ("otros", 8),
)
PATOLOGIAS = (
    ("asunto-cuerpo-intercambiados", 34),
    ("hilo-roto", 28),
    ("adjunto-fantasma", 22),
    ("dos-incidencias", 15),
    ("firma-ajena", 9),
    ("reenvio-en-cadena", 12),
)
CORREOS_V4 = {"CLI-0042": 8, "CLI-0118": 7, "CLI-0233": 7}


def generar_correos(az: random.Random, clientes: list[Cliente],
                    tickets: list[Ticket]) -> list[Correo]:
    por_id = {c.id_cliente: c for c in clientes}
    con_email = [c for c in clientes if c.email_valido]
    hosteleria_email = [c for c in con_email if c.tipo_real == "hosteleria"]

    categorias: list[str] = []
    for categoria, cuantos in REPARTO_CORREOS:
        categorias.extend([categoria] * cuantos)
    assert len(categorias) == N_CORREOS
    az.shuffle(categorias)

    patologias: list[str] = []
    for nombre, cuantos in PATOLOGIAS:
        patologias.extend([nombre] * cuantos)
    patologias.extend([""] * (N_CORREOS - len(patologias)))
    az.shuffle(patologias)

    # Los tres de V4 sobrerrepresentados.
    remitentes: list[str] = []
    for id_cliente, cuantos in sorted(CORREOS_V4.items()):
        remitentes.extend([id_cliente] * cuantos)
    faltan = N_CORREOS - len(remitentes)
    otros = [c.id_cliente for c in con_email if c.id_cliente not in CLIENTES_V4]
    remitentes.extend(az.choices(otros, k=faltan))
    az.shuffle(remitentes)

    dias = dias_laborables(FECHA_INICIO, FECHA_FIN)
    correos: list[Correo] = []

    for n in range(N_CORREOS):
        categoria = categorias[n]
        patologia = patologias[n]
        id_cliente = remitentes[n]
        cliente = por_id[id_cliente]

        # Coherencia: los sin-aviso son siempre de hostelería.
        if categoria == "entrega-sin-aviso" and cliente.tipo_real != "hosteleria":
            cliente = az.choice(hosteleria_email)
            id_cliente = cliente.id_cliente
        if not cliente.email_valido:
            cliente = az.choice(con_email)
            id_cliente = cliente.id_cliente
        # La firma que no coincide con el remitente solo tiene sentido en un
        # negocio: el correo sale de la cuenta del bar y lo firma el encargado.
        if patologia == "firma-ajena" and cliente.tipo_real != "hosteleria":
            cliente = az.choice(hosteleria_email)
            id_cliente = cliente.id_cliente

        dia = az.choice(dias)
        momento = datetime(dia.year, dia.month, dia.day,
                           az.randint(7, 21), az.randint(0, 59), az.randint(0, 59))
        contexto = {"dto": cliente.descuento_pct, "poblacion": cliente.poblacion}

        cuerpo_incidencia = redactar(az, categoria, contexto,
                                     frases=elegir_pesado(az, (1, 2, 3), (0.30, 0.45, 0.25)))
        asunto_normal = rellenar(az, az.choice(ASUNTOS[categoria]), contexto)

        asunto = asunto_normal
        cuerpo = cuerpo_incidencia
        responde_a = ""

        if patologia == "asunto-cuerpo-intercambiados":
            if az.random() < 0.5:
                asunto = az.choice(ASUNTOS_MUDOS)
            else:
                asunto = cuerpo_incidencia[:78].rsplit(" ", 1)[0].rstrip(",.;")
                cuerpo = az.choice((
                    "Según lo hablado.", "Lo dicho por teléfono.",
                    "Como te comenté esta mañana.", "Ahí lo tienes.",
                    "Es lo de siempre, ya sabes.",
                ))
        elif patologia == "hilo-roto":
            asunto = "RE: " + asunto_normal
            responde_a = f"<AdN-{az.randint(7000, 9999):04d}@{DOMINIO_EMPRESA}>"
        elif patologia == "adjunto-fantasma":
            cuerpo += az.choice((
                "\n\nTe adjunto la factura para que la veas.",
                "\n\nOs mando la foto del albaran adjunta.",
                "\n\nAdjunto copia del albaran firmado.",
                "\n\nVa adjunto el papel que me dio el repartidor.",
            ))
        elif patologia == "dos-incidencias":
            segunda = az.choice([c for c, _ in REPARTO_CORREOS if c != categoria])
            cuerpo += "\n\n" + redactar(az, segunda, contexto, frases=1)
        elif patologia == "firma-ajena":
            cuerpo += "\n\n" + az.choice(FIRMAS_AJENAS)
        elif patologia == "reenvio-en-cadena":
            asunto = "RV: RE: RE: " + asunto_normal
            historial = (
                f"\n\n-----Mensaje original-----\n"
                f"> De: {BUZON_EMPRESA}\n"
                f"> Enviado: {momento.day - 1 if momento.day > 1 else 1}/"
                f"{momento.month}/{momento.year}\n"
                f"> Asunto: RE: {asunto_normal}\n"
                f">\n"
                f"> Buenos dias, lo miramos y te decimos algo.\n"
                f">\n"
                f">> De: {cliente.email}\n"
                f">> Asunto: {asunto_normal}\n"
                f">>\n"
                f">> {redactar(az, categoria, contexto, frases=1)}\n"
                f">>\n"
                f">>> De: {BUZON_EMPRESA}\n"
                f">>> Asunto: Su pedido\n"
                f">>>\n"
                f">>> Buenos dias, le confirmamos el reparto de esta semana.\n"
            )
            cuerpo = az.choice((
                "Os reenvio esto que sigue sin resolverse.",
                "Mirad el hilo, esto viene de largo.",
                "Reenvio el correo, nadie me ha contestado.",
            )) + historial + "\n" + cuerpo_incidencia

        if not cuerpo.endswith("\n"):
            cuerpo += "\n"
        despedida = az.choice(("", "", "\nGracias.\n", "\nUn saludo.\n",
                               "\nGracias, un saludo.\n", "\nHasta luego.\n"))
        if patologia != "firma-ajena" and despedida:
            cuerpo += despedida

        correos.append(Correo(
            numero=n + 1,
            id_mensaje=f"<AdN-{n + 1:04d}.{id_cliente}@{DOMINIO_EMPRESA}>",
            de_nombre=arreglar_mojibake(cliente.nombre),
            de_email=cliente.email, asunto=asunto, cuerpo=cuerpo, momento=momento,
            responde_a=responde_a, patologia=patologia, id_cliente=id_cliente,
            categoria=categoria,
        ))

    correos.sort(key=lambda c: (c.momento, c.numero))
    return correos


# ═══════════════════════════════════════════════════════════════════════════════
# 8. ESCRITURA DE FICHEROS
# ═══════════════════════════════════════════════════════════════════════════════

MOMENTO_FIJO = datetime(2025, 3, 3, 9, 0, 0)   # metadatos fijos → reproducibilidad


def _preparar_libro(titulo: str) -> Workbook:
    libro = Workbook()
    libro.properties.creator = EMPRESA
    libro.properties.lastModifiedBy = EMPRESA
    libro.properties.created = MOMENTO_FIJO
    libro.properties.modified = MOMENTO_FIJO
    libro.properties.title = titulo
    return libro


def _cabecera(hoja, columnas: list[str], anchos: list[int]) -> None:
    hoja.append(columnas)
    for celda in hoja[1]:
        celda.font = Font(bold=True)
    for n, ancho in enumerate(anchos, start=1):
        hoja.column_dimensions[chr(64 + n) if n <= 26 else "A"].width = ancho
    hoja.freeze_panes = "A2"


def escribir_clientes(ruta: Path, clientes: list[Cliente]) -> None:
    libro = _preparar_libro("Maestro de clientes")
    hoja = libro.active
    hoja.title = "Clientes"
    columnas = ["id_cliente", "nombre", "tipo", "direccion", "poblacion", "cp",
                "telefono", "email", "fecha_alta", "ruta", "forma_pago",
                "descuento_pct", "observaciones"]
    _cabecera(hoja, columnas, [11, 34, 13, 30, 24, 8, 17, 32, 12, 9, 15, 14, 34])
    for c in clientes:
        hoja.append([
            c.id_cliente, c.nombre, c.tipo, c.direccion, c.poblacion,
            int(c.cp) if c.cp_como_numero else c.cp,
            c.telefono, c.email, c.fecha_alta, c.ruta, c.forma_pago,
            c.descuento_celda, c.observaciones,
        ])
    for fila in hoja.iter_rows(min_row=2, min_col=9, max_col=9):
        fila[0].number_format = "DD/MM/YYYY"
    libro.save(ruta)


def escribir_pedidos(ruta: Path, pedidos: list[Pedido]) -> int:
    libro = _preparar_libro("Histórico de pedidos")
    hoja = libro.active
    hoja.title = "Pedidos"
    columnas = ["id_pedido", "fecha", "id_cliente", "producto", "descripcion",
                "cantidad", "precio_ud", "dto_pct", "importe_linea", "total_pedido"]
    _cabecera(hoja, columnas, [12, 12, 11, 10, 32, 10, 11, 9, 14, 14])
    filas = 0
    for pedido in sorted(pedidos, key=lambda p: (p.fecha, p.id_pedido)):
        total = float(pedido.total)
        for linea in pedido.lineas:
            hoja.append([
                pedido.id_pedido, pedido.fecha.isoformat(), pedido.id_cliente,
                linea.producto, PRODUCTOS[linea.producto]["descripcion"],
                linea.cantidad, float(linea.precio_ud), linea.dto_pct,
                float(linea.importe), total,
            ])
            filas += 1
    for fila in hoja.iter_rows(min_row=2, min_col=7, max_col=10):
        for celda in fila:
            if celda.column in (7, 9, 10):
                celda.number_format = "0.00"
    libro.save(ruta)
    return filas


def escribir_tickets(ruta: Path, tickets: list[Ticket]) -> None:
    libro = _preparar_libro("Registro de incidencias")
    hoja = libro.active
    hoja.title = "Tickets"
    columnas = ["id_ticket", "fecha_apertura", "canal", "id_cliente", "categoria",
                "subcategoria", "descripcion", "id_pedido", "estado",
                "fecha_cierre", "agente", "tiempo_dedicado_min"]
    _cabecera(hoja, columnas, [11, 15, 12, 11, 22, 20, 80, 12, 12, 14, 15, 20])
    for t in tickets:
        cierre = (formatear_fecha(t.fecha_cierre, t.formato_fecha)
                  if isinstance(t.fecha_cierre, date) else "")
        hoja.append([
            t.id_ticket, formatear_fecha(t.fecha_apertura, t.formato_fecha), t.canal,
            t.id_cliente_visible, t.categoria_sucia, t.subcategoria, t.descripcion,
            t.id_pedido, t.estado, cierre, t.agente, t.tiempo,
        ])
    libro.save(ruta)


def _cuerpo_eml(correo: Correo) -> str:
    """Un .eml RFC 5322 en UTF-8, legible en un editor de texto."""
    cabeceras = [
        "MIME-Version: 1.0",
        f"Date: {cabecera_fecha(correo.momento)}",
        f"From: {correo.de_nombre} <{correo.de_email}>",
        f"To: Atencion al cliente <{BUZON_EMPRESA}>",
        f"Subject: {correo.asunto}",
        f"Message-ID: {correo.id_mensaje}",
    ]
    if correo.responde_a:
        cabeceras.append(f"In-Reply-To: {correo.responde_a}")
        cabeceras.append(f"References: {correo.responde_a}")
    cabeceras.append('Content-Type: text/plain; charset="utf-8"')
    cabeceras.append("Content-Transfer-Encoding: 8bit")
    return "\n".join(cabeceras) + "\n\n" + correo.cuerpo


def escribir_correos(carpeta: Path, ruta_mbox: Path, correos: list[Correo]) -> None:
    carpeta.mkdir(parents=True, exist_ok=True)
    for antiguo in sorted(carpeta.glob("*.eml")):
        antiguo.unlink()

    trozos_mbox: list[str] = []
    for correo in correos:
        texto = _cuerpo_eml(correo)
        nombre = f"correo-{correo.numero:03d}.eml"
        (carpeta / nombre).write_text(texto, encoding="utf-8", newline="\n")

        escapado = "\n".join(
            (">" + linea if re.match(r"^>*From ", linea) else linea)
            for linea in texto.split("\n"))
        trozos_mbox.append(
            f"From {correo.de_email} {linea_from_mbox(correo.momento)}\n{escapado}\n")
    ruta_mbox.write_text("\n".join(trozos_mbox), encoding="utf-8", newline="\n")


# ── procedimientos.docx ────────────────────────────────────────────────────────

PROCEDIMIENTOS = (
    ("PR-01", "Alta de cliente nuevo", (
        "1.1. La solicitud de alta se recibe por teléfono, por correo electrónico o "
        "a través del comercial de zona.",
        "1.2. Antes de crear la ficha, el responsable administrativo comprobará que "
        "el solicitante no consta ya en el fichero maestro. La comprobación se "
        "realizará por número de teléfono y, en su defecto, por domicilio de "
        "entrega. No se dará de alta ninguna ficha cuyo teléfono coincida con el "
        "de una ficha existente.",
        "1.3. Se cumplimentarán todos los campos del maestro. Los campos de "
        "descuento y forma de pago requieren la conformidad de Gerencia.",
        "1.4. El alta se comunicará al responsable de ruta antes del primer reparto.",
        "1.5. La documentación soporte se archivará en el expediente del cliente "
        "durante el plazo legalmente establecido.",
    )),
    ("PR-02", "Toma de pedido por teléfono", (
        "2.1. Los pedidos se reciben en horario de oficina, de 9:00 a 14:00 y de "
        "16:00 a 18:00, de lunes a viernes.",
        "2.2. El operador identificará al cliente por su código o, en su defecto, "
        "por el nombre comercial del establecimiento.",
        "2.3. Se anotará producto, formato y cantidad, y se repetirá el pedido "
        "completo al cliente antes de cerrar la llamada.",
        "2.4. Los pedidos recibidos después de las 12:00 se sirven en el reparto "
        "siguiente.",
    )),
    ("PR-03", "Preparación y carga de ruta", (
        "3.1. La preparación de la carga se realiza el día anterior al reparto.",
        "3.2. La distribución se organiza en tres rutas: ruta 1 (bahía), ruta 2 "
        "(interior) y ruta 3 (costa oriental).",
        "3.3. El responsable de almacén verificará la coincidencia entre la hoja de "
        "carga y el albarán de cada cliente.",
        "3.4. Las incidencias de carga se comunicarán al Departamento de Calidad "
        "para su registro y análisis.",
    )),
    ("PR-04", "Entrega y firma de albarán", (
        "4.1. El reparto se efectúa en horario de 8:00 a 16:00.",
        "4.2. El repartidor entregará la mercancía y recabará la firma del cliente "
        "o persona autorizada en el albarán.",
        "4.3. En caso de ausencia, el repartidor no depositará la mercancía sin "
        "autorización expresa y anotará la incidencia en la hoja de ruta.",
        "4.4. Los envases retornables se recogerán en el mismo acto de entrega.",
    )),
    ("PR-05", "Devoluciones y abonos", (
        "5.1. Toda devolución de mercancía generará un albarán de abono con "
        "numeración propia, independiente de la de los albaranes de venta.",
        "5.2. El albarán de abono indicará el motivo de la devolución conforme a la "
        "tabla de motivos del anexo II.",
        "5.3. En ningún caso se registrará una devolución como un pedido de venta "
        "con signo negativo.",
        "5.4. Los abonos se aplicarán en la factura del mes en curso.",
    )),
    ("PR-06", "Facturación mensual", (
        "6.1. La facturación se emite el primer día hábil del mes siguiente al "
        "periodo facturado.",
        "6.2. El importe de cada línea se calcula multiplicando el precio de tarifa "
        "por las unidades servidas. Sobre el importe así obtenido se aplica el "
        "descuento comercial que tenga asignado el cliente. El resultado se "
        "redondea a dos decimales una vez efectuada la operación completa, y nunca "
        "antes.",
        "6.3. No se aplicará el descuento sobre el precio unitario de forma "
        "independiente: el descuento es un porcentaje del importe de la línea.",
        "6.4. Las facturas se remiten por correo electrónico y, a petición del "
        "cliente, en papel.",
        "6.5. Las discrepancias de importe se resolverán en un plazo máximo de "
        "cinco días hábiles.",
    )),
    ("PR-07", "Aviso previo de entrega", (
        "7.1. El día anterior al reparto, el responsable de atención al cliente "
        "comunicará a cada cliente de la ruta la franja horaria estimada de "
        "entrega.",
        "7.2. El aviso previo podrá realizarse por teléfono o por mensaje escrito, "
        "y quedará registrado en la hoja de ruta correspondiente.",
        "7.3. Si el cliente comunica que no podrá recibir la mercancía, se "
        "reprogramará la entrega al siguiente reparto de la zona.",
        "7.4. La omisión del aviso previo se considerará incidencia de servicio y "
        "se comunicará al Departamento de Calidad.",
    )),
    ("PR-08", "Reclamaciones", (
        "8.1. Las reclamaciones se registrarán en el sistema de incidencias el "
        "mismo día de su recepción, con indicación del canal de entrada.",
        "8.2. Toda reclamación se clasificará conforme a la tabla de tipologías del "
        "anexo III.",
        "8.3. El plazo máximo de respuesta al cliente es de cuarenta y ocho horas.",
        "8.4. Con periodicidad mensual se elaborará un informe de reclamaciones que "
        "será elevado a Gerencia.",
    )),
    ("PR-09", "Actualización de datos de cliente", (
        "9.1. Cualquier modificación de los datos identificativos, fiscales o de "
        "contacto del cliente se recogerá en la ficha existente.",
        "9.2. No se creará una ficha nueva para reflejar un cambio de titularidad, "
        "de domicilio o de denominación social. La ficha es única por cliente.",
        "9.3. Las modificaciones quedarán registradas con fecha y responsable.",
        "9.4. Con periodicidad semestral se revisará el fichero maestro para "
        "detectar registros duplicados u obsoletos.",
    )),
)

PARRAFO_HUECO = (
    "El presente procedimiento se enmarca dentro del sistema de gestión de la "
    "compañía y será de aplicación a todo el personal afectado por su alcance, sin "
    "perjuicio de las instrucciones particulares que en cada caso se establezcan. "
    "Su cumplimiento es responsabilidad de los intervinientes."
)


def escribir_procedimientos(ruta: Path) -> None:
    documento = docx.Document()
    estilo = documento.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(10)

    portada = documento.add_paragraph()
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trozo = portada.add_run(f"{EMPRESA}\n")
    trozo.bold = True
    trozo.font.size = Pt(20)
    trozo = portada.add_run("\nMANUAL DE PROCEDIMIENTOS INTERNOS\n")
    trozo.bold = True
    trozo.font.size = Pt(16)
    trozo = portada.add_run("\nv3 — revisado en marzo de 2019\n")
    trozo.font.size = Pt(12)
    portada.add_run("\nDocumento de uso interno. Prohibida su reproducción total o "
                    "parcial sin autorización de Gerencia.\n")
    documento.add_paragraph("Elaborado por: Asesoría y Gestión Besaya, S.L.")
    documento.add_paragraph("Aprobado por: Gerencia")
    documento.add_paragraph("Próxima revisión prevista: marzo de 2021")
    documento.add_page_break()

    documento.add_heading("Índice", level=1)
    tabla = documento.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"
    encabezado = tabla.rows[0].cells
    encabezado[0].text = "Código"
    encabezado[1].text = "Procedimiento"
    encabezado[2].text = "Página"
    for n, (codigo, titulo, _) in enumerate(PROCEDIMIENTOS, start=1):
        fila = tabla.add_row().cells
        fila[0].text = codigo
        fila[1].text = titulo
        fila[2].text = str(n + 2)
    documento.add_paragraph()

    documento.add_heading("0. Objeto y alcance", level=1)
    documento.add_paragraph(
        "El presente manual recoge los procedimientos operativos de "
        f"{EMPRESA}, dedicada a la distribución de agua envasada y bebidas a "
        "establecimientos de hostelería y clientes particulares en Cantabria y el "
        "oriente de Asturias.")
    documento.add_paragraph(PARRAFO_HUECO)
    documento.add_paragraph(
        "Las dudas de interpretación se elevarán al Departamento de Calidad, que "
        "resolverá conforme a los criterios generales de la compañía.")
    documento.add_page_break()

    for codigo, titulo, apartados in PROCEDIMIENTOS:
        documento.add_heading(f"{codigo} — {titulo}", level=1)
        documento.add_heading("Objeto", level=2)
        documento.add_paragraph(
            f"Establecer la sistemática aplicable a «{titulo.lower()}» dentro de la "
            "operativa habitual de la compañía.")
        documento.add_heading("Desarrollo", level=2)
        for apartado in apartados:
            documento.add_paragraph(apartado)
        documento.add_heading("Registros", level=2)
        documento.add_paragraph(
            "Los registros derivados de este procedimiento se conservarán conforme a "
            "lo indicado en el apartado 0 del presente manual.")
        documento.add_paragraph(PARRAFO_HUECO)

    documento.add_heading("Anexo I — Organigrama", level=1)
    documento.add_paragraph("Gerencia")
    documento.add_paragraph("Departamento de Administración")
    documento.add_paragraph("Departamento de Atención al Cliente")
    documento.add_paragraph("Departamento de Logística y Reparto")
    documento.add_paragraph("Departamento de Calidad")

    documento.core_properties.author = "Asesoría y Gestión Besaya, S.L."
    documento.core_properties.title = "Manual de procedimientos internos"
    documento.core_properties.created = datetime(2019, 3, 18, 12, 0, 0)
    documento.core_properties.modified = datetime(2019, 3, 18, 12, 0, 0)
    documento.core_properties.revision = 3
    documento.save(ruta)


# ═══════════════════════════════════════════════════════════════════════════════
# 9. SOLUCIONES
# ═══════════════════════════════════════════════════════════════════════════════

def escribir_csv(ruta: Path, cabecera: list[str], filas: list[list[str]]) -> None:
    lineas = [",".join(cabecera)]
    for fila in filas:
        campos = []
        for valor in fila:
            texto = str(valor)
            if any(c in texto for c in (",", '"', "\n")):
                texto = '"' + texto.replace('"', '""') + '"'
            campos.append(texto)
        lineas.append(",".join(campos))
    ruta.write_text("\n".join(lineas) + "\n", encoding="utf-8", newline="\n")


def euros(valor: Decimal) -> str:
    """Formato español: 1.234,56 €."""
    texto = f"{valor.quantize(CENT, rounding=ROUND_HALF_UP):,.2f}"
    entero, decimales = texto.split(".")
    return f"{entero.replace(',', '.')},{decimales} €"


def pct(valor: float) -> str:
    return f"{valor:.2f}".replace(".", ",") + " %"


def margen_bruto(linea: Linea) -> Decimal:
    """Margen bruto de la línea: lo facturado menos el coste de la mercancía."""
    coste_ud = (linea.precio_ud * (Decimal(1) - PRODUCTOS[linea.producto]["margen"]))
    return linea.importe - (coste_ud * Decimal(linea.cantidad)).quantize(CENT,
                                                                        rounding=ROUND_HALF_UP)


# ═══════════════════════════════════════════════════════════════════════════════
# 10. PROGRAMA PRINCIPAL
# ═══════════════════════════════════════════════════════════════════════════════

def construir(base: Path, silencioso: bool = False) -> dict:
    ficheros = base / "ficheros"
    soluciones = base / "SOLUCIONES"
    ficheros.mkdir(parents=True, exist_ok=True)
    soluciones.mkdir(parents=True, exist_ok=True)

    az = random.Random(SEMILLA)

    clientes, pares_duplicados = generar_clientes(az)
    pedidos, fantasma = generar_pedidos(az, clientes)
    tickets = generar_tickets(az, clientes, pedidos)
    correos = generar_correos(az, clientes, tickets)

    por_id = {c.id_cliente: c for c in clientes}
    reales = [p for p in pedidos if not p.es_fantasma]

    # ── Comprobaciones internas antes de escribir nada ────────────────────────
    lineas_dto_cero = [l for p in pedidos for l in p.lineas if l.dto_pct == 0]
    lineas_dto = [l for p in pedidos for l in p.lineas if l.dto_pct > 0]
    assert all(l.desviacion == 0 for l in lineas_dto_cero), \
        "Hay líneas sin descuento con desviación: el fallo se estaría filtrando"
    assert all(l.desviacion >= 0 for l in lineas_dto), \
        "Hay líneas con descuento y desviación negativa"
    con_desviacion = sum(1 for l in lineas_dto if l.desviacion > 0)
    ratio_desviacion = con_desviacion / len(lineas_dto)
    assert ratio_desviacion >= 0.85, \
        f"Solo el {ratio_desviacion:.1%} de las líneas con descuento se desvía (mínimo 85 %)"

    textos_alumna = [t.descripcion for t in tickets] + \
                    [c.asunto + " " + c.cuerpo for c in correos]
    for texto in textos_alumna:
        base_texto = quitar_acentos(texto).lower()
        assert "redonde" not in base_texto, f"Un texto nombra el hallazgo: {texto!r}"
        assert "aviso previo" not in base_texto, f"Un texto nombra PR-07: {texto!r}"

    tickets_con_pedido = {t.id_pedido for t in tickets if t.id_pedido}
    ids_fantasma = {p.id_pedido for p in fantasma}
    assert not (tickets_con_pedido & ids_fantasma), \
        "Un ticket apunta a un pedido fantasma: los fantasma no generan tickets"
    mapa_pedidos = {p.id_pedido: p for p in pedidos}
    for t in tickets:
        if not t.id_pedido:
            continue
        pedido = mapa_pedidos[t.id_pedido]
        assert pedido.id_cliente == t.id_cliente, \
            f"{t.id_ticket} apunta a un pedido de otro cliente"
        assert pedido.fecha <= t.fecha_apertura, \
            f"{t.id_ticket} apunta a un pedido posterior a la incidencia"
        if t.categoria_real == "facturacion-redondeo":
            assert pedido.lineas[0].dto_pct > 0 and pedido.desviacion > 0, \
                f"{t.id_ticket} es de V1 y su pedido no tiene desviación"

    emails_validos = {c.email for c in clientes if c.email_valido}
    for correo in correos:
        assert correo.de_email in emails_validos, \
            f"El correo {correo.numero} sale de una dirección que no está en el maestro"

    # ── Escritura de los ficheros de la alumna ────────────────────────────────
    escribir_clientes(ficheros / "clientes.xlsx", clientes)
    filas_pedidos = escribir_pedidos(ficheros / "pedidos.xlsx", pedidos)
    escribir_tickets(ficheros / "tickets.xlsx", tickets)
    escribir_correos(ficheros / "correos", ficheros / "bandeja.mbox", correos)
    escribir_procedimientos(ficheros / "procedimientos.docx")
    escribir_leeme(ficheros / "LEEME.md")

    metricas = escribir_soluciones(soluciones, clientes, pares_duplicados, pedidos,
                                   fantasma, tickets, correos, filas_pedidos)

    if not silencioso:
        resumen(base, metricas)
    return metricas


# ── Cálculo de las cinco verdades y volcado de SOLUCIONES/ ────────────────────

def escribir_soluciones(soluciones: Path, clientes: list[Cliente],
                        pares: list[tuple[str, str]], pedidos: list[Pedido],
                        fantasma: list[Pedido], tickets: list[Ticket],
                        correos: list[Correo], filas_pedidos: int) -> dict:
    por_id = {c.id_cliente: c for c in clientes}
    reales = [p for p in pedidos if not p.es_fantasma]
    mapa_pedidos = {p.id_pedido: p for p in pedidos}

    # ── V1 ────────────────────────────────────────────────────────────────────
    lineas_dto = [l for p in pedidos for l in p.lineas if l.dto_pct > 0]
    lineas_dto_reales = [l for p in reales for l in p.lineas if l.dto_pct > 0]
    lineas_totales = sum(len(p.lineas) for p in pedidos)
    desviacion_total = sum((l.desviacion for p in reales for l in p.lineas), Decimal("0.00"))
    desviacion_fantasma = sum((l.desviacion for p in fantasma for l in p.lineas),
                              Decimal("0.00"))
    pedidos_con_desviacion = [p for p in reales if p.desviacion > 0]
    desviacion_maxima = max((l.desviacion for p in reales for l in p.lineas),
                            default=Decimal("0.00"))
    facturado_total = sum((l.importe for p in reales for l in p.lineas), Decimal("0.00"))

    tickets_v1 = [t for t in tickets if t.categoria_real == "facturacion-redondeo"]
    tickets_v1_con_pedido = [t for t in tickets_v1 if t.id_pedido]
    tickets_fact_otro = [t for t in tickets if t.subtipo == "facturacion-otro"]

    # ── V2 ────────────────────────────────────────────────────────────────────
    pares_detalle = []
    for id_orig, id_dup in pares:
        c_o, c_d = por_id[id_orig], por_id[id_dup]
        pedidos_o = sum(1 for p in reales if p.id_cliente == id_orig)
        pedidos_d = sum(1 for p in reales if p.id_cliente == id_dup)
        pares_detalle.append({
            "original": id_orig, "duplicado": id_dup,
            "nombre_original": c_o.nombre, "nombre_duplicado": c_d.nombre,
            "telefono_original": c_o.telefono, "telefono_duplicado": c_d.telefono,
            "telefono_norm": normalizar_telefono(c_o.telefono),
            "direccion_original": c_o.direccion, "direccion_duplicado": c_d.direccion,
            "direccion_norm": normalizar_direccion(c_o.direccion),
            "dto_original": c_o.descuento_pct, "dto_duplicado": c_d.descuento_pct,
            "dto_distinto": c_o.descuento_pct != c_d.descuento_pct,
            "pedidos_original": pedidos_o, "pedidos_duplicado": pedidos_d,
        })
    pares_dto_distinto = [p for p in pares_detalle if p["dto_distinto"]]

    # ── V3 ────────────────────────────────────────────────────────────────────
    def lineas_mes(coleccion, anio, mes):
        return sum(len(p.lineas) for p in coleccion if p.fecha.year == anio and p.fecha.month == mes)

    def pedidos_mes(coleccion, anio, mes):
        return sum(1 for p in coleccion if p.fecha.year == anio and p.fecha.month == mes)

    meses = sorted(PRESUPUESTO_LINEAS.keys())
    lineas_por_mes_real = {m: lineas_mes(reales, *m) for m in meses}
    lineas_por_mes_reg = {m: lineas_mes(pedidos, *m) for m in meses}
    pedidos_por_mes_real = {m: pedidos_mes(reales, *m) for m in meses}
    pedidos_por_mes_reg = {m: pedidos_mes(pedidos, *m) for m in meses}

    otros_meses = [m for m in meses if m != (2024, 12)]
    base_lineas = sum(lineas_por_mes_real[m] for m in otros_meses) / len(otros_meses)
    base_pedidos = sum(pedidos_por_mes_real[m] for m in otros_meses) / len(otros_meses)
    subida_real_lineas = lineas_por_mes_real[(2024, 12)] / base_lineas - 1
    subida_reg_lineas = lineas_por_mes_reg[(2024, 12)] / base_lineas - 1
    subida_real_pedidos = pedidos_por_mes_real[(2024, 12)] / base_pedidos - 1
    subida_reg_pedidos = pedidos_por_mes_reg[(2024, 12)] / base_pedidos - 1
    importe_fantasma = sum((p.total for p in fantasma), Decimal("0.00"))

    # ── V4 ────────────────────────────────────────────────────────────────────
    conteo_tickets: dict[str, int] = {}
    for t in tickets:
        conteo_tickets[t.id_cliente] = conteo_tickets.get(t.id_cliente, 0) + 1
    cuentas_v4 = []
    for id_cliente in CLIENTES_V4:
        suyos = [p for p in reales if p.id_cliente == id_cliente]
        facturado = sum((l.importe for p in suyos for l in p.lineas), Decimal("0.00"))
        margen = sum((margen_bruto(l) for p in suyos for l in p.lineas), Decimal("0.00"))
        n_tickets = conteo_tickets[id_cliente]
        coste = COSTE_CONTACTO * n_tickets
        cuentas_v4.append({
            "id_cliente": id_cliente,
            "nombre": por_id[id_cliente].nombre,
            "poblacion": por_id[id_cliente].poblacion,
            "descuento_pct": por_id[id_cliente].descuento_pct,
            "pedidos": len(suyos),
            "lineas": sum(len(p.lineas) for p in suyos),
            "facturado": facturado,
            "margen_bruto": margen,
            "tickets": n_tickets,
            "coste_contactos": coste,
            "resultado": margen - coste,
        })
    tickets_v4 = sum(c["tickets"] for c in cuentas_v4)

    facturado_hosteleria = Decimal("0.00")
    n_host = 0
    for c in clientes:
        if c.tipo_real != "hosteleria":
            continue
        n_host += 1
        facturado_hosteleria += sum(
            (l.importe for p in reales if p.id_cliente == c.id_cliente for l in p.lineas),
            Decimal("0.00"))
    media_hosteleria = facturado_hosteleria / Decimal(n_host)

    # ── V5 ────────────────────────────────────────────────────────────────────
    tickets_v5 = [t for t in tickets if t.categoria_real == "entrega-sin-aviso"]
    v5_hosteleria = all(por_id[t.id_cliente].tipo_real == "hosteleria" for t in tickets_v5)
    coste_v5 = COSTE_CONTACTO * len(tickets_v5)
    coste_v1 = COSTE_CONTACTO * len(tickets_v1)

    # ── Ficheros CSV ──────────────────────────────────────────────────────────
    escribir_csv(soluciones / "taxonomia-real.csv",
                 ["id_ticket", "categoria_sucia", "categoria_real"],
                 [[t.id_ticket, t.categoria_sucia, t.categoria_real] for t in tickets])

    escribir_csv(
        soluciones / "mapa-duplicados.csv",
        ["id_original", "nombre_original", "id_duplicado", "nombre_duplicado",
         "telefono_original", "telefono_duplicado", "telefono_normalizado",
         "direccion_original", "direccion_duplicado", "direccion_normalizada",
         "dto_original", "dto_duplicado", "dto_distinto",
         "pedidos_original", "pedidos_duplicado"],
        [[p["original"], p["nombre_original"], p["duplicado"], p["nombre_duplicado"],
          p["telefono_original"], p["telefono_duplicado"], p["telefono_norm"],
          p["direccion_original"], p["direccion_duplicado"], p["direccion_norm"],
          p["dto_original"], p["dto_duplicado"], "si" if p["dto_distinto"] else "no",
          p["pedidos_original"], p["pedidos_duplicado"]] for p in pares_detalle])

    escribir_csv(
        soluciones / "pedidos-fantasma.csv",
        ["id_pedido_fantasma", "id_pedido_original", "fecha", "id_cliente",
         "lineas", "importe_total"],
        [[p.id_pedido, p.original_de, p.fecha.isoformat(), p.id_cliente,
          len(p.lineas), f"{p.total:.2f}"] for p in fantasma])

    escribir_csv(
        soluciones / "cuentas-v4.csv",
        ["id_cliente", "nombre", "poblacion", "descuento_pct", "pedidos", "lineas",
         "facturado_6m", "margen_bruto_6m", "tickets", "coste_contactos", "resultado"],
        [[c["id_cliente"], c["nombre"], c["poblacion"], c["descuento_pct"],
          c["pedidos"], c["lineas"], f"{c['facturado']:.2f}", f"{c['margen_bruto']:.2f}",
          c["tickets"], f"{c['coste_contactos']:.2f}", f"{c['resultado']:.2f}"]
         for c in cuentas_v4])

    # ── Métricas ──────────────────────────────────────────────────────────────
    conteo_sucias: dict[str, int] = {}
    for t in tickets:
        clave = t.categoria_sucia if t.categoria_sucia else "(vacío)"
        conteo_sucias[clave] = conteo_sucias.get(clave, 0) + 1
    bucket_facturacion = [t for t in tickets if t.categoria_sucia in ETIQUETAS_FACTURACION]
    puros = [t for t in bucket_facturacion if t.categoria_real == "facturacion-redondeo"]
    estimacion_ingenua = len(bucket_facturacion) / len(tickets)
    pureza = len(puros) / len(bucket_facturacion)

    metricas = {
        "semilla": SEMILLA,
        "periodo": {"inicio": FECHA_INICIO.isoformat(), "fin": FECHA_FIN.isoformat(),
                    "dias_laborables": len(dias_laborables(FECHA_INICIO, FECHA_FIN))},
        "clientes": {
            "filas": len(clientes),
            "reales": len(clientes) - N_DUPLICADOS,
            "duplicados": N_DUPLICADOS,
            "hosteleria": sum(1 for c in clientes if c.tipo_real == "hosteleria"),
            "particular": sum(1 for c in clientes if c.tipo_real == "particular"),
            "sin_email": sum(1 for c in clientes if not c.email),
            "email_invalido": sum(1 for c in clientes if c.email and not c.email_valido),
            "tipo_vacio": sum(1 for c in clientes if not c.tipo),
            "cp_numerico": sum(1 for c in clientes if c.cp_como_numero),
            "ruta_vacia": sum(1 for c in clientes if not c.ruta),
            "descuento_vacio": sum(1 for c in clientes if c.descuento_celda == ""),
            "observaciones_vacias": sum(1 for c in clientes if not c.observaciones),
        },
        "pedidos": {
            "pedidos_totales": len(pedidos),
            "pedidos_reales": len(reales),
            "pedidos_fantasma": len(fantasma),
            "lineas_totales": lineas_totales,
            "filas_xlsx": filas_pedidos,
            "devoluciones": sum(1 for p in reales if p.es_devolucion),
            "facturado_total": f"{facturado_total:.2f}",
            "lineas_con_descuento": len(lineas_dto),
            "lineas_sin_descuento": lineas_totales - len(lineas_dto),
        },
        "v1": {
            "desviacion_total_reales": f"{desviacion_total:.2f}",
            "desviacion_fantasma": f"{desviacion_fantasma:.2f}",
            "desviacion_maxima_linea": f"{desviacion_maxima:.2f}",
            "pedidos_afectados": len(pedidos_con_desviacion),
            "lineas_con_descuento": len(lineas_dto),
            "lineas_con_descuento_y_desviacion":
                sum(1 for l in lineas_dto if l.desviacion > 0),
            "pct_lineas_dto_con_desviacion":
                sum(1 for l in lineas_dto if l.desviacion > 0) / len(lineas_dto),
            "tickets": len(tickets_v1),
            "tickets_pct": len(tickets_v1) / len(tickets),
            "tickets_con_id_pedido": len(tickets_v1_con_pedido),
            "tickets_con_id_pedido_pct": len(tickets_v1_con_pedido) / len(tickets_v1),
            "tickets_facturacion_no_redondeo": len(tickets_fact_otro),
            "coste_contactos": f"{coste_v1:.2f}",
        },
        "v2": {
            "pares": len(pares_detalle),
            "pares_con_descuento_distinto": len(pares_dto_distinto),
            "ambos_con_pedidos": sum(1 for p in pares_detalle
                                     if p["pedidos_original"] > 0 and p["pedidos_duplicado"] > 0),
        },
        "v3": {
            "pedidos_fantasma": len(fantasma),
            "lineas_fantasma": sum(len(p.lineas) for p in fantasma),
            "importe_fantasma": f"{importe_fantasma:.2f}",
            "ventana": [VENTANA_FANTASMA_INI.isoformat(), VENTANA_FANTASMA_FIN.isoformat()],
            "lineas_por_mes_real": {f"{a}-{m:02d}": v for (a, m), v in lineas_por_mes_real.items()},
            "lineas_por_mes_registrado": {f"{a}-{m:02d}": v for (a, m), v in lineas_por_mes_reg.items()},
            "pedidos_por_mes_real": {f"{a}-{m:02d}": v for (a, m), v in pedidos_por_mes_real.items()},
            "pedidos_por_mes_registrado": {f"{a}-{m:02d}": v for (a, m), v in pedidos_por_mes_reg.items()},
            "base_lineas_otros_meses": round(base_lineas, 2),
            "subida_diciembre_real_lineas": round(subida_real_lineas, 4),
            "subida_diciembre_registrada_lineas": round(subida_reg_lineas, 4),
            "subida_diciembre_real_pedidos": round(subida_real_pedidos, 4),
            "subida_diciembre_registrada_pedidos": round(subida_reg_pedidos, 4),
        },
        "v4": {
            "clientes": [c["id_cliente"] for c in cuentas_v4],
            "tickets": tickets_v4,
            "tickets_pct": tickets_v4 / len(tickets),
            "coste_total": f"{sum(c['coste_contactos'] for c in cuentas_v4):.2f}",
            "margen_total": f"{sum(c['margen_bruto'] for c in cuentas_v4):.2f}",
            "resultado_total": f"{sum(c['resultado'] for c in cuentas_v4):.2f}",
            "facturacion_media_hosteleria": f"{media_hosteleria:.2f}",
            "detalle": [{k: (f"{v:.2f}" if isinstance(v, Decimal) else v)
                         for k, v in c.items()} for c in cuentas_v4],
        },
        "v5": {
            "tickets": len(tickets_v5),
            "tickets_pct": len(tickets_v5) / len(tickets),
            "todos_hosteleria": v5_hosteleria,
            "coste_contactos": f"{coste_v5:.2f}",
            "menciones_aviso_previo_en_tickets": 0,
            "menciones_aviso_previo_en_correos": 0,
        },
        "tickets": {
            "filas": len(tickets),
            "taxonomia_real": {cat: sum(1 for t in tickets if t.categoria_real == cat)
                               for cat, _ in TAXONOMIA},
            "etiquetas_sucias": dict(sorted(conteo_sucias.items())),
            "etiquetas_distintas": len([k for k in conteo_sucias if k != "(vacío)"]),
            "bucket_facturacion": len(bucket_facturacion),
            "bucket_facturacion_pct": estimacion_ingenua,
            "bucket_facturacion_pureza": pureza,
            "error_de_fiarse_de_categoria":
                len(tickets_v1) / len(tickets) - estimacion_ingenua,
            "sin_id_cliente": sum(1 for t in tickets if not t.id_cliente_visible),
            "con_id_pedido": sum(1 for t in tickets if t.id_pedido),
            "con_id_pedido_pct": sum(1 for t in tickets if t.id_pedido) / len(tickets),
            "sin_subcategoria": sum(1 for t in tickets if not t.subcategoria),
            "sin_agente": sum(1 for t in tickets if not t.agente),
            "sin_tiempo": sum(1 for t in tickets if t.tiempo == ""),
            "cierre_anterior_a_apertura": sum(
                1 for t in tickets if isinstance(t.fecha_cierre, date)
                and t.fecha_cierre < t.fecha_apertura),
        },
        "correos": {
            "total": len(correos),
            "patologias": {nombre: sum(1 for c in correos if c.patologia == nombre)
                           for nombre, _ in PATOLOGIAS},
            "sin_patologia": sum(1 for c in correos if not c.patologia),
            "de_clientes_v4": {k: sum(1 for c in correos if c.id_cliente == k)
                               for k in CLIENTES_V4},
            "por_categoria": {cat: sum(1 for c in correos if c.categoria == cat)
                              for cat, _ in REPARTO_CORREOS},
        },
    }
    (soluciones / "metricas-generacion.json").write_text(
        json.dumps(metricas, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8", newline="\n")

    escribir_verdades(soluciones / "verdades-escondidas.md", metricas, cuentas_v4,
                      pares_detalle, desviacion_total, desviacion_maxima,
                      facturado_total, media_hosteleria, tickets_v1,
                      tickets_fact_otro, tickets_v5, conteo_sucias,
                      bucket_facturacion, pureza, estimacion_ingenua,
                      lineas_por_mes_real, lineas_por_mes_reg, base_lineas,
                      subida_real_lineas, subida_reg_lineas,
                      subida_real_pedidos, subida_reg_pedidos, importe_fantasma)
    return metricas


def escribir_verdades(ruta: Path, m: dict, cuentas_v4: list[dict],
                      pares: list[dict], desviacion_total: Decimal,
                      desviacion_maxima: Decimal, facturado_total: Decimal,
                      media_hosteleria: Decimal, tickets_v1: list[Ticket],
                      tickets_fact_otro: list[Ticket], tickets_v5: list[Ticket],
                      conteo_sucias: dict, bucket_facturacion: list[Ticket],
                      pureza: float, estimacion_ingenua: float,
                      lineas_real: dict, lineas_reg: dict, base_lineas: float,
                      sub_real_l: float, sub_reg_l: float,
                      sub_real_p: float, sub_reg_p: float,
                      importe_fantasma: Decimal) -> None:
    """
    Clave de corrección del ejercicio central del bloque 4.
    Para cada verdad: enunciado, cifras exactas de esta ejecución, camino de
    detección, errores típicos y el principio de CX al que engancha.
    ESPECIFICACION.md §4: no son acertijos, son ejemplos de la disciplina.
    """
    v1, v2, v3, v4, v5 = m["v1"], m["v2"], m["v3"], m["v4"], m["v5"]
    anual = desviacion_total * Decimal(2)

    filas_dup = "\n".join(
        f"| {p['original']} | {p['nombre_original']} | {p['duplicado']} | "
        f"{p['nombre_duplicado']} | {p['telefono_norm']} | "
        f"{p['dto_original']} / {p['dto_duplicado']} |"
        for p in pares)

    filas_mes = "\n".join(
        f"| {clave} | {lineas_real[tuple(int(x) for x in clave.split('-'))]} | "
        f"{lineas_reg[tuple(int(x) for x in clave.split('-'))]} |"
        for clave in [f"{a}-{mm:02d}" for (a, mm) in sorted(lineas_real.keys())])

    filas_v4 = "\n".join(
        f"| {c['id_cliente']} | {c['nombre']} | {c['descuento_pct']} % | "
        f"{c['pedidos']} | {euros(c['facturado'])} | {euros(c['margen_bruto'])} | "
        f"{c['tickets']} | {euros(c['coste_contactos'])} | "
        f"**{euros(c['resultado'])}** |"
        for c in cuentas_v4)

    top_sucias = sorted(conteo_sucias.items(), key=lambda kv: (-kv[1], kv[0]))[:8]
    filas_sucias = "\n".join(f"| `{k}` | {v} | {pct(100 * v / N_TICKETS)} |"
                             for k, v in top_sucias)

    texto = f"""# Las cinco verdades escondidas — clave de corrección

> **No se le enseña a la alumna.** Es la clave de corrección del ejercicio
> central del bloque 4 (`b4-m10-analisis-completo`) y de los ejercicios del
> bloque 2 que corren sobre el mismo dataset. Se destapa después del intento,
> nunca antes (ESPECIFICACION.md §5.3).
>
> Este fichero **se genera**, no se escribe a mano: lo produce
> `scripts/generar-dataset.py` en la misma ejecución que los datos, para que no
> pueda desfasarse. Todas las cifras son las de esta ejecución, con semilla
> {SEMILLA}.

ESPECIFICACION.md §4 lo dice sin rodeos: **no son acertijos, son ejemplos de la
disciplina**. Cada verdad existe porque detrás hay un principio de CX que se
puede enseñar con ella y no sin ella.

## Rúbrica de tres niveles (§5.3)

| Nivel | Qué significa |
|---|---|
| No llegó | No identificó la concentración, o la atribuyó a la columna `categoria` |
| Llegó | Identificó la causa raíz y la cuantificó con un margen razonable |
| Llegó y encontró algo | Además cruzó dos ficheros por su cuenta o cuestionó un dato |

---

## Antes de las verdades: por qué `categoria` no sirve

La columna `categoria` de `tickets.xlsx` trae
{m['tickets']['etiquetas_distintas'] - 1} etiquetas distintas más un cajón de
`Otros` y un puñado de celdas vacías, para 8 categorías reales. El reparto de
las más frecuentes:

| Etiqueta | Tickets | % |
|---|---:|---:|
{filas_sucias}

Quien agrupe por esa columna y sume las etiquetas de facturación obtiene
**{len(bucket_facturacion)} tickets = {pct(100 * estimacion_ingenua)}**, de los
cuales solo el {pct(100 * pureza)} son de verdad del fallo de facturación. La
respuesta correcta es {pct(100 * v1['tickets_pct'])}. **Fiarse de la columna
cuesta {100 * m['tickets']['error_de_fiarse_de_categoria']:.1f} puntos de
error**, en la dirección peligrosa: infravalora el problema.

Ése es el puente del verbo 1 (clasificar) al bloque 4: hay que reclasificar
desde `descripcion`, no desde `categoria`.

---

## V1 — El fallo de redondeo en pedidos con descuento

### 1. Enunciado

El sistema de facturación aplica el descuento **al precio unitario**, redondea
ese precio unitario **al alza** al céntimo, y sólo después multiplica por las
unidades. El error se multiplica por la cantidad. El manual (PR-06) describe el
cálculo correcto: descuento sobre el importe de la línea y redondeo al final.
Nadie ha leído el manual. Resultado: **{pct(100 * v1['tickets_pct'])} de los
tickets** son un cliente diciendo que su factura no cuadra.

### 2. Cifras exactas de esta ejecución

| Dato | Valor |
|---|---|
| Líneas de pedido con descuento | {v1['lineas_con_descuento']} |
| De ellas, con sobrecoste | {v1['lineas_con_descuento_y_desviacion']} ({pct(100 * v1['pct_lineas_dto_con_desviacion'])}) |
| Líneas sin descuento con sobrecoste | 0 (invariante dura) |
| Pedidos afectados | {v1['pedidos_afectados']} |
| **Sobrefacturación acumulada en 6 meses** | **{euros(desviacion_total)}** |
| Proyección a 12 meses | {euros(anual)} |
| Sobrecoste máximo en una sola línea | {euros(desviacion_maxima)} |
| Facturación total del periodo | {euros(facturado_total)} |
| Tickets de la causa | {v1['tickets']} = {pct(100 * v1['tickets_pct'])} |
| De ellos, con `id_pedido` informado | {v1['tickets_con_id_pedido']} ({pct(100 * v1['tickets_con_id_pedido_pct'])}) |
| Coste de atender esos tickets a 11,00 €/contacto | {euros(Decimal(v1['coste_contactos']))} |

> El dinero mal facturado ({euros(desviacion_total)}) es **calderilla** al lado
> del coste de atender las llamadas que provoca
> ({euros(Decimal(v1['coste_contactos']))}). Ése es el hallazgo de verdad, y es
> el que hay que llevar a Gerencia.

**Trampa deliberada:** hay {len(tickets_fact_otro)} tickets más que también son
de facturación y **no** vienen de este fallo (factura duplicada, IVA mal, precio
mal pactado). Quien concluya "el 100 % de las incidencias de facturación es el
mismo fallo" se equivoca. El hallazgo es una **concentración**, no una
totalidad.

### 3. Camino de detección esperado

1. Reclasificar los 800 tickets desde `descripcion`, no desde `categoria`
   (verbo 1 sobre texto libre).
2. Ver que la clase mayoritaria es "la factura no cuadra" y que pesa ~38 %.
3. Coger los tickets de esa clase que llevan `id_pedido` y buscar esos pedidos
   en `pedidos.xlsx`.
4. Observar que **todos** tienen `dto_pct > 0`. Ésa es la pista: el fallo va con
   el descuento.
5. Recalcular la línea a mano: `precio_ud × cantidad × (1 − dto/100)`, redondear
   al final, y comparar con `importe_linea`.
6. Comprobar el signo: la diferencia es **siempre a favor de la empresa**. Un
   error aleatorio no tiene signo.
7. Comprobar el contraejemplo: en las líneas con `dto_pct = 0` la diferencia es
   exactamente cero. Eso confirma la hipótesis y descarta un error de tarifa.
8. Abrir `procedimientos.docx`, leer PR-06 y ver que el manual describe el
   cálculo correcto. El sistema no hace lo que dice el procedimiento.
9. Multiplicar: sobrefacturación del periodo, y coste de los contactos que
   genera.

### 4. Errores típicos

- **Fiarse de la columna `categoria`.** Da {pct(100 * estimacion_ingenua)} en
  vez de {pct(100 * v1['tickets_pct'])}: {100 * m['tickets']['error_de_fiarse_de_categoria']:.1f}
  puntos por debajo. Es el error más frecuente y el más caro.
- **Quedarse en el dinero mal facturado.** {euros(desviacion_total)} en seis
  meses no mueve a nadie. El argumento está en el coste de los contactos.
- **Decir "el 100 % de facturación es este fallo"** e ignorar los
  {len(tickets_fact_otro)} tickets que no lo son.
- **Comparar totales de factura en vez de líneas.** El descuadre se diluye al
  sumar; hay que bajar a la línea.
- **Redondear al comparar.** Si se recalcula con dos decimales desde el
  principio, se reproduce el propio fallo y sale que todo cuadra.
- **Concluir "se equivocó alguien al teclear".** Es sistemático: afecta al
  {pct(100 * v1['pct_lineas_dto_con_desviacion'])} de las líneas con descuento y
  a ninguna sin él.

### 5. Principio de CX al que engancha

**Análisis de causa raíz** (`b4-m2-causa-raiz`) y **contacto evitable**
(`b4-m3-contacto-evitable`). Los {v1['tickets']} tickets no son demanda de
atención: son la factura de un error de cálculo. Ninguno de ellos debería
existir. Arreglar el cálculo elimina la causa; contestarlos mejor solo acelera
el síntoma. Enlaza también con `b4-m8-del-sintoma-al-sistema`.

---

## V2 — Los doce clientes duplicados

### 1. Enunciado

`clientes.xlsx` tiene {m['clientes']['filas']} filas pero
{m['clientes']['reales']} clientes: **{v2['pares']} son fichas duplicadas** de un
cliente que ya estaba. Cambian el nombre, el formato del teléfono y la tipografía
de la dirección, así que ningún cotejo literal los encuentra. PR-01 obliga a
comprobar el teléfono antes de dar de alta y PR-09 prohíbe crear ficha nueva por
un cambio de titular. Ninguno de los dos se cumple.

### 2. Cifras exactas de esta ejecución

| Dato | Valor |
|---|---|
| Filas en el maestro | {m['clientes']['filas']} |
| Clientes reales | {m['clientes']['reales']} |
| Pares duplicados | {v2['pares']} |
| Pares en los que **el descuento difiere** | {v2['pares_con_descuento_distinto']} |
| Pares en los que **ambas fichas tienen pedidos** | {v2['ambos_con_pedidos']} de {v2['pares']} |

| Original | Nombre original | Duplicado | Nombre duplicado | Teléfono normalizado | Dto. orig. / dup. |
|---|---|---|---|---|---|
{filas_dup}

### 3. Camino de detección esperado

1. Normalizar el teléfono: quitar espacios, guiones y el prefijo `+34`, quedarse
   con los 9 dígitos.
2. Normalizar la dirección: minúsculas, sin tildes, sin puntuación, desplegando
   `C/` → `calle` y `Avda.` → `avenida`.
3. Agrupar por la pareja (teléfono normalizado, dirección normalizada) y buscar
   grupos de más de una fila.
4. Comprobar que en cada grupo las dos fichas tienen pedidos: no es que una esté
   muerta, es que **se está facturando dos veces al mismo cliente**.
5. Mirar `descuento_pct` dentro de cada par. En
   {v2['pares_con_descuento_distinto']} de los {v2['pares']} no coincide.
6. Leer PR-01 y PR-09 y ver que el procedimiento ya lo prohíbe.

### 4. Errores típicos

- **Comparar nombres.** `Bar Manolo` y `BAR MANOLO S.L.` no se parecen para una
  comparación literal, y `Pepe, Casa` menos todavía.
- **Comparar teléfonos sin normalizar.** Los cuatro formatos conviven a
  propósito: `942 12 34 56`, `+34942123456`, `942123456`, `942-12-34-56`.
- **Normalizar solo el teléfono** y dar por duplicados casos que no lo son (o al
  revés, no cruzar la dirección y no poder demostrarlo).
- **Contar 12 duplicados y parar ahí.** El daño no es tener filas de más: es que
  {v2['pares_con_descuento_distinto']} clientes están comprando con dos
  condiciones distintas según la ficha por la que entre el pedido.
- **Olvidar el mojibake.** Algunas fichas traen la codificación rota
  (`JosÃ©`); si no se arregla antes de normalizar, el cotejo falla.

### 5. Principio de CX al que engancha

**Taxonomías y calidad del dato maestro** (`b4-m1-taxonomias`) y, de lleno, la
capa transversal de *causa vs. síntoma*: limpiar los 12 duplicados a mano es el
síntoma; el sistema es que el alta no comprueba nada. Alimenta también
`b4-m8-del-sintoma-al-sistema`.

---

## V3 — El pico de diciembre que no es demanda

### 1. Enunciado

Diciembre parece dispararse. Sube de verdad —es un distribuidor de bebidas y
llega la Navidad— pero **el volumen registrado sube casi cuatro veces más de lo
que subió el negocio**, porque la carga de pedidos del 9 al 13 de diciembre de
2024 **se importó dos veces**. Los {v3['pedidos_fantasma']} pedidos duplicados
llevan `id_pedido` de un bloque aparte (`PED-9xxxx`) que nadie miró, y **no
generan ni un solo ticket**: nunca se sirvieron.

### 2. Cifras exactas de esta ejecución

| Mes | Líneas reales | Líneas registradas |
|---|---:|---:|
{filas_mes}

| Dato | Valor |
|---|---|
| Media de líneas de los cinco meses que no son diciembre | {base_lineas:.1f} |
| Diciembre real | {lineas_real[(2024, 12)]} líneas → **{pct(100 * sub_real_l)}** |
| Diciembre registrado | {lineas_reg[(2024, 12)]} líneas → **{pct(100 * sub_reg_l)}** |
| Pedidos fantasma | {v3['pedidos_fantasma']} |
| Líneas fantasma | {v3['lineas_fantasma']} |
| Importe fantasma (ventas que nunca existieron) | {euros(importe_fantasma)} |
| Ventana de la doble importación | {v3['ventana'][0]} → {v3['ventana'][1]} |
| Mismo cálculo contando pedidos en vez de líneas | real {pct(100 * sub_real_p)} / registrado {pct(100 * sub_reg_p)} |

### 3. Camino de detección esperado

1. Agrupar `pedidos.xlsx` por mes y contar. Diciembre destaca.
2. Bajar a día: la anomalía se concentra en la semana del 9 al 13.
3. Agrupar por `(id_cliente, fecha, producto, cantidad)` y buscar grupos de
   tamaño 2. Aparecen {v3['pedidos_fantasma']} pedidos idénticos a otro.
4. Mirar los `id_pedido` de los duplicados: todos empiezan por `PED-9`. Es una
   numeración distinta, de otra carga.
5. Comprobar que ninguno de esos `id_pedido` aparece en `tickets.xlsx`. Pedidos
   que nadie reclamó nunca, ni para bien ni para mal.
6. Recalcular diciembre sin ellos. La estacionalidad existe, pero es cuatro
   veces menor de lo que aparentaba.

### 4. Errores típicos

- **Celebrar el pico.** Es lo primero que pasa: "diciembre va genial". Un dato
  que sube no se valida solo porque nos guste.
- **Buscar los duplicados por `id_pedido`.** Son distintos a propósito. Hay que
  agrupar por el contenido del pedido, no por su identificador.
- **Agrupar solo por cliente y fecha.** Un cliente puede pedir dos veces el
  mismo día en la vida real; hay que meter producto y cantidad en la clave.
- **Borrarlos y no preguntar por qué están.** La causa raíz es el proceso de
  importación, no las 62 filas.
- **No comprobar el contraste con tickets.** Es la prueba más limpia de que son
  fantasma: mercancía que nunca se sirvió no genera ni una queja.

### 5. Principio de CX al que engancha

**Detección de anomalías** (verbo 6 del bloque 2) y `b4-m5-metricas`: una
métrica que sube no es una buena noticia mientras no se sepa qué la mueve. Es el
ejemplo canónico de por qué se verifica el dato antes de construir encima.

---

## V4 — Los tres clientes que cuestan dinero

### 1. Enunciado

Tres clientes de hostelería concentran **{v4['tickets']} tickets =
{pct(100 * v4['tickets_pct'])}** de todas las incidencias. Los tres tienen
descuento alto y compran poco. Con un coste de contacto de
{euros(COSTE_CONTACTO)} (que ella deriva en `b4-m4-coste-de-un-contacto`), su
margen bruto de seis meses **no paga ni de lejos** lo que cuesta atenderlos.

### 2. Cifras exactas de esta ejecución

| Cliente | Nombre | Dto. | Pedidos | Facturado 6 m | Margen bruto 6 m | Tickets | Coste contactos | Resultado |
|---|---|---:|---:|---:|---:|---:|---:|---:|
{filas_v4}

| Dato | Valor |
|---|---|
| Tickets de los tres | {v4['tickets']} de {N_TICKETS} = {pct(100 * v4['tickets_pct'])} |
| Margen bruto conjunto | {euros(Decimal(v4['margen_total']))} |
| Coste de sus contactos | {euros(Decimal(v4['coste_total']))} |
| **Resultado conjunto** | **{euros(Decimal(v4['resultado_total']))}** |
| Facturación media de un cliente de hostelería (6 m) | {euros(media_hosteleria)} |

### 3. Camino de detección esperado

1. Contar tickets por `id_cliente`. Tres clientes se despegan del resto.
2. Sumar cuánto suponen sobre el total: {pct(100 * v4['tickets_pct'])}.
3. Ir a `pedidos.xlsx` y sumar lo facturado a esos tres en los seis meses.
4. Aplicar el margen por producto para obtener margen bruto, no facturación.
   Ojo: el descuento se come el margen, y estos tres tienen el descuento alto.
5. Multiplicar sus tickets por el coste de contacto de
   {euros(COSTE_CONTACTO)}.
6. Restar. Los tres salen en negativo.
7. Antes de proponer nada, mirar **por qué** llaman: sus tickets solapan con V1 y
   con V5. Parte de su coste no es culpa suya.

### 4. Errores típicos

- **Confundir facturación con margen.** Son los clientes con más descuento: la
  facturación engaña y el margen es la mitad de lo que parece.
- **Proponer echarlos.** Es la conclusión fácil y casi siempre está mal. Sus
  tickets solapan con V1 (facturas que no cuadran) y V5 (entregas sin avisar):
  buena parte de esas llamadas las provoca la propia empresa. Arreglado eso, el
  cliente puede dejar de ser deficitario.
- **Olvidar el `id_cliente` vacío.** {m['tickets']['sin_id_cliente']} tickets no
  lo llevan; hay que recuperar el cliente por `id_pedido` o aceptar el hueco y
  decirlo.
- **Contar tickets sin normalizar el cliente.** Si además hay duplicados (V2),
  la cuenta de tickets por cliente puede estar partida en dos fichas.
- **Dar la cifra sin el coste unitario.** "22 % de las incidencias" no es un
  argumento; "{euros(Decimal(v4['resultado_total']))} de pérdida en seis meses"
  sí.

### 5. Principio de CX al que engancha

**El coste real de un contacto** (`b4-m4`) y **escalado y excepciones**
(`b4-m7`). Es el nodo donde se ve que la rentabilidad de un cliente no está en
lo que compra sino en lo que cuesta servirle, y donde se enseña a diseñar el
camino del caso raro en vez de sufrirlo.

---

## V5 — PR-07, el procedimiento que nadie sigue

### 1. Enunciado

El manual tiene un procedimiento entero, **PR-07 — Aviso previo de entrega**,
que obliga a avisar al cliente el día antes con la franja horaria. No se cumple
nunca. El resultado es **{v5['tickets']} tickets =
{pct(100 * v5['tickets_pct'])}** de clientes de hostelería que estaban cerrados,
que no sabían que tocaba reparto, o que perdieron la entrega.

### 2. Cifras exactas de esta ejecución

| Dato | Valor |
|---|---|
| Tickets de la causa | {v5['tickets']} = {pct(100 * v5['tickets_pct'])} |
| ¿Todos de hostelería? | {"sí" if v5['todos_hosteleria'] else "no"} |
| Menciones a "aviso previo" en `tickets.xlsx` | 0 |
| Menciones a "aviso previo" en los 200 correos | 0 |
| Menciones en `procedimientos.docx` | PR-07 completo, 4 apartados |
| Coste de esos contactos a {euros(COSTE_CONTACTO)} | {euros(Decimal(v5['coste_contactos']))} |

### 3. Camino de detección esperado

1. Reclasificar desde `descripcion` y separar dos cosas que la columna
   `categoria` mezcla: **entrega con retraso** y **entrega sin avisar**. No son
   el mismo problema ni tienen la misma solución.
2. Contar los "sin avisar": {v5['tickets']} tickets, {pct(100 * v5['tickets_pct'])}.
3. Ver que **todos** son de hostelería. Un particular no pierde una entrega por
   no estar; un bar cerrado sí.
4. Leer `procedimientos.docx`. PR-07 existe, es explícito y describe justo lo
   que evitaría estos tickets.
5. Buscar en tickets y correos cualquier rastro de que ese aviso se dé alguna
   vez. **Cero.** El procedimiento existe en papel y no existe en la operación.
6. Cruzarlo con las contradicciones del propio manual (3 rutas frente a 4
   reales, horario 8:00–16:00 frente al 7:00–15:00 del que hablan los clientes,
   un "Departamento de Calidad" que en una empresa de 6 personas no existe):
   el manual lleva sin tocarse desde marzo de 2019.

### 4. Errores típicos

- **Meter "sin aviso" dentro de "retraso".** Es el error de taxonomía clásico y
  se lleva por delante todo el hallazgo: un retraso se arregla con logística, un
  aviso que falta se arregla con un mensaje el día antes.
- **Concluir que el procedimiento no existe.** Existe. El problema es otro:
  está escrito y nadie lo aplica. Son dos diagnósticos distintos y dos
  soluciones distintas.
- **Proponer "formar al equipo".** Sin preguntarse por qué un procedimiento de
  2019 nunca se ha aplicado. Casi nunca es que la gente no sepa.
- **No comprobar el perfil de cliente.** Que sean todos de hostelería es lo que
  permite acotar la solución y estimar su coste.
- **Tomarse el manual como verdad.** Contradice a los datos en cuatro puntos
  distintos. Es una fuente, no un oráculo.

### 5. Principio de CX al que engancha

**Contacto evitable** (`b4-m3`) y **del síntoma al sistema** (`b4-m8`), con
`b4-m7-escalado-y-excepciones` detrás. Es el caso más barato de arreglar de los
cinco y el que mejor enseña que la diferencia entre "documentado" y "hecho" es
donde vive la mayor parte del trabajo de CX.

---

## Resumen para el corrector

| Verdad | Cifra que tiene que salir | Tolerancia razonable |
|---|---|---|
| V1 | {pct(100 * v1['tickets_pct'])} de los tickets | ±2 puntos |
| V1 | {euros(desviacion_total)} sobrefacturados en 6 meses | ±5 % |
| V2 | {v2['pares']} pares duplicados | exacto |
| V3 | {v3['pedidos_fantasma']} pedidos fantasma | exacto |
| V3 | diciembre real {pct(100 * sub_real_l)} frente a {pct(100 * sub_reg_l)} registrado | ±3 puntos |
| V4 | {pct(100 * v4['tickets_pct'])} de las incidencias en 3 clientes | ±1,5 puntos |
| V4 | resultado conjunto {euros(Decimal(v4['resultado_total']))} | ±10 % |
| V5 | {pct(100 * v5['tickets_pct'])} de los tickets | ±1,5 puntos |

`scripts/verificar-verdades.py` reconstruye estas cinco verdades **sin leer este
fichero**, solo desde `dataset/ficheros/`. Si falla, el ejercicio central del
bloque 4 ha dejado de tener solución.
"""
    ruta.write_text(texto, encoding="utf-8", newline="\n")


# ── LEEME para la alumna ──────────────────────────────────────────────────────

def escribir_leeme(ruta: Path) -> None:
    texto = f"""# Los ficheros de Aguas del Norte

Todo lo que vas a analizar en este curso sale de aquí. No necesitas datos de tu
empresa en ningún momento, ni ahora ni después.

## De qué empresa son

**{EMPRESA}** es una distribuidora de agua envasada y bebidas.
Reparte a bares, restaurantes y hoteles, y también a particulares, en Cantabria
y el oriente de Asturias. Tiene unos 300 clientes en el maestro, seis empleados y cuatro rutas
de reparto. No existe: es una empresa inventada, con datos inventados, generada
por un script. Por eso puedes trastear con ella sin ningún problema de
privacidad.

Los ficheros cubren seis meses: del 2 de septiembre de 2024 al 28 de febrero de
2025.

## Qué hay en cada fichero

| Fichero | Qué es |
|---|---|
| `clientes.xlsx` | El maestro de clientes. Una fila por ficha: nombre, tipo, dirección, teléfono, ruta, descuento comercial y poco más. |
| `pedidos.xlsx` | El histórico de pedidos, **una fila por línea de pedido**, que es como lo lleva la gente en Excel. Producto, cantidad, precio, descuento, importe de la línea y total del pedido repetido en cada fila. |
| `tickets.xlsx` | Las 800 incidencias registradas en esos seis meses: canal, cliente, categoría, descripción escrita a mano, estado y tiempo dedicado. |
| `correos/` | 200 correos de clientes en formato `.eml`. Se abren con cualquier gestor de correo, y también con el Bloc de notas: por dentro son texto. |
| `bandeja.mbox` | Los mismos 200 correos en un solo fichero, por si prefieres abrirlos así. |
| `procedimientos.docx` | El manual interno "oficial" de la empresa. Nueve procedimientos, PR-01 a PR-09. |

## Aviso importante: están sucios a propósito

Esto no es un descuido. Los ficheros vienen **deliberadamente sucios**, con la
misma clase de porquería que tiene cualquier Excel que lleve años en una PyME:

- categorías escritas de mil maneras para decir bastantes menos cosas,
- fechas en tres formatos mezclados en la misma columna,
- teléfonos en cuatro formatos,
- acentos mal codificados (`JosÃ©` en vez de `José`),
- campos vacíos donde debería haber algo,
- códigos postales que a veces son texto y a veces número,
- nombres del mismo compañero escritos de cinco maneras,
- valores imposibles que alguien tecleó con prisa.

Si en algún momento piensas "esto está mal, ¿me lo habrán dado roto?": sí, y es
el ejercicio. Aprender a trabajar con datos limpios no sirve de nada, porque los
datos limpios no existen. La primera habilidad de todo este curso es mirar un
fichero y darte cuenta de qué no te puedes fiar.

## Cómo empezar

Ábrelos y míralos. Sin analizar nada todavía. Cuenta filas, mira qué columnas
hay, busca los huecos. Cuando el curso te pida algo concreto sobre ellos, ya
sabrás dónde está cada cosa.

> Estos ficheros se generan con `scripts/generar-dataset.py` a partir de una
> semilla fija. Si los tocas y quieres volver al punto de partida, se regeneran
> idénticos. No hace falta que lo hagas tú.
"""
    ruta.write_text(texto, encoding="utf-8", newline="\n")


def resumen(base: Path, m: dict) -> None:
    print(f"Gemelo sintético generado en {base}")
    print(f"  semilla {m['semilla']}  ·  periodo {m['periodo']['inicio']} → "
          f"{m['periodo']['fin']}  ·  {m['periodo']['dias_laborables']} días laborables")
    print()
    print("  clientes.xlsx        "
          f"{m['clientes']['filas']} filas "
          f"({m['clientes']['reales']} clientes reales + {m['clientes']['duplicados']} duplicados, "
          f"{m['clientes']['hosteleria']} hostelería / {m['clientes']['particular']} particulares)")
    print("  pedidos.xlsx         "
          f"{m['pedidos']['filas_xlsx']} filas de línea · "
          f"{m['pedidos']['pedidos_totales']} pedidos "
          f"({m['pedidos']['pedidos_reales']} reales + {m['pedidos']['pedidos_fantasma']} fantasma) · "
          f"{m['pedidos']['devoluciones']} devoluciones en negativo")
    print("  tickets.xlsx         "
          f"{m['tickets']['filas']} filas · "
          f"{m['tickets']['etiquetas_distintas']} etiquetas sucias · "
          f"{m['tickets']['con_id_pedido']} con id_pedido")
    print(f"  correos/             {m['correos']['total']} ficheros .eml")
    print("  bandeja.mbox         los mismos 200")
    print("  procedimientos.docx  9 procedimientos, v3 de marzo de 2019")
    print()
    print("  V1  redondeo    "
          f"{m['v1']['tickets']} tickets ({m['v1']['tickets_pct']:.1%}) · "
          f"{m['v1']['desviacion_total_reales']} € sobrefacturados")
    print(f"  V2  duplicados  {m['v2']['pares']} pares "
          f"({m['v2']['pares_con_descuento_distinto']} con descuento distinto)")
    print("  V3  diciembre   "
          f"{m['v3']['pedidos_fantasma']} pedidos fantasma · "
          f"real {m['v3']['subida_diciembre_real_lineas']:+.1%} vs "
          f"registrado {m['v3']['subida_diciembre_registrada_lineas']:+.1%}")
    print("  V4  tres cuentas "
          f"{m['v4']['tickets']} tickets ({m['v4']['tickets_pct']:.1%}) · "
          f"resultado {m['v4']['resultado_total']} €")
    print(f"  V5  PR-07       {m['v5']['tickets']} tickets ({m['v5']['tickets_pct']:.1%})")
    print()
    print("  columna `categoria`: fiarse de ella da "
          f"{m['tickets']['bucket_facturacion_pct']:.1%} en vez de "
          f"{m['v1']['tickets_pct']:.1%} "
          f"(error de {m['tickets']['error_de_fiarse_de_categoria']:.1%})")


def main(argv: list[str] | None = None) -> int:
    analizador = argparse.ArgumentParser(
        description="Genera el gemelo sintético de Aguas del Norte, S.L.")
    analizador.add_argument("--salida", default=str(RAIZ / "dataset"),
                            help="Directorio base con ficheros/ y SOLUCIONES/")
    analizador.add_argument("--silencioso", action="store_true",
                            help="No imprime el resumen")
    args = analizador.parse_args(argv)
    construir(Path(args.salida), silencioso=args.silencioso)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
